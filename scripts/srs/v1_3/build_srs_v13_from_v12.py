from __future__ import annotations

import argparse
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parents[3]
SOURCE_DOCX = (
    BASE
    / "output"
    / "doc"
    / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.2.docx"
)
TARGET_DOCX = (
    BASE
    / "output"
    / "doc"
    / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.3.docx"
)
DIAGRAM_DIR = BASE / "docs" / "source" / "diagrams" / "rendered" / "v1_3"


@dataclass(frozen=True)
class FigureReplacement:
    caption: str
    image_name: str
    new_caption: str | None = None
    width_inches: float = 5.6


def normalize_text(text: str) -> str:
    return " ".join(text.replace("\u3000", " ").replace("\u00A0", " ").split())


def find_paragraph_exact(document: DocumentObject, text: str) -> Paragraph:
    needle = normalize_text(text)
    for paragraph in document.paragraphs:
        if normalize_text(paragraph.text) == needle:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def paragraph_index(document: DocumentObject, target: Paragraph) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError(f"Paragraph not found in document: {target.text!r}")


def get_previous_paragraph(document: DocumentObject, paragraph: Paragraph) -> Paragraph | None:
    index = paragraph_index(document, paragraph)
    if index == 0:
        return None
    return document.paragraphs[index - 1]


def has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def ensure_paragraph_properties(paragraph: Paragraph) -> OxmlElement:
    paragraph_properties = paragraph._p.find(qn("w:pPr"))
    if paragraph_properties is None:
        paragraph_properties = OxmlElement("w:pPr")
        paragraph._p.insert(0, paragraph_properties)
    return paragraph_properties


def set_keep_with_next(paragraph: Paragraph, value: bool = True) -> None:
    paragraph_properties = ensure_paragraph_properties(paragraph)
    existing = paragraph_properties.find(qn("w:keepNext"))
    if value:
        if existing is None:
            paragraph_properties.append(OxmlElement("w:keepNext"))
    elif existing is not None:
        paragraph_properties.remove(existing)


def set_keep_together(paragraph: Paragraph, value: bool = True) -> None:
    paragraph_properties = ensure_paragraph_properties(paragraph)
    existing = paragraph_properties.find(qn("w:keepLines"))
    if value:
        if existing is None:
            paragraph_properties.append(OxmlElement("w:keepLines"))
    elif existing is not None:
        paragraph_properties.remove(existing)


def set_page_break_before(paragraph: Paragraph, value: bool = True) -> None:
    paragraph_properties = ensure_paragraph_properties(paragraph)
    existing = paragraph_properties.find(qn("w:pageBreakBefore"))
    if value:
        if existing is None:
            paragraph_properties.append(OxmlElement("w:pageBreakBefore"))
    elif existing is not None:
        paragraph_properties.remove(existing)


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", center: bool = False) -> Paragraph:
    new_paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if center:
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_paragraph


def add_picture_to_paragraph(paragraph: Paragraph, image_path: Path, width_inches: float) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_together(paragraph, True)
    set_keep_with_next(paragraph, True)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def set_caption_text(paragraph: Paragraph, caption: str) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_together(paragraph, True)
    paragraph.add_run(caption)


def insert_figure_after(paragraph: Paragraph, image_path: Path, caption: str, width_inches: float) -> Paragraph:
    picture_paragraph = insert_paragraph_after(paragraph, center=True)
    add_picture_to_paragraph(picture_paragraph, image_path, width_inches)
    caption_paragraph = insert_paragraph_after(picture_paragraph, center=True)
    set_caption_text(caption_paragraph, caption)
    return caption_paragraph


def replace_figure(document: DocumentObject, spec: FigureReplacement) -> str:
    caption_paragraph = find_paragraph_exact(document, spec.caption)
    picture_paragraph = get_previous_paragraph(document, caption_paragraph)
    if picture_paragraph is None or not has_drawing(picture_paragraph):
        raise ValueError(f"Image paragraph missing before caption: {spec.caption}")

    image_path = DIAGRAM_DIR / spec.image_name
    if not image_path.exists():
        raise FileNotFoundError(f"Missing image: {image_path}")

    add_picture_to_paragraph(picture_paragraph, image_path, spec.width_inches)
    set_caption_text(caption_paragraph, spec.new_caption or spec.caption)
    return f"replaced {spec.caption}"


def find_table_by_header(document: DocumentObject, headers: tuple[str, ...]):
    needle = tuple(normalize_text(item) for item in headers)
    for table in document.tables:
        if not table.rows:
            continue
        row = tuple(normalize_text(cell.text) for cell in table.rows[0].cells)
        if row == needle:
            return table
    raise ValueError(f"Table not found for header: {headers}")


def update_analysis_summary_table(document: DocumentObject) -> None:
    table = find_table_by_header(document, ("实体类别", "代表实体", "说明"))
    rows = (
        ("主数据", "StudentProfile", "承载学号、姓名、年级、班级和政治面貌等基础信息。"),
        ("知识治理", "KnowledgeEntry / KnowledgeSource / TemplateAsset", "支撑政策问答、来源版本和模板下载。"),
        ("党团流程", "PartyMemberStatus / PartyWorkflowEvent", "描述阶段状态、关键事件和时间线。"),
        ("申请审批", "CommonRequest / CommonRequestAttachment / ApprovalTask / ApprovalAction", "承载申请单、附件、审批任务和动作留痕。"),
        ("通知管理", "NoticeMessage / NoticeDelivery", "支撑通知圈选、分发和接收记录。"),
        ("学业分析", "CurriculumRuleSet / TermCourseOffering / AcademicGapResult", "支撑培养方案比对、缺口提示和课程类型建议。"),
        ("导入导出与审计", "ImportBatch / DocumentAuditLog", "承载批量交换、错误处理和可追溯日志。"),
    )
    for row_index, values in enumerate(rows, start=1):
        while len(table.rows) <= row_index:
            table.add_row()
        for col_index, value in enumerate(values):
            table.rows[row_index].cells[col_index].text = value


def update_analysis_section(document: DocumentObject) -> list[str]:
    logs: list[str] = []

    class_caption = find_paragraph_exact(document, "图 3-3 核心业务分析类图")
    class_picture = get_previous_paragraph(document, class_caption)
    if class_picture is None or not has_drawing(class_picture):
        raise ValueError("Missing picture paragraph for 图 3-3")
    add_picture_to_paragraph(class_picture, DIAGRAM_DIR / "fig-3-3-core-analysis-class.png", 5.7)
    set_caption_text(class_caption, "图 3-3 核心业务分析类图")
    update_analysis_summary_table(document)
    logs.append("replaced 图 3-3")

    intro = find_paragraph_exact(document, "关键动态交互采用请假申请与后台审核作为典型时序进行说明。")
    clear_paragraph_content(intro)
    intro.add_run("以下进一步按五个模块补充代表性动态模型，分别覆盖知识问答、党团流程、通知推送、电子证明审批和学业分析场景。")
    set_page_break_before(intro, True)

    knowledge_caption = find_paragraph_exact(document, "图 3-4 学生提交请假申请至后台审核时序图")
    knowledge_picture = get_previous_paragraph(document, knowledge_caption)
    if knowledge_picture is None or not has_drawing(knowledge_picture):
        raise ValueError("Missing picture paragraph for 图 3-4")

    add_picture_to_paragraph(knowledge_picture, DIAGRAM_DIR / "fig-3-4-knowledge-qa-sequence.png", 5.6)
    set_caption_text(knowledge_caption, "图 3-4 知识问答与模板下载时序图")
    logs.append("replaced 图 3-4")

    cursor = knowledge_caption
    for caption, image_name, width in (
        ("图 3-5 党团事务流程状态图", "fig-3-5-party-workflow-state.png", 5.2),
        ("图 3-6 官方信息汇聚与精准推送时序图", "fig-3-6-notice-push-sequence.png", 5.6),
        ("图 3-7 电子证明生成与审批时序图", "fig-3-7-document-approval-sequence.png", 5.55),
        ("图 3-8 学业分析与预警活动图", "fig-3-8-academic-warning-activity.png", 4.0),
    ):
        cursor = insert_figure_after(cursor, DIAGRAM_DIR / image_name, caption, width)
        logs.append(f"inserted {caption}")

    return logs


def build_v13(overwrite_target: bool) -> list[str]:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"Missing source docx: {SOURCE_DOCX}")
    if TARGET_DOCX.exists() and not overwrite_target:
        raise FileExistsError(f"Target exists: {TARGET_DOCX}")

    TARGET_DOCX.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)

    document = Document(TARGET_DOCX)
    logs: list[str] = []
    for spec in (
        FigureReplacement("图 2-1 软件产品与外部环境关系图", "fig-2-1-system-context.png"),
        FigureReplacement("图 3-1 学生侧全局用例图", "fig-3-1-student-usecase.png", "图 3-1 学生侧五模块用例图"),
        FigureReplacement("图 3-2 管理侧全局用例图", "fig-3-2-admin-usecase.png", "图 3-2 管理侧五模块用例图"),
    ):
        logs.append(replace_figure(document, spec))

    logs.extend(update_analysis_section(document))
    document.save(TARGET_DOCX)
    return logs


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Incrementally derive SRS v1.3 from the manually edited v1.2 document.")
    parser.add_argument("--overwrite-target", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    logs = build_v13(overwrite_target=args.overwrite_target)
    print(f"[OK] saved {TARGET_DOCX}")
    for line in logs:
        print(f"[OK] {line}")


if __name__ == "__main__":
    main()
