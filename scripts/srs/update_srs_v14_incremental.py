from __future__ import annotations

import argparse
import re
import shutil
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Literal

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parents[2]
DEFAULT_SOURCE_DOCX = (
    BASE
    / "output"
    / "doc"
    / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.3.docx"
)
DEFAULT_TARGET_DOCX = (
    BASE
    / "output"
    / "doc"
    / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.4.docx"
)
DEFAULT_DIAGRAM_DIR = BASE / "docs" / "source" / "diagrams" / "rendered" / "v1_4"

SPACE_RE = re.compile(r"\s+")


@dataclass(frozen=True)
class ParagraphSelector:
    text: str
    mode: Literal["exact", "prefix"] = "exact"


@dataclass(frozen=True)
class FigureSpec:
    caption: str
    image_name: str
    width_inches: float = 5.55
    page_break_before: bool = False


@dataclass(frozen=True)
class ReplaceFigureOperation:
    op_id: str
    caption_selectors: tuple[ParagraphSelector, ...]
    new_caption: str
    image_name: str
    width_inches: float = 5.55


@dataclass(frozen=True)
class TextUpdateOperation:
    op_id: str
    selectors: tuple[ParagraphSelector, ...]
    new_text: str
    center: bool = False


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ").replace("\u00A0", " ")
    return SPACE_RE.sub(" ", text).strip()


def paragraph_text_matches(paragraph: Paragraph, selector: ParagraphSelector) -> bool:
    paragraph_text = normalize_text(paragraph.text)
    needle = normalize_text(selector.text)
    if selector.mode == "exact":
        return paragraph_text == needle
    if selector.mode == "prefix":
        return paragraph_text.startswith(needle)
    raise ValueError(f"Unsupported selector mode: {selector.mode}")


def find_paragraphs(document: DocumentObject, selector: ParagraphSelector) -> list[Paragraph]:
    return [paragraph for paragraph in document.paragraphs if paragraph_text_matches(paragraph, selector)]


def find_first_paragraph(
    document: DocumentObject,
    selectors: Iterable[ParagraphSelector],
) -> Paragraph | None:
    for selector in selectors:
        matches = find_paragraphs(document, selector)
        if matches:
            return matches[0]
    return None


def paragraph_index(document: DocumentObject, target: Paragraph) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError(f"Paragraph not found in document: {target.text!r}")


def get_previous_paragraph(document: DocumentObject, paragraph: Paragraph) -> Paragraph | None:
    index = paragraph_index(document, paragraph)
    if index == 0:
        return None
    return document.paragraphs[index - 1]


def has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def ensure_paragraph_properties(paragraph: Paragraph) -> OxmlElement:
    paragraph_properties = paragraph._p.find(qn("w:pPr"))
    if paragraph_properties is None:
        paragraph_properties = OxmlElement("w:pPr")
        paragraph._p.insert(0, paragraph_properties)
    return paragraph_properties


def set_keep_with_next(paragraph: Paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_keep_together(paragraph: Paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_together = value


def set_page_break_before(paragraph: Paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.page_break_before = value


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    center: bool = False,
) -> Paragraph:
    new_paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if center:
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_paragraph


def insert_cloned_paragraph_before(
    anchor_paragraph: Paragraph,
    template_paragraph: Paragraph,
    text: str = "",
    *,
    bold: bool = False,
    center: bool = False,
) -> Paragraph:
    new_paragraph_element = deepcopy(template_paragraph._p)
    anchor_paragraph._p.addprevious(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, anchor_paragraph._parent)
    set_plain_text(new_paragraph, text, center=center, bold=bold)
    return new_paragraph


def clone_paragraph_after(
    paragraph: Paragraph,
    template_paragraph: Paragraph,
    text: str = "",
    *,
    bold: bool = False,
    center: bool = False,
) -> Paragraph:
    new_paragraph_element = deepcopy(template_paragraph._p)
    paragraph._p.addnext(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)
    set_plain_text(new_paragraph, text, center=center, bold=bold)
    return new_paragraph


def insert_cloned_table_before(
    anchor_paragraph: Paragraph,
    template_table: Table,
    rows: list[list[str]],
) -> Table:
    new_table_element = deepcopy(template_table._tbl)
    anchor_paragraph._p.addprevious(new_table_element)
    new_table = Table(new_table_element, template_table._parent)
    set_table_rows(new_table, rows)
    return new_table


def add_picture_to_paragraph(paragraph: Paragraph, image_path: Path, width_inches: float) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_together(paragraph, True)
    set_keep_with_next(paragraph, True)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def set_caption_text(paragraph: Paragraph, caption: str) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_together(paragraph, True)
    paragraph.add_run(caption)


def set_plain_text(paragraph: Paragraph, text: str, center: bool = False, bold: bool = False) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else None
    run = paragraph.add_run(text)
    run.bold = bold


def insert_page_break_after(paragraph: Paragraph) -> Paragraph:
    page_break_paragraph = insert_paragraph_after(paragraph)
    page_break_paragraph.paragraph_format.space_before = 0
    page_break_paragraph.paragraph_format.space_after = 0
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    return page_break_paragraph


def insert_figure_after(
    paragraph: Paragraph,
    image_path: Path,
    caption: str,
    width_inches: float,
    placeholder: bool,
    page_break_before: bool = False,
) -> Paragraph:
    cursor = paragraph
    if page_break_before:
        cursor = insert_page_break_after(cursor)

    picture_paragraph = insert_paragraph_after(cursor)
    if image_path.exists():
        add_picture_to_paragraph(picture_paragraph, image_path, width_inches)
    elif placeholder:
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_keep_together(picture_paragraph, True)
        set_keep_with_next(picture_paragraph, True)
        picture_paragraph.add_run(f"[待插图：{caption}]")
    else:
        raise FileNotFoundError(f"Image not found: {image_path}")

    caption_paragraph = insert_paragraph_after(picture_paragraph)
    set_caption_text(caption_paragraph, caption)
    return caption_paragraph


def remove_captioned_figure(document: DocumentObject, caption_paragraph: Paragraph) -> None:
    previous_paragraph = get_previous_paragraph(document, caption_paragraph)
    if previous_paragraph is not None and has_drawing(previous_paragraph):
        delete_paragraph(previous_paragraph)
    delete_paragraph(caption_paragraph)


def remove_paragraphs_by_selectors(document: DocumentObject, selectors: Iterable[ParagraphSelector]) -> int:
    removed_count = 0
    for selector in selectors:
        matches = list(find_paragraphs(document, selector))
        for paragraph in matches:
            if normalize_text(paragraph.text).startswith("图 "):
                remove_captioned_figure(document, paragraph)
            else:
                delete_paragraph(paragraph)
            removed_count += 1
    return removed_count


def copy_source_docx(source_docx: Path, target_docx: Path, overwrite_target: bool) -> None:
    if source_docx.resolve() == target_docx.resolve():
        raise ValueError("Source and target DOCX must be different.")
    if not source_docx.exists():
        raise FileNotFoundError(f"Source DOCX not found: {source_docx}")
    if target_docx.exists() and not overwrite_target:
        raise FileExistsError(
            f"Target DOCX already exists: {target_docx}. Use --overwrite-target to replace it."
        )
    target_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, target_docx)


def table_header_matches(table: Table, expected_header: tuple[str, ...]) -> bool:
    if not table.rows:
        return False
    actual = tuple(normalize_text(cell.text) for cell in table.rows[0].cells)
    return actual == tuple(normalize_text(cell) for cell in expected_header)


def find_table_by_header(document: DocumentObject, expected_header: tuple[str, ...]) -> Table | None:
    for table in document.tables:
        if table_header_matches(table, expected_header):
            return table
    return None


def set_table_rows(table: Table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)

    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        for cell_index, value in enumerate(row_values):
            row.cells[cell_index].text = value


def upsert_table_row(table: Table, key: str, row_values: list[str]) -> str:
    for row in table.rows[1:]:
        if normalize_text(row.cells[0].text) == normalize_text(key):
            for cell_index, value in enumerate(row_values):
                row.cells[cell_index].text = value
            return "updated"

    new_row = table.add_row()
    for cell_index, value in enumerate(row_values):
        new_row.cells[cell_index].text = value
    return "inserted"


TEXT_UPDATES: tuple[TextUpdateOperation, ...] = (
    TextUpdateOperation(
        op_id="environment-intro",
        selectors=(ParagraphSelector("全量需求范围完整覆盖原始需求文档的 5 大模块", mode="prefix"),),
        new_text=(
            "全量需求范围完整覆盖原始需求文档与补充需求中的 7 个模块，技术拆解主线扩展到知识、流程、审批、通知、学业分析、"
            "奖励荣誉展示、学生画像与审计留痕等闭环。"
        ),
    ),
    TextUpdateOperation(
        op_id="function-overview-intro",
        selectors=(ParagraphSelector("本节严格按照需求文档给出的五大功能模块展开", mode="prefix"),),
        new_text=(
            "本节严格按照需求文档与补充需求给出的七个功能模块展开，并将会议纪要中的边界、约束、默认方案和实现前提并入同一套描述。"
            "任何后续技术拆解都不得被理解为对原始七个模块的删减或替代。"
        ),
    ),
    TextUpdateOperation(
        op_id="support-capabilities-intro",
        selectors=(ParagraphSelector("为支撑上述五大模块", mode="prefix"),),
        new_text="为支撑上述七个模块，系统还必须提供以下共性支撑能力：",
    ),
    TextUpdateOperation(
        op_id="fr-intro",
        selectors=(ParagraphSelector("在开发拆解层面，上述五大模块进一步分解为下列 FR", mode="prefix"),),
        new_text="在开发拆解层面，上述七个模块进一步分解为下列 FR。支撑型 FR 不构成新的业务模块，而是服务于七个模块落地。",
    ),
    TextUpdateOperation(
        op_id="traceability-summary",
        selectors=(ParagraphSelector("需求文档中的五大模块已在正文中逐项展开", mode="prefix"),),
        new_text="需求文档及补充需求中的七个模块已在正文中逐项展开，且每个模块都映射到明确的 FR、验收条目和原型页面。",
    ),
    TextUpdateOperation(
        op_id="support-note",
        selectors=(ParagraphSelector("资料导入导出、角色权限、字段脱敏和审计日志属于全局支撑能力", mode="prefix"),),
        new_text="资料导入导出、角色权限、字段脱敏和审计日志属于全局支撑能力，服务于七个模块，不应被误读为替代业务模块。",
    ),
    TextUpdateOperation(
        op_id="use-case-intro",
        selectors=(ParagraphSelector("系统的主要参与者包括学生", mode="prefix"),),
        new_text=(
            "系统的主要参与者包括学生、班团骨干 / 团支书 / 党支部书记、辅导员 / 班主任 / 审批老师、团委老师、超级管理员和学院领导。"
            "为减少单图跨域连线和提升传统 SRS 阅读性，用例模型拆分为学生侧与管理侧两张全局用例图，其覆盖范围已扩展至补充需求中的七个业务模块。"
        ),
    ),
    TextUpdateOperation(
        op_id="representative-usecase-note",
        selectors=(ParagraphSelector("以下选取五个代表性用例", mode="prefix"),),
        new_text="以下选取七个代表性用例，分别对应七个模块的核心用户旅程。",
    ),
    TextUpdateOperation(
        op_id="analysis-intro",
        selectors=(
            ParagraphSelector("分析模型从静态结构和动态交互两个视角描述核心业务对象及其关系", mode="prefix"),
        ),
        new_text="分析模型从静态结构和动态交互两个视角描述核心业务对象及其关系，并确保七个模块在数据对象层与交互层均可追溯。",
    ),
    TextUpdateOperation(
        op_id="analysis-summary",
        selectors=(ParagraphSelector("分析模型与五大模块及其 FR 保持一致", mode="prefix"),),
        new_text="分析模型与七个模块及其 FR 保持一致，可直接作为数据库设计和接口拆分的基础。",
    ),
    TextUpdateOperation(
        op_id="schedule-summary",
        selectors=(ParagraphSelector("一期开发周期建议按 12 周执行，以知识库、流程、审批、通知、审计五个核心闭环为主线推进", mode="prefix"),),
        new_text=(
            "一期开发周期建议按 12 周执行，以知识库、流程、审批、通知、学业分析、奖励荣誉展示、学生画像与审计控制等闭环为主线推进，"
            "同时完整纳入原始需求及补充需求中的新增能力。"
        ),
    ),
    TextUpdateOperation(
        op_id="acceptance-intro",
        selectors=(ParagraphSelector("项目验收必须在 Kingbase 测试环境完成，并以五大核心模块的功能和非功能结果作为判定依据", mode="prefix"),),
        new_text="项目验收必须在 Kingbase 测试环境完成，并以七个核心模块的功能和非功能结果作为判定依据。",
    ),
    TextUpdateOperation(
        op_id="prototype-intro",
        selectors=(ParagraphSelector("本项目当前采用线框图级文本原型说明核心界面，并按《需求文档.md》的五大模块分别给出代表页面", mode="prefix"),),
        new_text="本项目当前采用线框图级文本原型说明核心界面，并按《需求文档.md》与补充需求中的七个模块分别给出代表页面。",
    ),
)


REPLACE_OPERATIONS: tuple[ReplaceFigureOperation, ...] = (
    ReplaceFigureOperation(
        op_id="fig-3-1",
        caption_selectors=(
            ParagraphSelector("图 3-1 学生侧七模块用例图"),
            ParagraphSelector("图 3-1 学生侧五模块用例图"),
            ParagraphSelector("图 3-1 学生侧全局用例图"),
        ),
        new_caption="图 3-1 学生侧七模块用例图",
        image_name="fig-3-1-student-usecase.png",
        width_inches=5.5,
    ),
    ReplaceFigureOperation(
        op_id="fig-3-2",
        caption_selectors=(
            ParagraphSelector("图 3-2 管理侧七模块用例图"),
            ParagraphSelector("图 3-2 管理侧五模块用例图"),
            ParagraphSelector("图 3-2 管理侧全局用例图"),
        ),
        new_caption="图 3-2 管理侧七模块用例图",
        image_name="fig-3-2-admin-usecase.png",
        width_inches=5.5,
    ),
)


ANALYSIS_ANCHOR = ParagraphSelector(
    "分析模型从静态结构和动态交互两个视角描述核心业务对象及其关系",
    mode="prefix",
)
ANALYSIS_CLEANUP = (
    ParagraphSelector("以下进一步按五个模块补充代表性动态模型", mode="prefix"),
    ParagraphSelector("以下进一步按七个模块补充代表性动态模型", mode="prefix"),
    ParagraphSelector("图 3-3 核心业务分析类图"),
    ParagraphSelector("图 3-4 奖励荣誉与学生画像扩展类图"),
    ParagraphSelector("图 3-5 知识问答与模板下载时序图"),
    ParagraphSelector("图 3-6 党团事务流程状态图"),
    ParagraphSelector("图 3-7 官方信息汇聚与精准推送时序图"),
    ParagraphSelector("图 3-8 电子证明生成与审批时序图"),
    ParagraphSelector("图 3-9 学业分析与预警活动图"),
    ParagraphSelector("图 3-10 荣誉榜单浏览与详情查看时序图"),
    ParagraphSelector("图 3-11 学生画像查看与纠错申诉时序图"),
    ParagraphSelector("图 3-11 学生画像查看与纠错 / 成长补录时序图"),
    ParagraphSelector("图 3-4 学生提交请假申请至后台审核时序图"),
    ParagraphSelector("图 3-3 党团事务流程状态图"),
    ParagraphSelector("图 3-5 官方信息汇聚与精准推送时序图"),
    ParagraphSelector("图 3-6 核心业务分析类图"),
    ParagraphSelector("图 3-7 电子证明生成与审批时序图"),
    ParagraphSelector("图 3-8 学业分析与预警活动图"),
)
ANALYSIS_FIGURES: tuple[FigureSpec, ...] = (
    FigureSpec("图 3-3 核心业务分析类图", "fig-3-3-core-analysis-class.png", 5.5),
    FigureSpec("图 3-4 奖励荣誉与学生画像扩展类图", "fig-3-4-extension-analysis-class.png", 5.5, True),
    FigureSpec("图 3-5 知识问答与模板下载时序图", "fig-3-5-knowledge-qa-sequence.png", 5.55, True),
    FigureSpec("图 3-6 党团事务流程状态图", "fig-3-6-party-workflow-state.png", 5.45),
    FigureSpec("图 3-7 官方信息汇聚与精准推送时序图", "fig-3-7-notice-push-sequence.png", 5.55),
    FigureSpec("图 3-8 电子证明生成与审批时序图", "fig-3-8-document-approval-sequence.png", 5.55),
    FigureSpec("图 3-9 学业分析与预警活动图", "fig-3-9-academic-warning-activity.png", 4.0),
    FigureSpec("图 3-10 荣誉榜单浏览与详情查看时序图", "fig-3-10-honor-display-sequence.png", 5.55, True),
    FigureSpec("图 3-11 学生画像查看与纠错 / 成长补录时序图", "fig-3-11-student-profile-sequence.png", 5.55, True),
)
POST_INSERT_CLEANUP = (
    ParagraphSelector("图 3-4 知识问答与模板下载时序图"),
    ParagraphSelector("图 3-5 党团事务流程状态图"),
    ParagraphSelector("图 3-6 官方信息汇聚与精准推送时序图"),
)

MODULE_TABLE_ROWS = (
    [
        "M-06",
        "奖励荣誉展示",
        "荣誉公示、榜样风采、分类筛选、荣誉数据维护",
        "学生、辅导员、学院领导、管理员",
        "FR-017",
    ],
    [
        "M-07",
        "学生画像与信息管理",
        "多维度信息聚合、动态档案更新、全景视图展示、画像数据治理",
        "学生（查看本人）、辅导员、班主任、学院领导",
        "FR-018",
    ],
)

FR_TABLE_ROWS = (
    [
        "FR-017",
        "奖励荣誉公示与榜样展示",
        "模块六",
        "系统应支持录入或导入校级及以上正式荣誉（含国奖、校优等），展示内容需包含事迹摘要与授予单位；管理员可对过期或撤销荣誉进行归档处理。",
    ],
    [
        "FR-018",
        "学生画像聚合与全景视图",
        "模块七",
        "系统应聚合学籍核心字段（学号、专业、年级）与动态成长字段（科研、竞赛、实践）；基础数据由教务同步；辅导员与院领导可按权限查看画像全貌，学生仅可查看本人信息。",
    ],
)

USE_CASE_SUMMARY_ROWS = (
    ["模块六", "浏览荣誉榜单并查看榜样事迹", "学生、辅导员、学院领导、管理员", "荣誉公示、分类筛选、事迹展示、荣誉数据维护"],
    ["模块七", "查看学生综合画像与成长档案", "学生、辅导员、班主任、学院领导", "多维度信息聚合、动态档案更新、全景视图、数据治理"],
)

ENTITY_TABLE_ROWS = (
    [
        "奖励荣誉",
        "HonorRecord / HonorCategory / HonorDisplayConfig",
        "承载荣誉名称、个人或集体展示主体、授予单位、文号、授权与审核状态、公示日期、事迹摘要、获奖感言、关联学生及展示有效期。",
    ],
    [
        "学生画像扩展",
        "StudentProfileExtension / ResearchExperience / CompetitionAward / SocialPractice",
        "承载科研项目、学科竞赛、社会实践、志愿服务、干部任职等动态成长字段，以及来源、录入人和最后更新时间。",
    ],
)

MODULE_OVERVIEW_BLOCKS = (
    (True, "模块六：奖励荣誉展示"),
    (False, "荣誉公示与风采展示：集中收录并展示校级及以上正式评定的学生个人或集体荣誉称号与奖项，涵盖国家奖学金、国家励志奖学金、校级优秀学生干部、优秀共青团员、先进班集体、学科竞赛国家级 / 省部级奖项等类别。"),
    (False, "榜样标杆树立：系统以图文结合、分类筛选的方式公开展示获奖者或获奖集体的基本信息、主要事迹摘要及获奖感言，突出先进典型的引领作用。展示内容须经学生本人授权及学院审核，确保隐私保护合规。"),
    (False, "规则与边界：仅展示经过官方红头文件或证书确认的正式荣誉，不收录口头表扬、临时奖励或非官方评选结果。展示范围默认仅限在读学生，已毕业学生荣誉原则上不再增补展示，历史数据保留查阅入口但默认不主动推送。"),
    (False, "管理维护：管理员可依据最终公示名单批量导入或人工新增荣誉记录，支持关联学生学籍库自动核验身份。荣誉条目必须标注授予单位、文号或证书编号、公示日期及有效期，并支持后续撤下或归档操作。"),
    (True, "模块七：学生画像与信息管理"),
    (False, "多维度信息整合：基于学籍核心数据构建学生基础档案，覆盖学号、姓名、专业、年级、班级等静态信息，并扩展收录科研项目参与、学科竞赛获奖、社会实践经历、志愿服务时长、学生干部任职等动态成长数据，形成学生个人成长画像。"),
    (False, "数据来源与更新：基础学籍信息由教务系统或学院管理员批量同步导入；扩展字段支持辅导员或学生本人经审核后补录更新。所有非学籍字段均需标注数据来源、录入人与最后更新时间。"),
    (False, "画像展示与应用：在管理端提供按维度筛选、检索和导出功能，便于辅导员、学院领导快速了解学生综合表现。学生端可查看本人已收录的画像信息，并有权对非结构化描述字段提出修正申诉。"),
    (False, "责任边界与安全约束：本模块仅作为信息聚合与展示工具，不包含对学生综合素质的自动评分、排名或评价结论输出。画像内容涉及科研未公开成果、竞赛未获奖记录或个人敏感信息的，严格遵循最小必要原则设置查看权限，敏感字段默认对非授权角色隐藏。所有数据操作全程留痕，确保可审计、可追溯。"),
)

USECASE_INSERTIONS = (
    (
        "代表用例六：浏览荣誉榜单并查看榜样事迹",
        [
            ["项目", "内容"],
            ["用例名称", "浏览荣誉榜单并查看榜样事迹"],
            ["主要参与者", "学生"],
            ["次要参与者", "辅导员 / 管理员"],
            ["前置条件", "管理员已导入或录入经学生本人授权（如适用）且经学院审核通过的校级及以上正式荣誉数据；荣誉条目已关联学生学籍信息或集体主体信息并完成身份核验，且位于公示有效期内；学生已登录系统。"],
            ["后置条件", "学生成功查看荣誉榜单及榜样事迹详情；系统记录访问行为用于展示热度统计。"],
            ["相关需求", "FR-017"],
        ],
        [
            ["步骤", "参与者", "系统响应"],
            ["1", "学生", "进入荣誉展示模块首页。"],
            ["2", "系统", "展示荣誉榜单，支持按类别和学年筛选。"],
            ["3", "学生", "点击某一荣誉条目。"],
            ["4", "系统", "展示获奖者或集体基本信息、荣誉名称、授予单位、事迹摘要及获奖感言（如有）。"],
            ["5", "学生", "浏览完毕，返回或继续查看。"],
        ],
        [
            ["编号", "场景", "处理方式"],
            ["K1", "荣誉已过期或归档", "标注“历史荣誉”，不再主动推送。"],
            ["K2", "获奖者已毕业", "保留荣誉展示，隐藏个人联系方式。"],
        ],
    ),
    (
        "代表用例七：查看学生综合画像与成长档案",
        [
            ["项目", "内容"],
            ["用例名称", "查看学生综合画像与成长档案"],
            ["主要参与者", "辅导员"],
            ["次要参与者", "学生 / 学院领导"],
            ["前置条件", "学籍信息已同步，扩展字段已录入；辅导员具备所带班级查看权限。"],
            ["后置条件", "辅导员完成画像查阅；系统记录操作日志；学生本人可通过独立入口发起信息纠错或成长补录申请。"],
            ["相关需求", "FR-018"],
        ],
        [
            ["步骤", "参与者", "系统响应"],
            ["1", "辅导员", "搜索并进入目标学生画像页。"],
            ["2", "系统", "聚合展示学籍静态字段与科研、竞赛、实践等动态字段。"],
            ["3", "辅导员", "浏览各维度信息，敏感字段默认脱敏。"],
            ["4", "辅导员", "结束查看，日志留痕。"],
            ["5", "学生", "在本人端查看画像并进入“信息纠错 / 成长补录”入口。"],
            ["6", "系统", "隐藏管理元数据，记录申诉说明或补录内容并提交至辅导员审核处理。"],
        ],
        [
            ["编号", "场景", "处理方式"],
            ["K1", "越权查看非管辖学生", "拒绝访问并记录越权尝试。"],
            ["K2", "学生本人查看画像或提交补录申请", "隐藏管理元数据，提供纠错申诉与成长补录入口。"],
        ],
    ),
)

ACCEPTANCE_PARAGRAPHS = (
    "模块六验收：成功导入不少于 30 条校级及以上正式荣誉记录，覆盖国奖、校优等至少 3 种类别；随机抽检 10 条荣誉条目，可正确展示获奖者姓名、专业年级、荣誉名称、授予单位、文号或证书编号、公示日期及事迹摘要；荣誉榜单支持按类别和学年筛选，筛选结果准确无误；已过期或归档荣誉默认不主动展示，历史查看页面标注“历史荣誉”字样；获奖者已毕业时，个人联系方式等敏感信息不得展示。",
    "模块七验收：预置不少于 50 名学生的学籍主数据后，成功补录科研项目、学科竞赛、社会实践等扩展字段各不少于 20 条；辅导员登录后可正确查看所带班级学生的全景画像，学籍静态字段与动态成长字段聚合展示无错漏；扩展字段均标注数据来源与最后更新时间；辅导员越权查看非管辖学生时系统拒绝访问并记录日志；学生本人登录仅可查看本人画像，管理元数据（如录入人、来源系统）不予展示，纠错申诉入口可正常使用。",
)

PROTOTYPE_BLOCKS = (
    (True, False, "5.6 奖励荣誉展示页面"),
    (False, False, "页面顶部展示荣誉榜单标题及分类筛选栏，支持按“国家奖学金”“校级优秀学生”“学科竞赛”“全部”等类别切换，并提供学年下拉筛选器。"),
    (False, False, "主内容区以卡片或图文列表形式展示荣誉条目，每条包含获奖者姓名、专业年级、荣誉名称、授予单位、公示日期及缩略事迹摘要。卡片右上角标注荣誉类别标签。"),
    (False, False, "点击任一荣誉卡片进入详情页，展示获奖者基本信息、荣誉名称、授予单位、文号或证书编号、完整事迹描述及感言（如有）。页面底部提供“返回榜单”入口。"),
    (False, False, "历史荣誉入口位于页面侧边或底部，点击后进入归档荣誉列表，每条均标注“历史荣誉，仅供参考”字样。获奖者已毕业时，详情页不展示联系方式等个人敏感字段。"),
    (False, False, "管理端对应页面支持荣誉条目批量导入、单条录入、分类维护、展示有效期设置及归档 / 撤销操作，并记录维护人与更新时间。"),
    (True, False, "5.7 学生画像与信息管理页面"),
    (False, True, "辅导员 / 管理端视图"),
    (False, False, "页面顶部提供学生搜索框，支持按学号、姓名、专业、年级或班级快速定位目标学生。"),
    (False, False, "进入画像详情后，页面左侧或上方展示学生身份摘要卡，包含学号、姓名、专业、年级、班级、政治面貌等静态字段及学生证件照位。"),
    (False, False, "主内容区按维度卡片聚合展示动态成长信息，包括科研项目参与、学科竞赛获奖、社会实践经历、志愿服务时长、学生干部任职等。每项记录均标注数据来源、录入人与最后更新时间。"),
    (False, False, "敏感字段（如家庭联系方式、身份证号）默认部分脱敏展示，并提供“申请查看完整信息”入口，点击后触发审批留痕。"),
    (False, False, "页面底部提供“导出画像快照”“添加成长记录”“信息纠错”等操作入口。"),
    (False, True, "学生端视图"),
    (False, False, "学生登录后仅可查看本人画像。页面展示学籍静态信息及已收录的动态成长记录，不显示数据来源、录入人等管理元数据。"),
    (False, False, "页面底部提供“信息纠错申诉”按钮，点击后可填写申诉说明并提交至辅导员审核处理。"),
)


def apply_text_updates(document: DocumentObject) -> list[str]:
    logs: list[str] = []
    for operation in TEXT_UPDATES:
        paragraph = find_first_paragraph(document, operation.selectors)
        if paragraph is None:
            logs.append(f"[MISS] {operation.op_id}: paragraph anchor not found")
            continue
        set_plain_text(paragraph, operation.new_text, center=operation.center)
        logs.append(f"[OK] {operation.op_id}: updated paragraph {paragraph_index(document, paragraph)}")
    return logs


def apply_traceability_table_updates(document: DocumentObject) -> list[str]:
    logs: list[str] = []

    modules_table = find_table_by_header(document, ("模块编号", "模块名称", "主要功能点", "主要用户", "对应 FR"))
    if modules_table is None:
        logs.append("[MISS] module-table: header not found")
    else:
        for row in MODULE_TABLE_ROWS:
            result = upsert_table_row(modules_table, row[0], row)
            logs.append(f"[OK] module-table: {result} {row[0]}")

    fr_table = find_table_by_header(document, ("FR", "标题", "对应能力", "关键规则摘要"))
    if fr_table is None:
        logs.append("[MISS] fr-table: header not found")
    else:
        for row in FR_TABLE_ROWS:
            result = upsert_table_row(fr_table, row[0], row)
            logs.append(f"[OK] fr-table: {result} {row[0]}")

    usecase_table = find_table_by_header(document, ("模块", "代表用例", "主要参与者", "覆盖内容"))
    if usecase_table is None:
        logs.append("[MISS] usecase-summary-table: header not found")
    else:
        for row in USE_CASE_SUMMARY_ROWS:
            result = upsert_table_row(usecase_table, row[0], row)
            logs.append(f"[OK] usecase-summary-table: {result} {row[0]}")

    entity_table = find_table_by_header(document, ("实体类别", "代表实体", "说明"))
    if entity_table is None:
        logs.append("[MISS] analysis-entity-table: header not found")
    else:
        for row in ENTITY_TABLE_ROWS:
            result = upsert_table_row(entity_table, row[0], row)
            logs.append(f"[OK] analysis-entity-table: {result} {row[0]}")

    return logs


def apply_module_overview_expansion(document: DocumentObject) -> list[str]:
    logs: list[str] = []
    anchor = find_first_paragraph(
        document,
        (ParagraphSelector("责任边界：本模块只提供辅助展示、风险提示与人工复核建议", mode="prefix"),),
    )
    heading_template = find_first_paragraph(document, (ParagraphSelector("模块五：学业情况分析与预警"),))
    body_template = find_first_paragraph(document, (ParagraphSelector("培养方案比对：", mode="prefix"),))

    if anchor is None or heading_template is None or body_template is None:
        return ["[MISS] module-overview-expansion: anchor or template not found"]

    cursor = anchor
    inserted = 0
    for is_heading, text in MODULE_OVERVIEW_BLOCKS:
        template = heading_template if is_heading else body_template
        cursor = clone_paragraph_after(cursor, template, text)
        inserted += 1

    return [f"[OK] module-overview-expansion: inserted {inserted} paragraphs after paragraph {paragraph_index(document, anchor)}"]


def apply_usecase_expansion(document: DocumentObject) -> list[str]:
    logs: list[str] = []
    anchor = find_first_paragraph(document, (ParagraphSelector("3.3 软件需求的分析模型"),))
    heading_template = find_first_paragraph(document, (ParagraphSelector("代表用例五：上传成绩数据并查看学业缺口"),))
    main_flow_template = find_first_paragraph(document, (ParagraphSelector("主事件流："),))
    exception_template = find_first_paragraph(document, (ParagraphSelector("异常 / 备选分支："),))
    detail_table_template = find_table_by_header(document, ("项目", "内容"))
    flow_table_template = find_table_by_header(document, ("步骤", "参与者", "系统响应"))
    exception_table_template = find_table_by_header(document, ("编号", "场景", "处理方式"))

    if (
        anchor is None
        or heading_template is None
        or main_flow_template is None
        or exception_template is None
        or detail_table_template is None
        or flow_table_template is None
        or exception_table_template is None
    ):
        return ["[MISS] usecase-expansion: anchor or template not found"]

    inserted_blocks = 0
    for title, detail_rows, flow_rows, exception_rows in USECASE_INSERTIONS:
        insert_cloned_paragraph_before(anchor, heading_template, title)
        insert_cloned_table_before(anchor, detail_table_template, detail_rows)
        insert_cloned_paragraph_before(anchor, main_flow_template, "主事件流：")
        insert_cloned_table_before(anchor, flow_table_template, flow_rows)
        insert_cloned_paragraph_before(anchor, exception_template, "异常 / 备选分支：")
        insert_cloned_table_before(anchor, exception_table_template, exception_rows)
        inserted_blocks += 6

    logs.append(f"[OK] usecase-expansion: inserted {inserted_blocks} blocks before paragraph {paragraph_index(document, anchor)}")
    return logs


def apply_acceptance_expansion(document: DocumentObject) -> list[str]:
    logs: list[str] = []
    anchor = find_first_paragraph(document, (ParagraphSelector("模块五验收：", mode="prefix"),))
    template = anchor
    if anchor is None or template is None:
        return ["[MISS] acceptance-expansion: anchor not found"]

    cursor = anchor
    for text in ACCEPTANCE_PARAGRAPHS:
        cursor = clone_paragraph_after(cursor, template, text)

    logs.append(f"[OK] acceptance-expansion: inserted {len(ACCEPTANCE_PARAGRAPHS)} paragraphs after paragraph {paragraph_index(document, anchor)}")
    return logs


def apply_prototype_expansion(document: DocumentObject) -> list[str]:
    logs: list[str] = []
    anchor = find_first_paragraph(document, (ParagraphSelector("上述原型描述用于指导后续 UI 详细设计和高保真原型制作。"),))
    heading_template = find_first_paragraph(document, (ParagraphSelector("5.5 学业情况分析与预警页面"),))
    body_template = find_first_paragraph(document, (ParagraphSelector("页面顶部展示专业、年级、培养方案版本", mode="prefix"),))

    if anchor is None or heading_template is None or body_template is None:
        return ["[MISS] prototype-expansion: anchor or template not found"]

    inserted = 0
    for is_heading, is_bold, text in PROTOTYPE_BLOCKS:
        template = heading_template if is_heading else body_template
        insert_cloned_paragraph_before(anchor, template, text, bold=is_bold)
        inserted += 1

    logs.append(f"[OK] prototype-expansion: inserted {inserted} paragraphs before paragraph {paragraph_index(document, anchor)}")
    return logs


def apply_replace_operations(
    document: DocumentObject,
    diagram_dir: Path,
    placeholder_missing: bool,
) -> list[str]:
    logs: list[str] = []
    for operation in REPLACE_OPERATIONS:
        caption_paragraph = find_first_paragraph(document, operation.caption_selectors)
        if caption_paragraph is None:
            logs.append(f"[MISS] {operation.op_id}: caption anchor not found")
            continue

        image_path = diagram_dir / operation.image_name
        image_paragraph = get_previous_paragraph(document, caption_paragraph)
        if image_paragraph is None or not has_drawing(image_paragraph):
            logs.append(f"[MISS] {operation.op_id}: image paragraph not found before caption")
            continue

        if image_path.exists():
            add_picture_to_paragraph(image_paragraph, image_path, operation.width_inches)
            set_caption_text(caption_paragraph, operation.new_caption)
            logs.append(
                f"[OK] {operation.op_id}: replaced at paragraph {paragraph_index(document, caption_paragraph)}"
            )
            continue

        if placeholder_missing:
            clear_paragraph_content(image_paragraph)
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_keep_together(image_paragraph, True)
            set_keep_with_next(image_paragraph, True)
            image_paragraph.add_run(f"[待插图：{operation.new_caption}]")
            set_caption_text(caption_paragraph, operation.new_caption)
            logs.append(
                f"[HOLD] {operation.op_id}: image missing, inserted placeholder at paragraph {paragraph_index(document, caption_paragraph)}"
            )
        else:
            logs.append(f"[SKIP] {operation.op_id}: image missing {image_path.name}")
    return logs


def apply_analysis_group(
    document: DocumentObject,
    diagram_dir: Path,
    placeholder_missing: bool,
) -> list[str]:
    anchor_paragraph = find_first_paragraph(document, (ANALYSIS_ANCHOR,))
    if anchor_paragraph is None:
        return ["[MISS] analysis-group: anchor paragraph not found"]

    missing_images = [
        figure.image_name for figure in ANALYSIS_FIGURES if not (diagram_dir / figure.image_name).exists()
    ]
    if missing_images and not placeholder_missing:
        return [f"[SKIP] analysis-group: missing images {', '.join(missing_images)}"]

    removed_count = remove_paragraphs_by_selectors(document, ANALYSIS_CLEANUP)
    cursor = anchor_paragraph
    intro = insert_paragraph_after(
        cursor,
        "以下进一步按七个模块补充代表性静态与动态模型，先给出包含个人 / 集体荣誉、公示授权、身份核验和成长留痕字段的扩展类图，再给出知识问答、党团流程、通知推送、电子证明审批、学业分析、荣誉展示和学生画像场景的代表性动态图。",
    )
    cursor = intro
    for figure in ANALYSIS_FIGURES:
        cursor = insert_figure_after(
            cursor,
            diagram_dir / figure.image_name,
            figure.caption,
            figure.width_inches,
            placeholder_missing,
            figure.page_break_before,
        )
    cursor = insert_page_break_after(cursor)
    post_removed = remove_paragraphs_by_selectors(document, POST_INSERT_CLEANUP)
    return [
        f"[OK] analysis-group: inserted {len(ANALYSIS_FIGURES)} figures after paragraph {paragraph_index(document, anchor_paragraph)}; cleaned {removed_count} paragraphs; post-cleaned {post_removed} duplicate paragraphs"
    ]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Incrementally derive SRS v1.4 from the manually refined v1.3 DOCX."
    )
    parser.add_argument("--source-docx", default=str(DEFAULT_SOURCE_DOCX))
    parser.add_argument("--target-docx", default=str(DEFAULT_TARGET_DOCX))
    parser.add_argument("--diagram-dir", default=str(DEFAULT_DIAGRAM_DIR))
    parser.add_argument("--overwrite-target", action="store_true")
    parser.add_argument("--placeholder-missing", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    source_docx = Path(args.source_docx).resolve()
    target_docx = Path(args.target_docx).resolve()
    diagram_dir = Path(args.diagram_dir).resolve()

    copy_source_docx(source_docx, target_docx, overwrite_target=args.overwrite_target)

    document = Document(target_docx)
    logs = []
    logs.extend(apply_text_updates(document))
    logs.extend(apply_module_overview_expansion(document))
    logs.extend(apply_traceability_table_updates(document))
    logs.extend(apply_usecase_expansion(document))
    logs.extend(apply_replace_operations(document, diagram_dir, args.placeholder_missing))
    logs.extend(apply_analysis_group(document, diagram_dir, args.placeholder_missing))
    logs.extend(apply_acceptance_expansion(document))
    logs.extend(apply_prototype_expansion(document))
    document.save(target_docx)

    print(f"[OK] saved incremental v1.4: {target_docx}")
    for line in logs:
        print(line)


if __name__ == "__main__":
    main()
