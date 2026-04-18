from __future__ import annotations

import importlib.util
import json
import shutil
import subprocess
import zipfile
import sys
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont


BASE = Path(r"D:\Codes\super-ruc")
SRC_DOCX = BASE / "output" / "doc" / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.4.docx"
OUT_DOCX = BASE / "output" / "doc" / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.5.docx"
OUT_PDF = BASE / "output" / "doc" / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.5.pdf"
WORK_DIR = BASE / "tmp" / "docs" / "v1_5"
HELPER_PATH = BASE / "scripts" / "srs" / "update_srs_v14_incremental.py"
CLASS_DIAGRAM_PATH = BASE / "scripts" / "srs" / "draw_orthogonal_class_diagram.py"
SOURCE_SEQUENCE = BASE / "docs" / "source" / "diagrams" / "mermaid" / "sequence-diagram.mmd"
SOURCE_PROFILE_SEQUENCE = BASE / "docs" / "source" / "diagrams" / "mermaid" / "student-profile-sequence.mmd"
MMDc_SCRIPT = Path(r"C:\Users\znnnnnh2\scoop\apps\nodejs-lts\current\bin\mmdc.ps1")
PUPPETEER_BROWSER = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load module from {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def set_run_text(paragraph, run_index: int, text: str) -> None:
    paragraph.runs[run_index].text = text


def find_exception_table(document: Document):
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header == ["编号", "场景", "处理方式"]:
            for row in table.rows[1:]:
                if row.cells[0].text.strip() == "P3":
                    return table
    raise ValueError("P3 exception table not found")


def find_design_table(document: Document):
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header == ["设计主题", "要求"]:
            return table
    raise ValueError("Design constraint table not found")


def replace_docx_media(docx_path: Path, replacements: dict[str, Path]) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename in replacements:
                data = replacements[info.filename].read_bytes()
            dst.writestr(info, data)
    temp_path.replace(docx_path)


def render_mermaid(source: Path, output_png: Path) -> None:
    config_path = output_png.with_suffix(".puppeteer.json")
    config_path.write_text(
        json.dumps(
            {
                "args": ["--no-sandbox", "--disable-setuid-sandbox"],
                "executablePath": str(PUPPETEER_BROWSER),
                "headless": True,
            }
        ),
        encoding="utf-8",
    )
    subprocess.run(
        [
            "powershell",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(MMDc_SCRIPT),
            "-i",
            str(source),
            "-o",
            str(output_png),
            "-p",
            str(config_path),
            "-b",
            "white",
        ],
        check=True,
    )


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    subprocess.run(
        [
            "pandoc",
            str(docx_path),
            "-o",
            str(pdf_path),
            "--pdf-engine=xelatex",
            "-V",
            "mainfont=Times New Roman",
            "-V",
            "CJKmainfont=SimSun",
        ],
        check=True,
    )


def get_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


LINE = "#000000"
TEXT = "#000000"
BAR_FILL = "#F2F2F2"
PARTICIPANT_FILL = "#FFFFFF"

FONT_PARTICIPANT = get_font(18, bold=True)
FONT_PARTICIPANT_SUB = get_font(14)
FONT_MESSAGE = get_font(15)
FONT_BAR = get_font(14, bold=True)


class SequenceDiagram:
    def __init__(self, participants: list[str], width: int = 2200, height: int = 3200):
        self.participants = participants
        self.width = width
        self.height = height
        self.left_margin = 90
        self.right_margin = 90
        self.top_margin = 18
        self.box_width = 210
        self.box_height = 78
        self.gap_after_box = 24
        self.img = Image.new("RGB", (width, height), "white")
        self.draw = ImageDraw.Draw(self.img)
        self.xs: list[int] = []
        self._draw_participants()
        self.current_y = self.top_margin + self.box_height + self.gap_after_box

    def _draw_participants(self) -> None:
        count = len(self.participants)
        usable_width = self.width - self.left_margin - self.right_margin
        step = usable_width / (count - 1) if count > 1 else 0
        bottom = self.top_margin + self.box_height
        for index, label in enumerate(self.participants):
            x = int(round(self.left_margin + index * step))
            self.xs.append(x)
            half_w = self.box_width // 2
            box = (x - half_w, self.top_margin, x + half_w, bottom)
            self.draw.rounded_rectangle(box, radius=8, fill=PARTICIPANT_FILL, outline=LINE, width=2)
            parts = label.split("\n")
            title = parts[0]
            sub = "\n".join(parts[1:]) if len(parts) > 1 else ""
            title_box = self.draw.textbbox((0, 0), title, font=FONT_PARTICIPANT)
            title_w = title_box[2] - title_box[0]
            title_h = title_box[3] - title_box[1]
            title_x = x - title_w / 2
            title_y = self.top_margin + 8
            self.draw.text((title_x, title_y), title, font=FONT_PARTICIPANT, fill=TEXT)
            if sub:
                sub_box = self.draw.multiline_textbbox((0, 0), sub, font=FONT_PARTICIPANT_SUB, spacing=2)
                sub_w = sub_box[2] - sub_box[0]
                sub_h = sub_box[3] - sub_box[1]
                sub_x = x - sub_w / 2
                sub_y = title_y + title_h + 4
                self.draw.multiline_text((sub_x, sub_y), sub, font=FONT_PARTICIPANT_SUB, fill=TEXT, spacing=2, align="center")
        self._draw_lifelines(bottom=self.height - 42)

    def _draw_lifelines(self, bottom: int) -> None:
        y = self.top_margin + self.box_height + 8
        for x in self.xs:
            self._draw_dashed_vertical(x, y, bottom)

    def _draw_dashed_vertical(self, x: int, y1: int, y2: int, dash: int = 12, gap: int = 9) -> None:
        current = y1
        while current < y2:
            segment_end = min(current + dash, y2)
            self.draw.line((x, current, x, segment_end), fill=LINE, width=2)
            current += dash + gap

    def _draw_horizontal_arrow(self, x1: int, x2: int, y: int, dashed: bool = False) -> None:
        if x1 == x2:
            return
        start, end = (x1, x2) if x1 < x2 else (x2, x1)
        arrow_at_right = x1 < x2
        pad = 14
        line_start = start + pad
        line_end = end - pad
        if dashed:
            dash = 12
            gap = 8
            current = line_start
            while current < line_end:
                seg_end = min(current + dash, line_end)
                self.draw.line((current, y, seg_end, y), fill=LINE, width=2)
                current += dash + gap
        else:
            self.draw.line((line_start, y, line_end, y), fill=LINE, width=2)
        head_x = line_end if arrow_at_right else line_start
        direction = "right" if arrow_at_right else "left"
        self._draw_arrow_head(head_x, y, direction)

    def _draw_arrow_head(self, x: int, y: int, direction: str) -> None:
        size = 9
        if direction == "right":
            pts = [(x, y), (x - size, y - size * 0.7), (x - size, y + size * 0.7)]
        else:
            pts = [(x, y), (x + size, y - size * 0.7), (x + size, y + size * 0.7)]
        self.draw.polygon(pts, fill=LINE)

    def add_bar(self, label: str, fill: str = BAR_FILL) -> None:
        bar_height = 34
        top = self.current_y
        left = self.left_margin - 18
        right = self.width - self.right_margin + 18
        self.draw.rounded_rectangle((left, top, right, top + bar_height), radius=8, fill=fill, outline=LINE, width=1)
        self.draw.text((left + 12, top + 8), label, font=FONT_BAR, fill=TEXT)
        self.current_y += bar_height + 16

    def add_message(self, sender: int, receiver: int, text: str, dashed: bool = False) -> None:
        label = text.replace("\\n", "\n")
        bbox = self.draw.multiline_textbbox((0, 0), label, font=FONT_MESSAGE, spacing=4)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]
        row_h = max(text_h + 18, 54)
        center_x = (self.xs[sender] + self.xs[receiver]) / 2
        text_x = int(round(center_x - text_w / 2))
        text_y = self.current_y + 2
        self.draw.multiline_text((text_x, text_y), label, font=FONT_MESSAGE, fill=TEXT, spacing=4, align="center")
        arrow_y = self.current_y + row_h - 12
        self._draw_horizontal_arrow(self.xs[sender], self.xs[receiver], arrow_y, dashed=dashed)
        self.current_y += row_h + 8

    def save(self, output: Path) -> None:
        cropped = self.img.crop((0, 0, self.width, min(self.height, self.current_y + 38)))
        cropped.save(output)


def draw_document_approval_sequence(output: Path) -> None:
    participants = [
        "学生",
        "学生端",
        "证明服务",
        "预览生成",
        "审批服务",
        "业务库",
        "班主任\n审批老师",
        "团委老师\n授权出件人",
    ]
    diagram = SequenceDiagram(participants, width=2200, height=3200)

    diagram.add_message(0, 1, "进入证明申请页")
    diagram.add_message(1, 2, "请求证明类型\\n模板与附件要求")
    diagram.add_message(2, 5, "读取模板配置\\n与学生基础信息范围")
    diagram.add_message(5, 2, "返回用途说明\\n模板与字段要求", dashed=True)
    diagram.add_message(2, 1, "展示证明类型\\n用途说明和边界提示", dashed=True)
    diagram.add_message(0, 1, "填写用途\\n选择模板或补充说明")

    diagram.add_bar("opt 上传已有附件")
    diagram.add_message(1, 2, "提交附件元数据")
    diagram.add_message(2, 5, "写入附件引用")

    diagram.add_message(1, 2, "提交证明申请")
    diagram.add_message(2, 5, "校验权限范围\\n附件规则和正式边界")

    diagram.add_bar("alt 必须走校级正式链路")
    diagram.add_message(2, 1, "返回“仅预检/归档/跟踪”提示\\n并引导正式渠道", dashed=True)
    diagram.add_message(1, 0, "展示边界说明", dashed=True)

    diagram.add_bar("else 校验通过")
    diagram.add_message(2, 3, "载入模板和学生数据\\n生成 PDF 预览")
    diagram.add_message(3, 2, "返回 preview_file_key", dashed=True)
    diagram.add_message(2, 5, "写入 CommonRequest(待审核)")
    diagram.add_message(2, 5, "写入 DocumentAuditLog(提交事件)")
    diagram.add_message(2, 4, "创建班主任审核任务")
    diagram.add_message(4, 5, "写入 ApprovalTask")
    diagram.add_message(4, 1, "返回申请编号\\n与预览文件引用", dashed=True)
    diagram.add_message(1, 0, "展示“待审核”和 PDF 预览", dashed=True)

    diagram.add_message(6, 1, "打开审核工作台")
    diagram.add_message(1, 4, "加载申请详情\\n附件/预览/历史动作")
    diagram.add_message(4, 5, "查询申请单、附件和审批历史")
    diagram.add_message(5, 4, "返回审核数据", dashed=True)
    diagram.add_message(4, 1, "展示统一审核视图", dashed=True)

    diagram.add_message(6, 1, "提交通过或驳回意见")
    diagram.add_message(1, 4, "写入初审动作")

    diagram.add_bar("alt 驳回")
    diagram.add_message(4, 5, "写入 ApprovalAction(驳回)")
    diagram.add_message(4, 5, "写入 ResubmissionSnapshot")
    diagram.add_message(4, 5, "更新 CommonRequest 为已驳回")
    diagram.add_message(4, 5, "写入 DocumentAuditLog(驳回事件)")
    diagram.add_message(1, 0, "通知在原表单基础上修改重提", dashed=True)

    diagram.add_bar("else 通过")
    diagram.add_message(4, 5, "写入 ApprovalAction(通过)")
    diagram.add_message(4, 5, "更新 CommonRequest 为待出件")
    diagram.add_message(4, 5, "创建团委/出件任务")

    diagram.add_message(7, 1, "审核证明并决定出件方式")
    diagram.add_message(1, 4, "提交终审结果")

    diagram.add_bar("alt 涉密转线下")
    diagram.add_message(4, 5, "写入 ApprovalAction(转线下处理)")
    diagram.add_message(4, 5, "更新 CommonRequest 为线下办理")
    diagram.add_message(4, 5, "写入 DocumentAuditLog(涉密线下处理)")
    diagram.add_message(1, 0, "返回带负责老师联系方式的线下提示卡片", dashed=True)

    diagram.add_bar("else 在线出件完成")
    diagram.add_message(4, 5, "写入 ApprovalAction(终审通过)")
    diagram.add_message(4, 5, "更新 CommonRequest 为已完成")
    diagram.add_message(4, 5, "写入 DocumentAuditLog(出件完成)")
    diagram.add_message(1, 0, "返回出件结果\\n与预览文件下载入口", dashed=True)

    diagram.save(output)


def draw_student_profile_sequence(output: Path) -> None:
    participants = [
        "辅导员",
        "管理端",
        "学生画像服务",
        "学籍与画像数据",
        "学生端",
        "纠错 / 补录服务",
        "审计日志",
        "学生",
    ]
    diagram = SequenceDiagram(participants, width=2200, height=3000)

    diagram.add_bar("rect 管理端查看")
    diagram.add_message(0, 1, "搜索并进入目标学生画像页")
    diagram.add_message(1, 2, "请求学生画像详情")
    diagram.add_message(2, 3, "聚合学籍静态字段与带来源留痕的动态成长字段")

    diagram.add_bar("alt 越权查看非管辖学生")
    diagram.add_message(2, 6, "记录越权尝试")
    diagram.add_message(6, 2, "留痕完成", dashed=True)
    diagram.add_message(2, 1, "拒绝访问", dashed=True)
    diagram.add_message(1, 0, "展示无权限提示", dashed=True)

    diagram.add_bar("else 权限校验通过")
    diagram.add_message(2, 6, "记录查看行为")
    diagram.add_message(6, 2, "留痕完成", dashed=True)
    diagram.add_message(2, 1, "返回全景画像与脱敏字段", dashed=True)
    diagram.add_message(1, 0, "展示画像与导出入口", dashed=True)

    diagram.add_bar("rect 学生本人查看")
    diagram.add_message(7, 4, "查看本人画像")
    diagram.add_message(4, 2, "请求本人画像")
    diagram.add_message(2, 3, "读取本人静态字段与动态记录")

    diagram.add_bar("alt 在读学生")
    diagram.add_message(2, 4, "返回本人画像并展示纠错与成长补录入口")
    diagram.add_message(4, 7, "展示画像、纠错申诉与成长补录入口", dashed=True)
    diagram.add_message(7, 4, "提交信息纠错或成长补录申请")
    diagram.add_message(4, 5, "提交申诉说明或补录内容")
    diagram.add_message(5, 3, "保存待审核申请与来源信息")
    diagram.add_message(3, 5, "保存成功", dashed=True)
    diagram.add_message(5, 6, "记录申请提交")
    diagram.add_message(6, 5, "留痕完成", dashed=True)
    diagram.add_message(5, 4, "返回提交成功", dashed=True)
    diagram.add_message(4, 7, "提示已提交至辅导员审核", dashed=True)

    diagram.add_bar("else 非在读/已归档学生")
    diagram.add_message(2, 4, "返回只读画像并隐藏纠错与补录入口")
    diagram.add_message(4, 7, "展示只读画像", dashed=True)

    diagram.save(output)


def main() -> None:
    helper = load_module(HELPER_PATH, "update_srs_v14_incremental")
    classmod = load_module(CLASS_DIAGRAM_PATH, "draw_orthogonal_class_diagram")

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    diag_dir = WORK_DIR / "diagrams"
    diag_dir.mkdir(parents=True, exist_ok=True)

    student_profile_mmd = diag_dir / "fig-3-11-student-profile-sequence.mmd"
    document_approval_mmd = diag_dir / "fig-3-8-document-approval-sequence.mmd"
    student_profile_png = diag_dir / "fig-3-11-student-profile-sequence.png"
    document_approval_png = diag_dir / "fig-3-8-document-approval-sequence.png"
    extension_class_png = diag_dir / "fig-3-4-extension-analysis-class.png"

    shutil.copy2(SRC_DOCX, OUT_DOCX)

    classmod.OUT = extension_class_png
    classmod.BOXES["StudentProfile"] = classmod.Box(
        870,
        70,
        320,
        "StudentProfile",
        [
            "+bigserial student_id",
            "+varchar student_no",
            "+varchar full_name",
            "+varchar grade_code",
            "+varchar class_code",
            "+varchar political_status",
            "+varchar enrollment_status",
        ],
    )
    classmod.main()

    original_sequence = SOURCE_SEQUENCE.read_text(encoding="utf-8")
    updated_sequence = original_sequence.replace(
        "    alt 转线下处理\n"
        "        ApprovalService->>DB: 写入 ApprovalAction(转线下处理)\n"
        "        ApprovalService->>DB: 更新 CommonRequest 为转线下处理\n"
        "        ApprovalService->>DB: 写入 DocumentAuditLog(线下处理)\n"
        "        Frontend-->>Student: 展示线下出件通知\n"
        "    else 在线出件完成\n",
        "    alt 涉密转线下\n"
        "        ApprovalService->>DB: 写入 ApprovalAction(转线下处理)\n"
        "        ApprovalService->>DB: 更新 CommonRequest 为线下办理\n"
        "        ApprovalService->>DB: 写入 DocumentAuditLog(涉密线下处理)\n"
        "        Frontend-->>Student: 返回带负责老师联系方式的线下提示卡片\n"
        "    else 在线出件完成\n",
    )
    if updated_sequence == original_sequence:
        raise RuntimeError("Failed to update document approval sequence source")
    document_approval_mmd.write_text(updated_sequence, encoding="utf-8")
    draw_document_approval_sequence(document_approval_png)

    original_profile = SOURCE_PROFILE_SEQUENCE.read_text(encoding="utf-8")
    updated_profile = original_profile.replace(
        "        ProfileService->>ProfileRepo: 读取本人静态字段与动态记录\n"
        "        ProfileService-->>StudentApp: 返回本人画像并隐藏管理元数据\n"
        "        StudentApp-->>Student: 展示画像、纠错申诉与成长补录入口\n"
        "        Student->>StudentApp: 提交信息纠错或成长补录申请\n"
        "        StudentApp->>UpdateService: 提交申诉说明或补录内容\n"
        "        UpdateService->>ProfileRepo: 保存待审核申请与来源信息\n"
        "        ProfileRepo-->>UpdateService: 保存成功\n"
        "        UpdateService->>AuditLog: 记录申请提交\n"
        "        AuditLog-->>UpdateService: 留痕完成\n"
        "        UpdateService-->>StudentApp: 返回提交成功\n"
        "        StudentApp-->>Student: 提示已提交至辅导员审核\n",
        "        ProfileService->>ProfileRepo: 读取本人静态字段与动态记录\n"
        "        alt 在读学生\n"
        "            ProfileService-->>StudentApp: 返回本人画像并展示纠错与成长补录入口\n"
        "            StudentApp-->>Student: 展示画像、纠错申诉与成长补录入口\n"
        "            Student->>StudentApp: 提交信息纠错或成长补录申请\n"
        "            StudentApp->>UpdateService: 提交申诉说明或补录内容\n"
        "            UpdateService->>ProfileRepo: 保存待审核申请与来源信息\n"
        "            ProfileRepo-->>UpdateService: 保存成功\n"
        "            UpdateService->>AuditLog: 记录申请提交\n"
        "            AuditLog-->>UpdateService: 留痕完成\n"
        "            UpdateService-->>StudentApp: 返回提交成功\n"
        "            StudentApp-->>Student: 提示已提交至辅导员审核\n"
        "        else 非在读/已归档学生\n"
        "            ProfileService-->>StudentApp: 返回只读画像并隐藏纠错与补录入口\n"
        "            StudentApp-->>Student: 展示只读画像\n"
        "        end\n",
    )
    if updated_profile == original_profile:
        raise RuntimeError("Failed to update student profile sequence source")
    student_profile_mmd.write_text(updated_profile, encoding="utf-8")
    draw_student_profile_sequence(student_profile_png)

    document = Document(OUT_DOCX)

    # Cover and date.
    set_run_text(document.paragraphs[0], 5, "5")
    date_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("日期："))
    set_run_text(date_paragraph, 1, "2026-04-15")

    # 2.2 limitations and constraints.
    audit_anchor = helper.find_first_paragraph(
        document,
        (helper.ParagraphSelector("短信通道属于可配置能力，不作为一期强制前提；站内通知和邮件优先。"),),
    )
    if audit_anchor is None:
        raise RuntimeError("2.2 audit anchor not found")
    helper.clone_paragraph_after(
        audit_anchor,
        audit_anchor,
        "审计约束：审批、配置、导入导出、通知发送、权限变化和敏感访问必须保留可查询日志，留存期不少于一个学期，超期日志需提供定期自动清理或冷备份归档至历史库的机制，防止生产主表膨胀。",
    )

    # 3.1 software function overview.
    proof_anchor = helper.find_first_paragraph(
        document,
        (
            helper.ParagraphSelector(
                "材料与附件：盖章场景原则上要求上传待盖章文件；证明场景支持上传附件或填写文本说明；",
                mode="prefix",
            ),
        ),
    )
    if proof_anchor is None:
        raise RuntimeError("3.1 proof anchor not found")
    helper.set_plain_text(
        proof_anchor,
        "材料与附件：盖章场景原则上要求上传待盖章文件；证明场景支持上传附件或填写文本说明；若审批老师在初审时研判该事项涉及敏感或涉密内容，有权在系统中直接勾选“转线下办理”；此时系统终止线上文件生成，保留日志记录，并向学生端下发包含后续指导与老师联系方式的系统提示。",
    )

    profile_anchor = helper.find_first_paragraph(
        document,
        (
            helper.ParagraphSelector(
                "责任边界与安全约束：本模块仅作为信息聚合与展示工具，不包含对学生综合素质的自动评分、排名或评价结论输出。",
                mode="prefix",
            ),
        ),
    )
    if profile_anchor is None:
        raise RuntimeError("3.1 profile anchor not found")
    helper.clone_paragraph_after(
        profile_anchor,
        profile_anchor,
        "账号与画像生命周期管理：系统需感知学生学籍状态变化。当学生毕业、休学或转出信息学院时，其系统账号自动降级为“历史归档”状态；冻结其修改、申诉与补录权限；对应画像数据移出当前的班级/学院日常管理统计视图，仅保留超级管理员或授权辅导员的历史查阅权限。",
    )

    import_export_anchor = helper.find_first_paragraph(
        document,
        (
            helper.ParagraphSelector(
                "Excel/Word/PDF 导入导出、模板化校验、错误报告、整批回滚和审计留痕能力。",
                mode="prefix",
            ),
        ),
    )
    if import_export_anchor is None:
        raise RuntimeError("3.1 import/export anchor not found")
    helper.set_plain_text(
        import_export_anchor,
        "Excel/Word/PDF 导入导出、模板化校验、整批回滚和审计留痕能力；当发生导入错误时，系统需提供带有具体错误原因（定位到行级与错误字段）的错误报告文件，或原样返回追加错误列的 Excel 供管理员下载修改。",
    )

    # 3.2 use case model.
    exception_table = find_exception_table(document)
    exception_table.rows[1].cells[2].text = (
        "审批老师在工作台标记“转线下处理”；系统终止线上流转，保留审批历史，并向学生推送线下办理提示及负责老师联系方式，不生成在线文件。"
    )

    # 3.3 analysis model.
    entity_table = helper.find_table_by_header(document, ("实体类别", "代表实体", "说明"))
    if entity_table is None:
        raise RuntimeError("Analysis entity table not found")
    for row in entity_table.rows[1:]:
        if row.cells[0].text.strip() == "主数据":
            row.cells[2].text = "承载学号、姓名、年级、班级、政治面貌和学籍状态等基础信息。"
            break
    else:
        raise RuntimeError("StudentProfile row not found in analysis entity table")

    # 4.1 performance.
    perf_anchor = helper.find_first_paragraph(
        document,
        (
            helper.ParagraphSelector(
                "标准格式 Excel 主数据 100 条批量导入应在 60 秒内完成成功提交或整批失败回滚，并输出可下载错误报告。",
                mode="prefix",
            ),
        ),
    )
    if perf_anchor is None:
        raise RuntimeError("4.1 performance anchor not found")
    helper.set_plain_text(
        perf_anchor,
        "标准格式 Excel 主数据 100 条批量导入应在 60 秒内完成成功提交或整批失败回滚，并在回滚时输出可精确定位错误行的可下载错误报告。",
    )

    # 4.2 design constraints table.
    design_table = find_design_table(document)
    updated_rows = [
        ["设计主题", "要求"],
        ["事务一致性", "申请状态、党团节点、导入批次和学业规则更新不得出现部分生效或静默丢失。"],
        ["审计留存", "关键审批、配置、导出和敏感访问记录至少保留 1 个学期，并支持按角色、时间和动作类型检索；架构设计上需引入定时任务（如定期的冷数据迁移脚本），将超出保留期的数据从高频查询表中剥离。"],
        ["权限控制", "学生默认只看本人；班主任默认看职责范围；高敏字段默认脱敏；高敏导出需额外授权。"],
        ["账号生命周期", "学籍状态（在读/毕业/休学/转出）必须作为全局拦截器或底层状态机的一部分，非在校状态账号自动切断业务提交流程，仅保留只读查询能力。"],
        ["内容治理", "知识条目必须携带来源、版本和更新时间；停用或过期内容不得继续作为当前有效口径对外展示。"],
        ["文件交换", "导入模板需固定字段结构；失败批次必须保障主库数据零污染（事务回滚），并直接提供行级别的错误诊断依据（如行号与错因对应表），避免用户人工排查。"],
        ["正式边界", "凡涉及校级正式流程边界的页面，必须显示风险提示和正式渠道引导文案。"],
        ["学业边界", "学业分析页面必须显式提示“仅供辅助参考”，不得替代学院或学校最终审核。"],
        ["易用性", "审批场景应在单一工作台内完成主要判断动作，避免跨页反复查找附件和历史。"],
    ]
    helper.set_table_rows(design_table, updated_rows)

    # 5.4 prototype.
    proof_proto = helper.find_first_paragraph(
        document,
        (
            helper.ParagraphSelector(
                "驳回后学生可回到原表单修改并重提；涉密场景应支持转线下办理并保留线上留痕。",
                mode="prefix",
            ),
        ),
    )
    if proof_proto is None:
        raise RuntimeError("5.4 prototype anchor not found")
    helper.set_plain_text(
        proof_proto,
        "驳回后学生可回到原表单修改并重提；涉密场景应支持转线下办理，并向学生返回线下办理提示与负责老师联系方式，同时保留线上留痕。",
    )

    profile_proto = helper.find_first_paragraph(
        document,
        (
            helper.ParagraphSelector(
                "学生登录后仅可查看本人画像。页面展示学籍静态信息及已收录的动态成长记录，不显示数据来源、录入人等管理元数据。",
                mode="prefix",
            ),
        ),
    )
    if profile_proto is None:
        raise RuntimeError("5.7 student profile prototype anchor not found")
    helper.clone_paragraph_after(
        profile_proto,
        profile_proto,
        "若学生处于毕业、休学或转出等历史归档状态，则页面仅保留只读查看，不显示信息纠错申诉或成长补录入口。",
    )

    # Revision history table.
    history_table = document.tables[0]
    history_table.rows[2].cells[0].text = "2"
    history_table.rows[2].cells[1].text = "2026-04-15"
    history_table.rows[2].cells[2].text = "项目组"
    history_table.rows[2].cells[3].text = "根据 fix.md 完成账号生命周期、审计归档、涉密转线下处理和导入错误报告等 v1.5 修订。"
    history_table.rows[2].cells[4].text = "V1.5"

    document.save(OUT_DOCX)

    replace_docx_media(
        OUT_DOCX,
        {
            "word/media/image5.png": extension_class_png,
            "word/media/image9.png": document_approval_png,
            "word/media/image12.png": student_profile_png,
        },
    )

    export_pdf(OUT_DOCX, OUT_PDF)

    print(f"[OK] docx: {OUT_DOCX}")
    print(f"[OK] pdf: {OUT_PDF}")
    print(f"[OK] diagrams: {diag_dir}")


if __name__ == "__main__":
    main()
