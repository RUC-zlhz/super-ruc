from __future__ import annotations

import json
import re
import shutil
import subprocess
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import pythoncom
import win32com.client
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


BASE = Path(r"D:\Codes\super-ruc")
SOURCE_DOCX = BASE / "output" / "doc" / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.5.docx"
SOURCE_PDF = BASE / "output" / "doc" / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.5.pdf"
WORK_DIR = BASE / "tmp" / "docs" / "v1_5"
WORK_DOCX = WORK_DIR / "work-v1.5.docx"
WORK_PDF = WORK_DIR / "work-v1.5.pdf"
BACKUP_DIR = WORK_DIR / "backups"
SVG_SRC_DIR = WORK_DIR / "svg-src"
SVG_OUT_DIR = WORK_DIR / "svg"
DIAGRAM_DIR = BASE / "docs" / "source" / "diagrams" / "mermaid"
PUPPETEER_CONFIG = WORK_DIR / "puppeteer.json"
MMDc_SCRIPT = Path(r"C:\Users\znnnnnh2\scoop\apps\nodejs-lts\current\bin\mmdc.ps1")
CHROME = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")

SPACE_RE = re.compile(r"\s+")

SPLIT_CAPTION_A = "图 3-2（a）管理侧七模块用例图（知识库、党团事务与通知推送）"
SPLIT_CAPTION_B = "图 3-2（b）管理侧七模块用例图（证明审批、预警分析、荣誉、画像与治理）"


@dataclass(frozen=True)
class FigureRenderSpec:
    caption: str
    source_path: Path | None
    output_name: str
    inline_source_text: str | None = None


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ").replace("\u00A0", " ").replace("\r", " ").replace("\x07", " ")
    return SPACE_RE.sub(" ", text).strip()


def paragraph_text_matches(paragraph: Paragraph, target: str) -> bool:
    return normalize_text(paragraph.text) == normalize_text(target)


def find_paragraph(document: DocumentObject, target: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph_text_matches(paragraph, target):
            return paragraph
    raise ValueError(f"Paragraph not found: {target}")


def paragraph_index(document: DocumentObject, target: Paragraph) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError(f"Paragraph not found in document: {target.text!r}")


def get_previous_paragraph(document: DocumentObject, paragraph: Paragraph) -> Paragraph:
    index = paragraph_index(document, paragraph)
    if index == 0:
        raise ValueError("Paragraph has no previous sibling")
    return document.paragraphs[index - 1]


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_caption_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False
    paragraph.add_run(text)


def clone_paragraph_after(paragraph: Paragraph, template: Paragraph) -> Paragraph:
    new_paragraph_element = deepcopy(template._p)
    paragraph._p.addnext(new_paragraph_element)
    return Paragraph(new_paragraph_element, paragraph._parent)


def prepare_docx_structure() -> None:
    document = Document(WORK_DOCX)
    split_a_exists = any(paragraph_text_matches(paragraph, SPLIT_CAPTION_A) for paragraph in document.paragraphs)
    split_b_exists = any(paragraph_text_matches(paragraph, SPLIT_CAPTION_B) for paragraph in document.paragraphs)
    if split_a_exists and split_b_exists:
        document.save(WORK_DOCX)
        return

    original_caption = find_paragraph(document, "图 3-2 管理侧七模块用例图")
    picture_paragraph = get_previous_paragraph(document, original_caption)

    set_caption_text(original_caption, SPLIT_CAPTION_A)

    second_picture = clone_paragraph_after(original_caption, picture_paragraph)
    clear_paragraph_content(second_picture)
    second_picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second_picture.paragraph_format.keep_together = True
    second_picture.paragraph_format.keep_with_next = True

    second_caption = clone_paragraph_after(second_picture, original_caption)
    set_caption_text(second_caption, SPLIT_CAPTION_B)

    document.save(WORK_DOCX)


def split_admin_usecase_upper() -> str:
    return """%%{
  init: {
    "theme": "base",
    "htmlLabels": false,
    "themeVariables": {
      "fontFamily": "Times New Roman, SimSun, serif",
      "primaryColor": "#ffffff",
      "primaryBorderColor": "#000000",
      "primaryTextColor": "#000000",
      "lineColor": "#000000",
      "clusterBkg": "#f0f0f0",
      "clusterBorder": "#7f8c8d",
      "fontSize": "16px",
      "lineWidth": "1.5px"
    },
    "flowchart": {
      "curve": "linear",
      "nodeSpacing": 50,
      "rankSpacing": 55,
      "htmlLabels": false
    }
  }
}%%
flowchart LR
    subgraph Actors[管理侧参与者]
        direction TB
        LeagueTeacher[团委老师]
        Leader[学院领导]
        Admin[超级管理员]
    end

    subgraph Platform[管理侧七模块用例图（上）]
        direction TB
        subgraph M1Group[模块一：智能问答与政策知识库]
            direction LR
            UC1([维护政策知识库])
            UC2([维护标准答案与模板])
        end

        subgraph M2Group[模块二：党团事务流程管理]
            direction LR
            UC3([维护党团节点与提醒])
            UC4([维护理论题库])
            UC5([审批党团流程])
        end

        subgraph M3Group[模块三：信息集成与精准推送]
            direction LR
            UC6([录入 / 汇聚官方通知])
            UC7([圈定人群并推送])
            UC8([查看触达结果])
        end
    end

    LeagueTeacher --- UC3
    LeagueTeacher --- UC4
    LeagueTeacher --- UC5
    LeagueTeacher --- UC6
    LeagueTeacher --- UC7
    LeagueTeacher --- UC8

    Leader --- UC5

    Admin --- UC1
    Admin --- UC2
    Admin --- UC6
    Admin --- UC7
"""


def split_admin_usecase_lower() -> str:
    return """%%{
  init: {
    "theme": "base",
    "htmlLabels": false,
    "themeVariables": {
      "fontFamily": "Times New Roman, SimSun, serif",
      "primaryColor": "#ffffff",
      "primaryBorderColor": "#000000",
      "primaryTextColor": "#000000",
      "lineColor": "#000000",
      "clusterBkg": "#f0f0f0",
      "clusterBorder": "#7f8c8d",
      "fontSize": "16px",
      "lineWidth": "1.5px"
    },
    "flowchart": {
      "curve": "linear",
      "nodeSpacing": 48,
      "rankSpacing": 55,
      "htmlLabels": false
    }
  }
}%%
flowchart LR
    subgraph Actors[管理侧参与者]
        direction TB
        Advisor[辅导员 / 班主任 / 审批老师]
        Leader[学院领导]
        Admin[超级管理员]
    end

    subgraph Platform[管理侧七模块用例图（下）]
        direction LR
        subgraph MainModules[核心业务模块]
            direction TB
            subgraph M4Group[模块四：电子证明生成与审批流程]
                direction LR
                UC9([审批事务申请])
                UC10([查看附件 / 历史 / PDF 预览])
                UC11([导出审批记录])
            end

            subgraph M5Group[模块五：学业情况分析与预警]
                direction LR
                UC12([维护培养方案与规则])
                UC13([查看学业风险与建议统计])
            end

            subgraph M6Group[模块六：奖励荣誉展示]
                direction LR
                UC14([查看荣誉榜单与详情])
                UC15([导入荣誉并核验学籍])
                UC16([维护授权 / 审核 / 展示有效期])
                UC17([归档 / 撤销荣誉条目])
            end

            subgraph M7Group[模块七：学生画像与信息管理]
                direction LR
                UC18([查看学生画像全景])
                UC19([导出画像快照])
                UC20([审核补录并处理纠错])
            end
        end

        subgraph SupportGroup[平台治理与运维]
            direction TB
            UC21([维护角色与字段权限])
            UC22([查看审计日志])
            UC23([导入 / 导出 Excel Word PDF])
        end
    end

    Advisor --- UC9
    Advisor --- UC10
    Advisor --- UC13
    Advisor --- UC14
    Advisor --- UC18
    Advisor --- UC20

    Leader --- UC9
    Leader --- UC13
    Leader --- UC14
    Leader --- UC18
    Leader --- UC19

    Admin --- UC12
    Admin --- UC15
    Admin --- UC16
    Admin --- UC18
    Admin --- UC21
    Admin --- UC22
    Admin --- UC23

    UC10 -. 支撑 .-> UC9
    UC11 -. 留痕 .-> UC22
    UC15 -. 导入依赖 .-> UC23
    UC16 -. 留痕 .-> UC22
    UC17 -. 留痕 .-> UC22
    UC18 -. 留痕 .-> UC22
    UC19 -. 导出留痕 .-> UC22
    UC20 -. 留痕 .-> UC22
    UC23 -. 留痕 .-> UC22
    UC13 -. 读取 .-> UC12
"""


def build_document_approval_sequence_text() -> str:
    original_sequence = (DIAGRAM_DIR / "sequence-diagram.mmd").read_text(encoding="utf-8")
    updated_sequence = original_sequence.replace(
        "    alt 转线下处理\n"
        "        ApprovalService->>DB: 写入 ApprovalAction(转线下处理)\n"
        "        ApprovalService->>DB: 更新 CommonRequest 为转线下处理\n"
        "        ApprovalService->>DB: 写入 DocumentAuditLog(线下处理)\n"
        "        Frontend-->>Student: 展示线下出件通知\n"
        "    else 在线出件完成\n",
        "    alt 涉密转线下\n"
        "        ApprovalService->>DB: 写入 ApprovalAction(转线下处理)\n"
        "        ApprovalService->>DB: 更新 CommonRequest 为线下办理\n"
        "        ApprovalService->>DB: 写入 DocumentAuditLog(涉密线下处理)\n"
        "        Frontend-->>Student: 返回带负责老师联系方式的线下提示卡片\n"
        "    else 在线出件完成\n",
    )
    if updated_sequence == original_sequence:
        raise RuntimeError("Failed to build document approval sequence source")
    return updated_sequence


def build_student_profile_sequence_text() -> str:
    original_profile = (DIAGRAM_DIR / "student-profile-sequence.mmd").read_text(encoding="utf-8")
    updated_profile = original_profile.replace(
        "        ProfileService->>ProfileRepo: 读取本人静态字段与动态记录\n"
        "        ProfileService-->>StudentApp: 返回本人画像并隐藏管理元数据\n"
        "        StudentApp-->>Student: 展示画像、纠错申诉与成长补录入口\n"
        "        Student->>StudentApp: 提交信息纠错或成长补录申请\n"
        "        StudentApp->>UpdateService: 提交申诉说明或补录内容\n"
        "        UpdateService->>ProfileRepo: 保存待审核申请与来源信息\n"
        "        ProfileRepo-->>UpdateService: 保存成功\n"
        "        UpdateService->>AuditLog: 记录申请提交\n"
        "        AuditLog-->>UpdateService: 留痕完成\n"
        "        UpdateService-->>StudentApp: 返回提交成功\n"
        "        StudentApp-->>Student: 提示已提交至辅导员审核\n",
        "        ProfileService->>ProfileRepo: 读取本人静态字段与动态记录\n"
        "        alt 在读学生\n"
        "            ProfileService-->>StudentApp: 返回本人画像并展示纠错与成长补录入口\n"
        "            StudentApp-->>Student: 展示画像、纠错申诉与成长补录入口\n"
        "            Student->>StudentApp: 提交信息纠错或成长补录申请\n"
        "            StudentApp->>UpdateService: 提交申诉说明或补录内容\n"
        "            UpdateService->>ProfileRepo: 保存待审核申请与来源信息\n"
        "            ProfileRepo-->>UpdateService: 保存成功\n"
        "            UpdateService->>AuditLog: 记录申请提交\n"
        "            AuditLog-->>UpdateService: 留痕完成\n"
        "            UpdateService-->>StudentApp: 返回提交成功\n"
        "            StudentApp-->>Student: 提示已提交至辅导员审核\n"
        "        else 非在读/已归档学生\n"
        "            ProfileService-->>StudentApp: 返回只读画像并隐藏纠错与补录入口\n"
        "            StudentApp-->>Student: 展示只读画像\n"
        "        end\n",
    )
    if updated_profile == original_profile:
        raise RuntimeError("Failed to build student profile sequence source")
    return updated_profile


def ensure_workdirs() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    SVG_SRC_DIR.mkdir(parents=True, exist_ok=True)
    SVG_OUT_DIR.mkdir(parents=True, exist_ok=True)


def inject_root_html_labels(text: str) -> str:
    if '"themeVariables"' in text and '"htmlLabels": false' not in text.split('"themeVariables"', 1)[0]:
        text = text.replace('"themeVariables": {', '"htmlLabels": false,\n    "themeVariables": {', 1)
    return text


def sanitize_mermaid_text(text: str) -> str:
    if "*** Add File:" in text:
        text = text.split("*** Add File:", 1)[0].rstrip() + "\n"
    return inject_root_html_labels(text)


def sanitize_svg_text(text: str) -> str:
    def fill_background(match: re.Match[str]) -> str:
        style = match.group(1)
        normalized = style.strip()
        if normalized and not normalized.endswith(";"):
            normalized += ";"
        if "fill:" not in normalized:
            normalized += "fill:#ffffff;"
        if "stroke:" not in normalized:
            normalized += "stroke:none;"
        return f'<rect class="background" style="{normalized}"'

    return re.sub(r'<rect class="background" style="([^"]*)"', fill_background, text)


def write_mermaid_source(name: str, text: str) -> Path:
    path = SVG_SRC_DIR / name
    path.write_text(sanitize_mermaid_text(text), encoding="utf-8")
    return path


def write_source_copy(name: str, source_path: Path) -> Path:
    text = source_path.read_text(encoding="utf-8")
    return write_mermaid_source(name, text)


def build_render_specs() -> list[FigureRenderSpec]:
    return [
        FigureRenderSpec("图 2-1 软件产品与外部环境关系图", DIAGRAM_DIR / "system-context.mmd", "fig-2-1-system-context.svg"),
        FigureRenderSpec("图 3-1 学生侧七模块用例图", DIAGRAM_DIR / "use-case-student.mmd", "fig-3-1-student-usecase.svg"),
        FigureRenderSpec(SPLIT_CAPTION_A, None, "fig-3-2a-admin-usecase.svg", split_admin_usecase_upper()),
        FigureRenderSpec(SPLIT_CAPTION_B, None, "fig-3-2b-admin-usecase.svg", split_admin_usecase_lower()),
        FigureRenderSpec("图 3-3 核心业务分析类图", DIAGRAM_DIR / "class-diagram.mmd", "fig-3-3-core-analysis-class.svg"),
        FigureRenderSpec("图 3-4 奖励荣誉与学生画像扩展类图", DIAGRAM_DIR / "extension-class-diagram.mmd", "fig-3-4-extension-analysis-class.svg"),
        FigureRenderSpec("图 3-5 知识问答与模板下载时序图", DIAGRAM_DIR / "knowledge-template-sequence.mmd", "fig-3-5-knowledge-template-sequence.svg"),
        FigureRenderSpec("图 3-6 党团事务流程状态图", DIAGRAM_DIR / "party-workflow-state.mmd", "fig-3-6-party-workflow-state.svg"),
        FigureRenderSpec("图 3-7 官方信息汇聚与精准推送时序图", DIAGRAM_DIR / "notice-push-sequence.mmd", "fig-3-7-notice-push-sequence.svg"),
        FigureRenderSpec("图 3-8 电子证明生成与审批时序图", None, "fig-3-8-document-approval-sequence.svg", build_document_approval_sequence_text()),
        FigureRenderSpec("图 3-9 学业分析与预警活动图", DIAGRAM_DIR / "academic-warning-activity.mmd", "fig-3-9-academic-warning-activity.svg"),
        FigureRenderSpec("图 3-10 荣誉榜单浏览与详情查看时序图", DIAGRAM_DIR / "honor-display-sequence.mmd", "fig-3-10-honor-display-sequence.svg"),
        FigureRenderSpec("图 3-11 学生画像查看与纠错 / 成长补录时序图", None, "fig-3-11-student-profile-sequence.svg", build_student_profile_sequence_text()),
    ]


def ensure_puppeteer_config() -> None:
    PUPPETEER_CONFIG.write_text(
        json.dumps(
            {
                "args": ["--no-sandbox", "--disable-setuid-sandbox"],
                "executablePath": str(CHROME),
                "headless": True,
            }
        ),
        encoding="utf-8",
    )


def render_mermaid(input_path: Path, output_path: Path) -> None:
    subprocess.run(
        [
            "powershell",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(MMDc_SCRIPT),
            "-i",
            str(input_path),
            "-o",
            str(output_path),
            "-p",
            str(PUPPETEER_CONFIG),
            "-b",
            "white",
        ],
        check=True,
    )


def generate_svg_assets() -> dict[str, Path]:
    ensure_puppeteer_config()
    caption_to_svg: dict[str, Path] = {}
    for spec in build_render_specs():
        if spec.inline_source_text is not None:
            source_copy = write_mermaid_source(spec.output_name.replace(".svg", ".mmd"), spec.inline_source_text)
        elif spec.source_path is not None:
            source_copy = write_source_copy(spec.output_name.replace(".svg", ".mmd"), spec.source_path)
        else:
            raise ValueError(f"Missing source for {spec.caption}")

        output_path = SVG_OUT_DIR / spec.output_name
        render_mermaid(source_copy, output_path)
        svg_text = sanitize_svg_text(output_path.read_text(encoding="utf-8"))
        output_path.write_text(svg_text, encoding="utf-8")
        if "foreignObject" in svg_text or "<div" in svg_text or "<span" in svg_text:
            raise ValueError(f"SVG still contains HTML labels: {output_path}")
        caption_to_svg[spec.caption] = output_path
    return caption_to_svg


def word_find_paragraph(document, text: str) -> tuple[int, object]:
    for index in range(1, document.Paragraphs.Count + 1):
        paragraph = document.Paragraphs(index)
        if normalize_text(paragraph.Range.Text) == normalize_text(text):
            return index, paragraph
    raise ValueError(f"Word paragraph not found: {text}")


def replace_inline_picture(paragraph, svg_path: Path, width_points: float) -> None:
    count = paragraph.Range.InlineShapes.Count
    for _ in range(count):
        paragraph.Range.InlineShapes(1).Delete()
    shape = paragraph.Range.InlineShapes.AddPicture(FileName=str(svg_path), LinkToFile=False, SaveWithDocument=True)
    shape.LockAspectRatio = True
    shape.Width = width_points


def replace_images_with_word(caption_to_svg: dict[str, Path]) -> None:
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        word.ScreenUpdating = False
        document = word.Documents.Open(
            FileName=str(WORK_DOCX),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
            Revert=False,
            Visible=False,
            OpenAndRepair=True,
            NoEncodingDialog=True,
        )

        ordered_captions = [
            "图 2-1 软件产品与外部环境关系图",
            "图 3-1 学生侧七模块用例图",
            SPLIT_CAPTION_A,
            SPLIT_CAPTION_B,
            "图 3-3 核心业务分析类图",
            "图 3-4 奖励荣誉与学生画像扩展类图",
            "图 3-5 知识问答与模板下载时序图",
            "图 3-6 党团事务流程状态图",
            "图 3-7 官方信息汇聚与精准推送时序图",
            "图 3-8 电子证明生成与审批时序图",
            "图 3-9 学业分析与预警活动图",
            "图 3-10 荣誉榜单浏览与详情查看时序图",
            "图 3-11 学生画像查看与纠错 / 成长补录时序图",
        ]

        for caption in ordered_captions:
            caption_index, _ = word_find_paragraph(document, caption)
            picture_paragraph = document.Paragraphs(caption_index - 1)
            if picture_paragraph.Range.InlineShapes.Count > 0:
                width_points = picture_paragraph.Range.InlineShapes(1).Width
            elif caption == SPLIT_CAPTION_B:
                width_points = 396.0
            else:
                width_points = 396.0
            replace_inline_picture(picture_paragraph, caption_to_svg[caption], width_points)

        document.Save()
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit(0)
            except Exception:
                pass
        pythoncom.CoUninitialize()


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "python",
            str(BASE / "scripts" / "srs" / "export_docx_pdf.py"),
            "--docx",
            str(docx_path),
            "--pdf",
            str(pdf_path),
        ],
        check=True,
    )
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)


def make_backups() -> None:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    shutil.copy2(SOURCE_DOCX, BACKUP_DIR / f"{SOURCE_DOCX.stem}.before-{timestamp}{SOURCE_DOCX.suffix}")
    if SOURCE_PDF.exists():
        shutil.copy2(SOURCE_PDF, BACKUP_DIR / f"{SOURCE_PDF.stem}.before-{timestamp}{SOURCE_PDF.suffix}")


def copy_outputs_back() -> None:
    shutil.copy2(WORK_DOCX, SOURCE_DOCX)
    shutil.copy2(WORK_PDF, SOURCE_PDF)


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    ensure_workdirs()
    make_backups()
    shutil.copy2(SOURCE_DOCX, WORK_DOCX)
    prepare_docx_structure()
    caption_to_svg = generate_svg_assets()
    replace_images_with_word(caption_to_svg)
    export_pdf(WORK_DOCX, WORK_PDF)
    copy_outputs_back()
    print(f"UPDATED {SOURCE_DOCX}")
    print(f"UPDATED {SOURCE_PDF}")


if __name__ == "__main__":
    main()
