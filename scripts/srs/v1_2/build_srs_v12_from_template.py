from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parents[3]
TEMPLATE = BASE / "软件需求规格说明书模板.docx"
OUT = BASE / "output" / "doc" / "软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.2.docx"
DIAG = BASE / "docs" / "source" / "diagrams" / "rendered" / "v1_2"

PROJECT = "信息学院学生综合服务与党团管理平台"
TODAY = "2026-04-14"
VERSION = "V1.2"

doc = Document(TEMPLATE)


def find_para_exact(text: str) -> Paragraph:
    found = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            found = paragraph
    if found is None:
        raise ValueError(f"Paragraph not found: {text}")
    return found


def find_para_startswith(prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    bold: bool = False,
    center: bool = False,
    style_name: str | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    if center:
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_para


def insert_heading_after(paragraph: Paragraph, text: str, level: int) -> Paragraph:
    return insert_paragraph_after(paragraph, text=text, style_name=f"Heading {level}")


def insert_table_after(paragraph: Paragraph, headers: list[str], rows: list[list[str]]) -> Table:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after_table(
    table: Table,
    text: str = "",
    bold: bool = False,
    center: bool = False,
    style_name: str | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    if style_name:
        new_para.style = style_name
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    if center:
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_para


def insert_picture_after(
    paragraph: Paragraph,
    image_path: Path,
    width_inches: float,
    caption: str,
    page_break_before: bool = False,
    page_break_after: bool = False,
) -> Paragraph:
    anchor = paragraph
    if page_break_before:
        break_para = insert_paragraph_after(paragraph)
        break_para.add_run().add_break(WD_BREAK.PAGE)
        anchor = break_para
    picture_para = insert_paragraph_after(anchor)
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_para.paragraph_format.keep_together = True
    picture_para.paragraph_format.keep_with_next = True
    picture_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    caption_para = insert_paragraph_after(picture_para, caption, center=True)
    caption_para.paragraph_format.keep_together = True
    if page_break_after:
        after_para = insert_paragraph_after(caption_para)
        after_para.add_run().add_break(WD_BREAK.PAGE)
        return after_para
    return caption_para


def add_bullets(paragraph: Paragraph, items: list[str]) -> Paragraph:
    cursor = paragraph
    for item in items:
        cursor = insert_paragraph_after(cursor, f"- {item}")
    return cursor


def clear_paragraph_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            paragraph._p.remove(child)


def apply_cover_content() -> None:
    doc_no = find_para_startswith("文档编号：")
    clear_paragraph_runs(doc_no)
    run = doc_no.add_run("文档编号：")
    run.bold = True
    run = doc_no.add_run(f"{PROJECT} ")
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run = doc_no.add_run("– SRS – ")
    run.bold = True
    run = doc_no.add_run(VERSION)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    doc_no.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    project_title = find_para_exact("<项目名称>")
    clear_paragraph_runs(project_title)
    run = project_title.add_run(PROJECT)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size = Pt(24)
    project_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    srs_title = find_para_exact("软件需求规格说明书")
    clear_paragraph_runs(srs_title)
    run = srs_title.add_run("软件需求规格说明书")
    run.bold = True
    run.font.size = Pt(26)
    srs_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    date_para = find_para_exact("日期：")
    clear_paragraph_runs(date_para)
    run = date_para.add_run("日期：")
    run.bold = True
    run.font.size = Pt(12)
    run = date_para.add_run(TODAY)
    run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_section(heading_text: str, next_heading_text: str | None, builder) -> None:
    start = find_para_exact(heading_text)
    end = find_para_exact(next_heading_text) if next_heading_text else None
    deleting = False
    paragraphs_to_delete: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        if paragraph == start:
            deleting = True
            continue
        if deleting:
            if end is not None and paragraph == end:
                break
            paragraphs_to_delete.append(paragraph)
    for paragraph in paragraphs_to_delete:
        delete_paragraph(paragraph)
    builder(start)


def cleanup_template_fragments() -> None:
    exact_fragments = {
        "项目名称、简称或代号",
        "用户单位",
        "开发单位",
        "大致功能和用途等>",
        "需交付哪些内容",
        "这些内容以何种形式交付：电子文件、打印材料 >",
    }
    to_delete: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("<") or text in exact_fragments:
            to_delete.append(paragraph)
    for paragraph in to_delete:
        delete_paragraph(paragraph)


def build_11(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(
        heading,
        "本文档用于定义信息学院学生综合服务与党团管理平台的一期软件需求规格，作为项目立项、方案评审、详细设计、开发实现、测试验收和后续运维的统一依据。",
    )
    insert_paragraph_after(
        cursor,
        "文档严格以《需求文档.md》《需求补充.md》、SRS 分析资料和现有两版需求规格说明书草稿为基础，明确系统范围、角色边界、功能要求、非功能约束、业务规则、交付要求和待确认事项，避免将讨论性建议误写为刚性承诺。",
    )


def build_12(heading: Paragraph) -> None:
    add_bullets(
        heading,
        [
            "信息学院团委老师、班主任、学院领导和业务负责人，用于确认业务边界、审批口径、权限规则与验收标准。",
            "产品、架构、前后端、数据库和测试人员，用于据此开展设计、编码、联调、测试和交付。",
            "运维和部署人员，用于理解 Kingbase 基线、文件流转机制、安全控制和上线前提。",
            "课程评审和答辩人员，用于快速掌握项目目标、实现范围、分析深度和交付质量。",
        ],
    )


def build_13(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(heading, f"项目名称：{PROJECT}")
    cursor = insert_paragraph_after(cursor, "项目定位：面向信息学院内部的一站式学生综合服务与党团管理平台，以减少重复答疑、沉淀规范流程、线上化高频办事、强化通知触达和形成客观留痕为核心目标。")
    cursor = insert_paragraph_after(cursor, "用户单位：信息学院团委及学生工作相关部门。")
    cursor = insert_paragraph_after(cursor, "开发单位：课程项目研发团队。")
    cursor = insert_paragraph_after(cursor, "适用范围：一期仅面向信息学院，不面向全校开放，不依托“微人大”生态，不默认调用校级 API。")
    insert_paragraph_after(cursor, "服务规模：管理端约 10 人，总使用规模约 1200 人；学生端高频场景优先面向移动端，管理端优先面向 PC。")


def build_14(heading: Paragraph) -> None:
    add_bullets(
        heading,
        [
            "第 1 章说明文档目标、读者对象、项目背景、术语和参考资料。",
            "第 2 章描述系统定位、外部环境、限制约束以及实施假设与前提。",
            "第 3 章描述功能概述、用例模型和分析模型，并给出正文级的功能追溯摘要。",
            "第 4 章描述性能要求、设计约束、界面要求、进度要求、交付要求和验收要求。",
            "第 5 章提供关键页面的线框图级文本原型说明。",
            "附录用于收口默认业务决策、权限矩阵、审批矩阵和待确认问题，便于详细设计阶段继续对齐。",
        ],
    )


def build_15(heading: Paragraph) -> None:
    intro = insert_paragraph_after(heading, "本文档中使用的关键术语如下：")
    table = insert_table_after(
        intro,
        ["术语", "说明"],
        [
            ["SRS", "Software Requirements Specification，软件需求规格说明书，是项目范围、约束和验收依据的正式文档。"],
            ["FR", "Functional Requirement，功能需求，描述系统必须具备的能力。"],
            ["NFR", "Non-Functional Requirement，非功能需求，描述性能、安全、可靠性等质量要求。"],
            ["Kingbase", "人大金仓数据库，是本项目唯一允许的生产数据库基线。"],
            ["受控 AI 匹配", "以标准答案、官方链接和知识条目为依据的受控问答机制，不允许无来源自由生成。"],
            ["正式流程边界", "学院平台与学校正式生效系统之间的职责边界；涉及校级正式生效的事项，平台可仅承担说明、预检、归档或跟踪。"],
            ["弱结论", "仅提供风险提示、缺口展示和人工核验建议，不输出毕业资格、课程替代或学分满足的强判定。"],
            ["离线文件流转", "在无校级 API 条件下，通过 Excel、Word、PDF 完成导入导出和数据交换的受控机制。"],
            ["审计日志", "记录审批、配置、导入导出和敏感访问等关键动作的可追溯日志。"],
        ],
    )
    insert_paragraph_after_table(table, "上述术语与缩写在全文中保持统一含义。")


def build_16(heading: Paragraph) -> None:
    add_bullets(
        heading,
        [
            "《需求文档.md》",
            "《需求补充.md》",
            "docs/srs/00-business-context.md",
            "docs/srs/01-customer-problems.md",
            "docs/srs/02-software-glance.md",
            "docs/srs/03-customer-needs.md",
            "docs/srs/04-software-vision.md",
            "docs/srs/traceability-matrix.md",
            "docs/srs/functional-requirements/_index.md 及各 FR 文件",
            "docs/srs/non-functional-requirements/_index.md 及各 NFR 文件",
            "docs/pending-business-decisions.md",
            "specs/001-student-service-platform/spec.md 及其分析、用例、UI/UX 补充文档",
            "output/doc/软件需求规格说明书-信息学院学生综合服务与党团管理平台-v1.0.docx",
            "学院学生综合服务与党团管理平台-软件需求规格说明书-v1.0.docx",
        ],
    )


def build_21(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(
        heading,
        "本系统定位为信息学院内部使用的学生综合服务与党团管理平台，采用学生端与管理端分离、前后端分离的体系结构。学生端优先支持微信小程序或等价移动 Web 场景，管理端优先支持 PC 管理后台。",
    )
    cursor = insert_paragraph_after(
        cursor,
        "系统服务边界以信息学院为单位，主要支撑政策咨询、知识模板下载、党团流程跟踪、常见事务申请与审批、通知汇聚与分发、文件导入导出、学业风险提示和审计留痕，不替代学校正式生效系统。",
    )
    table = insert_table_after(
        cursor,
        ["外部对象/系统", "关系说明"],
        [
            ["普通学生", "通过移动端或浏览器访问系统，完成政策查询、进度查看、申请提交、通知接收和学业风险查看。"],
            ["团支书/党支部书记/班团骨干", "在组织范围内参与流程推进、节点处理、信息协同和部分业务辅助。"],
            ["班主任/审批老师/团委老师", "负责审批、数据维护、通知发布、知识治理、导入导出和统计查看。"],
            ["学院领导", "以汇总统计、关键结果和例外情况查看为主，不直接承担常规维护。"],
            ["超级管理员", "负责权限、模板、标签、流程、审计和系统配置维护。"],
            ["学院提供的数据文件", "通过 Excel、Word、PDF 作为学生主数据、政策资料、模板资产和导入导出的主要媒介。"],
            ["官方网站/公众号/邮件或短信服务", "作为受控通知来源和消息分发渠道，但公众号不默认采用非授权爬取方式。"],
            ["校级正式系统", "仅作为边界参照或跳转目标；在未确认授权前，不默认对接其接口。"],
        ],
    )
    cursor = insert_paragraph_after_table(table, "系统与外部环境的关系如图所示：")
    cursor = insert_picture_after(
        cursor,
        DIAG / "system-context.png",
        5.6,
        "图 2-1 软件产品与外部环境关系图",
        page_break_before=False,
    )
    add_bullets(
        cursor,
        [
            "学院侧生态独立建设，不依托“微人大”现有生态，但可参考其流程、权限和日志设计思路。",
            "学生侧场景以“查、提、看状态”为主，管理侧场景以“筛、审、导、查日志”为主。",
            "全量需求范围完整覆盖原始需求文档的 5 大模块，技术拆解主线为知识、流程、审批、通知和审计 5 个闭环。",
        ],
    )


def build_22(heading: Paragraph) -> None:
    intro = insert_paragraph_after(heading, "系统设计和实施必须同时满足以下限制与约束：")
    table = insert_table_after(
        intro,
        ["约束类别", "约束说明"],
        [
            ["架构约束", "系统必须采用前端与后端分离架构；学生端与管理端的交互规则保持一致。"],
            ["数据库约束", "生产数据库必须使用 Kingbase，表结构、SQL、索引和迁移脚本均需兼容。"],
            ["集成约束", "在甲方未明确开放接口前，禁止假设存在校级 API；与外部系统的数据交换默认依赖 Excel、Word、PDF 离线流转。"],
            ["范围约束", "一期范围仅限信息学院，完整覆盖需求文档中的全部功能点，但不直接替代校级正式生效链路。"],
            ["内容治理约束", "政策问答必须优先基于标准答案、官方来源和文件链接；模糊场景必须转人工。"],
            ["安全约束", "身份证号、联系方式、生源地、户籍地、导师、处分记录等敏感字段必须加密存储、受控展示、受控导出并记录访问日志。"],
            ["审计约束", "审批、配置、导入导出、通知发送、权限变化和敏感访问必须保留可查询日志，留存期不少于一个学期。"],
            ["可靠性约束", "关键状态变更和主数据导入不得出现部分生效或静默丢失；错误批次需整批回滚并输出错误报告。"],
            ["文件约束", "政策文档与模板以 Word/PDF 为主，结构化数据以 Excel 为主，单文件建议支持约 30MB。"],
            ["业务约束", "学业模块只能输出弱结论，不得做毕业资格、课程替代或学分满足的强判定。"],
        ],
    )
    cursor = insert_paragraph_after_table(table, "实现和验收时需特别注意以下默认边界：")
    add_bullets(
        cursor,
        [
            "学生账号默认不得导出学院级原始数据。",
            "助教原则上不拥有正式审批权限。",
            "短信通道属于可配置能力，不作为一期强制前提；站内通知和邮件优先。",
            "未经授权的公众号抓取不纳入默认实现方案。",
        ],
    )


def build_23(heading: Paragraph) -> None:
    add_bullets(
        heading,
        [
            "学院能够提供一期所需的基础学生主数据、部分政策文件、模板文件和样例流程规则。",
            "角色账号可基于学号或教工号形成统一身份主体，再通过角色/岗位授予权限；是否绑定微信实名作为实现细节待确认。",
            "在校级正式流程边界尚未全部明确前，相关事项默认按“说明、预检、归档、跟踪”边界处理，不宣称学院平台直接形成校级正式生效结果。",
            "知识库、模板和通知来源需要由学院业务方持续维护，默认遵循“谁上传谁维护”的运营原则。",
            "若培养方案、开课信息和成绩规则暂未完全结构化，系统仍需交付学业分析模块、规则维护能力和样例数据演示，并保持弱结论边界。",
        ],
    )


def build_31(heading: Paragraph) -> None:
    intro = insert_paragraph_after(
        heading,
        "本节严格按照《需求文档.md》给出的五大功能模块展开，并将《需求补充.md》中的边界、约束、默认方案和实现前提并入同一套描述。任何后续技术拆解都不得被理解为对原始五大模块的删减或替代。",
    )
    cursor = insert_paragraph_after(intro, "模块一：智能问答与政策知识库", bold=True)
    cursor = add_bullets(
        cursor,
        [
            "智能咨询：录入学院/学校常用政策文件，围绕奖助学金、休学复学、宿舍调整、查档调档、节假日放假等主题提供关键词检索与 AI 匹配回复。",
            "回复原则：优先基于标准答案、权威文件和官方链接返回结果，避免自由生成式“幻觉”；敏感信息不直接展示，模糊场景必须转人工。",
            "模板下载：提供常用证明、请假条、活动预算表、简报等标准 Word/Excel 模板下载。",
            "知识维护：老师侧必须支持持续上传、停用、更新知识条目、来源文件和模板资产，并记录版本、维护人和更新时间。",
            "范围说明：本模块既覆盖学生自助查问，也覆盖管理员对知识来源、模板与标准答案的治理能力。",
        ],
    )
    cursor = insert_paragraph_after(cursor, "模块二：党团事务流程管理", bold=True)
    cursor = add_bullets(
        cursor,
        [
            "线性流程可视化：展示入党、入团等流程的标准阶段链路，例如入党申请人、积极分子、发展对象、预备党员、正式党员。",
            "进度追踪：学生可查询当前所处阶段、已完成动作、下一步要求和应提交材料。",
            "关键节点提醒：系统根据节点时间或相对时长自动触发提醒，并对异常超期形成待办。",
            "理论自测：支持导入官方题库，允许学生进行党建知识自主测试并留存成绩记录。",
            "组织协同：团支书、党支部书记、团委老师和学院领导可在权限范围内查看、审核和推进流程节点。",
        ],
    )
    cursor = insert_paragraph_after(cursor, "模块三：信息集成与精准推送", bold=True)
    cursor = add_bullets(
        cursor,
        [
            "信息聚类：支持对学校后勤、保卫处、就业办等部门的官方通知进行手工录入、合规汇聚或公开来源采集。",
            "标签化分发：按照“就业”“实习”“计算机类”等标签，以及年级、专业、毕业生身份等画像圈定目标群体。",
            "多渠道通知：支持通过小程序站内消息、邮件或短信向目标群体批量推送信息，并记录发送与接收状态。",
            "通知治理：必须区分官方渠道与非官方渠道，保证来源可信、口径清晰、可回查。",
        ],
    )
    cursor = insert_paragraph_after(cursor, "模块四：电子证明生成与审批流程", bold=True)
    cursor = add_bullets(
        cursor,
        [
            "证明自动生成：系统基于学生数据库字段和标准模板自动填充证明内容，生成 PDF 预览或待审批文件。",
            "电子审批流：数字化学院内部审批流程，覆盖请假申请、盖章申请、开具证明、活动预算表等场景。",
            "材料与附件：盖章场景原则上要求上传待盖章文件；证明场景支持上传附件或填写文本说明；涉密场景允许转线下处理但必须留痕。",
            "驳回、撤回与重提：被驳回后应在原表单基础上修改重提；允许设置撤回期限，并支持老师撤回审批结果后重批。",
        ],
    )
    cursor = insert_paragraph_after(cursor, "模块五：学业情况分析与预警", bold=True)
    cursor = add_bullets(
        cursor,
        [
            "培养方案比对：录入不同年级、专业、版本的培养方案规则，作为学业分析基线。",
            "成绩分析：学生上传成绩单或由学院导入成绩数据后，系统对比培养方案，展示未修满学分模块和风险点。",
            "选课建议：根据缺失模块给出本学期开设相关课程的类型级建议，不涉及选课竞争分析。",
            "责任边界：本模块只提供辅助展示、风险提示与人工复核建议，不输出毕业资格、课程替代或强结论。",
        ],
    )
    cursor = insert_paragraph_after(cursor, "为支撑上述五大模块，系统还必须提供以下共性支撑能力：", bold=True)
    cursor = add_bullets(
        cursor,
        [
            "统一身份、角色与字段级权限控制，支持学生、班团骨干、班主任、团委老师、学院领导和超级管理员分层使用。",
            "Excel/Word/PDF 导入导出、模板化校验、错误报告、整批回滚和审计留痕能力。",
            "知识来源、通知来源、模板资产、流程配置和规则配置的后台维护能力。",
            "高敏字段脱敏、访问日志、审批日志、通知日志和导出日志等全流程审计能力。",
        ],
    )
    table = insert_table_after(
        cursor,
        ["模块编号", "模块名称", "主要功能点", "主要用户", "对应 FR"],
        [
            ["M-01", "智能问答与政策知识库", "政策咨询、权威回复、模板下载、知识来源维护", "学生、老师、管理员", "FR-001、FR-002、FR-003"],
            ["M-02", "党团事务流程管理", "流程可视化、进度追踪、节点提醒、理论自测", "学生、团支书、党支部书记、团委老师", "FR-004、FR-005"],
            ["M-03", "信息集成与精准推送", "官方信息汇聚、标签化分发、多渠道通知", "管理员、学生", "FR-010、FR-011"],
            ["M-04", "电子证明生成与审批流程", "证明生成、PDF 预览、请假/盖章/预算表审批、驳回重提", "学生、班主任、审批老师、学院领导", "FR-006、FR-007、FR-008"],
            ["M-05", "学业情况分析与预警", "培养方案比对、成绩分析、选课类型建议、弱结论提示", "学生、业务负责人、学院领导", "FR-014、FR-015"],
        ],
    )
    cursor = insert_paragraph_after_table(table, "在开发拆解层面，上述五大模块进一步分解为下列 FR。支撑型 FR 不构成新的业务模块，而是服务于五大模块落地。")
    fr_table = insert_table_after(
        cursor,
        ["FR", "标题", "对应能力", "关键规则摘要"],
        [
            ["FR-001", "政策与流程查询", "模块一", "学生可查询政策、流程、资格和材料说明；查询结果必须清晰可读。"],
            ["FR-002", "权威答复治理", "模块一", "受控 AI 匹配必须基于标准答案或官方链接；模糊场景转人工。"],
            ["FR-003", "知识与模板维护及下载", "模块一", "老师可上传、停用、更新知识条目和模板；学生可下载标准模板，系统需保留版本信息。"],
            ["FR-004", "党团进度查看", "流程跟踪", "学生可查看当前阶段、历史节点和下一动作。"],
            ["FR-005", "党团提醒与理论自测", "模块二", "关键节点提醒可管理、可追踪、可形成异常待办；支持导入官方题库并记录学生自测结果。"],
            ["FR-006", "电子证明生成与常见事务在线提交", "模块四", "系统应能用学生数据填充标准模板生成 PDF 预览，并支持请假、盖章、证明等事项在线提交。"],
            ["FR-007", "申请审核工作台", "申请与审批", "审批人应在单一工作台查看申请详情、附件、历史和当前状态。"],
            ["FR-008", "驳回撤回与重提规则", "申请与审批", "驳回后允许保留原表单修改重提；支持限时撤回和重批。"],
            ["FR-009", "文件导入导出", "资料管理与审计", "导入导出需模板化校验、错误报告、整批回滚和日志留痕。"],
            ["FR-010", "官方信息聚合、标签与目标人群管理", "模块三", "通知应支持官方来源录入或合规汇聚、标签分类和按画像圈定目标人群。"],
            ["FR-011", "通知发送与接收记录", "模块三", "站内、邮件和短信的发送结果与接收状态需可回查。"],
            ["FR-012", "角色与字段级权限控制", "治理基线", "遵循最小权限原则，高敏字段默认脱敏或受控访问。"],
            ["FR-013", "审计日志跟踪", "治理基线", "关键动作必须记录操作者、时间、对象和结果，且学期内可查询。"],
            ["FR-014", "培养方案比对与学业缺口展示", "模块五", "系统应对比培养方案与成绩结果，展示未修满模块、学分缺口和风险提示，不输出毕业资格结论。"],
            ["FR-015", "培养方案规则维护与选课类型建议", "模块五", "支持按专业、年级、版本维护规则，并基于缺口给出课程类型级建议。"],
            ["FR-016", "学院运营统计看板", "统计看板", "汇总申请量、审批状态、通知触达、党团记录等运营数据。"],
        ],
    )
    cursor = insert_paragraph_after_table(fr_table, "功能追溯摘要如下：")
    add_bullets(
        cursor,
        [
            "需求文档中的五大模块已在正文中逐项展开，且每个模块都映射到明确的 FR、验收条目和原型页面。",
            "资料导入导出、角色权限、字段脱敏和审计日志属于全局支撑能力，服务于五大模块，不应被误读为替代业务模块。",
            "《需求补充.md》中的默认方案、边界条件和待确认事项已被吸收进对应模块说明和附录，避免遗漏原始需求。",
        ],
    )


def build_32(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(
        heading,
        "系统的主要参与者包括学生、团支书 / 党支部书记、班主任 / 审批老师、团委老师、超级管理员和学院领导。为减少单图跨域连线和提升传统 SRS 阅读性，用例模型拆分为学生侧与管理侧两张全局用例图，但其覆盖范围仍然严格对应需求文档中的五大模块。",
    )
    cursor = insert_picture_after(
        cursor,
        DIAG / "use-case-student.png",
        5.6,
        "图 3-1 学生侧全局用例图",
        page_break_before=False,
        page_break_after=False,
    )
    cursor = insert_picture_after(
        cursor,
        DIAG / "use-case-admin.png",
        5.6,
        "图 3-2 管理侧全局用例图",
        page_break_before=False,
        page_break_after=False,
    )
    summary = insert_table_after(
        cursor,
        ["模块", "代表用例", "主要参与者", "覆盖内容"],
        [
            ["模块一", "查询政策问答并下载模板", "学生、管理员", "政策咨询、权威回复、模板下载、知识维护"],
            ["模块二", "查看党团进度并完成理论自测", "学生、团支书、团委老师", "流程可视化、节点提醒、理论自测"],
            ["模块三", "汇聚官方通知并精准推送", "管理员、学生", "信息聚类、标签分发、多渠道通知"],
            ["模块四", "提交证明/请假申请并进入审批", "学生、审批老师", "电子证明、在线审批、驳回重提"],
            ["模块五", "上传成绩数据并查看学业缺口", "学生、业务负责人", "培养方案比对、成绩分析、课程建议"],
        ],
    )
    cursor = insert_paragraph_after_table(summary, "以下选取五个代表性用例，分别对应五大模块的核心用户旅程。")
    cursor = insert_paragraph_after(cursor, "代表用例一：查询政策问答并下载模板")
    table1 = insert_table_after(
        cursor,
        ["项目", "内容"],
        [
            ["用例名称", "查询政策问答并下载模板"],
            ["主要参与者", "学生"],
            ["次要参与者", "管理员 / 团委老师"],
            ["前置条件", "知识库已导入政策条目、标准答案、来源链接和模板资产；学生已登录并具备查询权限。"],
            ["后置条件", "学生获得权威答复、来源依据和可下载模板；系统记录查询关键词和点击行为用于优化。"],
            ["相关需求", "FR-001、FR-002、FR-003"],
        ],
    )
    cursor = insert_paragraph_after_table(table1, "主事件流：")
    flow1 = insert_table_after(
        cursor,
        ["步骤", "参与者", "系统响应"],
        [
            ["1", "学生", "输入关键词或选择常见问题分类。"],
            ["2", "系统", "检索标准答案、政策来源和关联模板，返回候选结果。"],
            ["3", "学生", "查看答案详情、来源链接、更新时间和适用条件。"],
            ["4", "系统", "展示模板下载入口；若问题敏感或模糊，则提示转人工。"],
            ["5", "学生", "下载模板或进入人工咨询入口。"],
        ],
    )
    cursor = insert_paragraph_after_table(flow1, "异常 / 备选分支：")
    alt1 = insert_table_after(
        cursor,
        ["编号", "场景", "处理方式"],
        [
            ["K1", "未检索到明确知识条目", "返回相近问题、官方链接和人工咨询提示，不生成自由结论。"],
            ["K2", "知识条目已过期或停用", "系统不再作为当前口径展示，仅提示查看最新版本或联系老师。"],
        ],
    )
    cursor = insert_paragraph_after_table(alt1, "代表用例二：查看党团进度并完成理论自测")
    table2 = insert_table_after(
        cursor,
        ["项目", "内容"],
        [
            ["用例名称", "查看党团进度并完成理论自测"],
            ["主要参与者", "学生"],
            ["次要参与者", "团支书 / 党支部书记、团委老师"],
            ["前置条件", "学生已存在有效党团阶段记录；节点规则、提醒规则和题库已配置。"],
            ["后置条件", "学生获知当前阶段、下一动作和提醒信息；如进行了自测，则系统留存答题结果。"],
            ["相关需求", "FR-004、FR-005"],
        ],
    )
    cursor = insert_paragraph_after_table(table2, "主事件流：")
    flow2 = insert_table_after(
        cursor,
        ["步骤", "参与者", "系统响应"],
        [
            ["1", "学生", "进入“我的党团进度”页面。"],
            ["2", "系统", "展示当前阶段、历史动作、下一节点要求和关键时间提醒。"],
            ["3", "学生", "查看需提交材料，或点击理论自测入口。"],
            ["4", "系统", "加载题库并提供答题界面，提交后返回分数和错题分析。"],
            ["5", "组织角色 / 团委老师", "在管理端查看提醒触发情况、自测完成情况和节点异常。"],
        ],
    )
    cursor = insert_paragraph_after_table(flow2, "异常 / 备选分支：")
    alt2 = insert_table_after(
        cursor,
        ["编号", "场景", "处理方式"],
        [
            ["D1", "学生暂无党团档案", "系统提示当前无可展示进度，并提供咨询或建档引导。"],
            ["D2", "题库未配置", "保留理论自测入口位但提示暂未开放，不影响进度查询主流程。"],
        ],
    )
    cursor = insert_paragraph_after_table(alt2, "代表用例三：汇聚官方通知并精准推送")
    table3 = insert_table_after(
        cursor,
        ["项目", "内容"],
        [
            ["用例名称", "汇聚官方通知并精准推送"],
            ["主要参与者", "管理员"],
            ["次要参与者", "学生"],
            ["前置条件", "管理员具备通知治理权限；通知来源、标签规则和目标人群画像字段已配置。"],
            ["后置条件", "目标学生收到通知并可回查来源、标签、发送时间和渠道状态。"],
            ["相关需求", "FR-010、FR-011"],
        ],
    )
    cursor = insert_paragraph_after_table(table3, "主事件流：")
    flow3 = insert_table_after(
        cursor,
        ["步骤", "参与者", "系统响应"],
        [
            ["1", "管理员", "录入或汇聚官方通知，并标注来源部门和发布时间。"],
            ["2", "系统", "支持添加就业、实习、年级、毕业生等标签。"],
            ["3", "管理员", "按标签和画像圈定目标人群，选择站内、邮件或短信渠道。"],
            ["4", "系统", "执行发送任务并记录每个渠道的发送结果与接收状态。"],
            ["5", "学生", "在通知中心查看详情、来源依据和接收时间。"],
        ],
    )
    cursor = insert_paragraph_after_table(flow3, "异常 / 备选分支：")
    alt3 = insert_table_after(
        cursor,
        ["编号", "场景", "处理方式"],
        [
            ["N1", "来源不属于官方渠道", "系统要求补充来源依据或标记为草稿，不允许直接群发。"],
            ["N2", "短信通道未启用", "系统保留站内和邮件发送，并记录短信未启用原因。"],
        ],
    )
    cursor = insert_paragraph_after_table(alt3, "代表用例四：提交证明/请假申请并进入审批")
    table4 = insert_table_after(
        cursor,
        ["项目", "内容"],
        [
            ["用例名称", "提交证明/请假申请并进入审批"],
            ["主要参与者", "学生"],
            ["次要参与者", "班主任 / 审批老师、团委老师、学院授权出件人"],
            ["前置条件", "学生完成身份认证并进入本人业务范围；证明模板、审批角色和边界提示已配置。"],
            ["后置条件", "生成可追踪申请记录；审批意见、状态变化和附件信息全部留痕；通过后生成预览或出件结果。"],
            ["相关需求", "FR-006、FR-007、FR-008"],
        ],
    )
    cursor = insert_paragraph_after_table(table4, "主事件流：")
    flow4 = insert_table_after(
        cursor,
        ["步骤", "参与者", "系统响应"],
        [
            ["1", "学生", "进入申请页面并选择证明、请假或盖章事项。"],
            ["2", "系统", "自动带出基础字段、模板、附件要求和 PDF 预览。"],
            ["3", "学生", "补充说明、上传附件或确认预览内容后提交。"],
            ["4", "审批人", "在工作台查看详情、附件、历史动作并执行通过或驳回。"],
            ["5", "系统", "更新状态、记录日志，并在通过后输出出件结果或下载入口。"],
        ],
    )
    cursor = insert_paragraph_after_table(flow4, "异常 / 备选分支：")
    alt4 = insert_table_after(
        cursor,
        ["编号", "场景", "处理方式"],
        [
            ["P1", "附件缺失或内容不合规", "审批人驳回并填写补充说明；系统保存驳回意见和表单快照；学生修改后重提。"],
            ["P2", "事项必须走校级正式链路", "系统展示“仅预检 / 归档 / 跟踪”边界提示，并引导至正式办理渠道。"],
            ["P3", "涉密内容不适合线上处理", "审批人标记转线下处理；系统保留申请和审批记录，但不生成在线文件。"],
        ],
    )
    cursor = insert_paragraph_after_table(alt4, "代表用例五：上传成绩数据并查看学业缺口")
    table5 = insert_table_after(
        cursor,
        ["项目", "内容"],
        [
            ["用例名称", "上传成绩数据并查看学业缺口"],
            ["主要参与者", "学生"],
            ["次要参与者", "业务负责人 / 管理员"],
            ["前置条件", "培养方案规则已导入；成绩单可由学生上传或由学院导入；课程基础数据可用。"],
            ["后置条件", "系统展示培养方案比对结果、未修满模块、风险提示和课程类型级建议。"],
            ["相关需求", "FR-014、FR-015"],
        ],
    )
    cursor = insert_paragraph_after_table(table5, "主事件流：")
    flow5 = insert_table_after(
        cursor,
        ["步骤", "参与者", "系统响应"],
        [
            ["1", "学生", "上传成绩单或进入个人学业分析页面。"],
            ["2", "系统", "解析成绩记录并与培养方案规则进行比对。"],
            ["3", "系统", "高亮缺失模块、风险点和尚未满足的课程类型。"],
            ["4", "学生", "查看课程类型级建议，并根据提示转人工咨询。"],
            ["5", "业务负责人", "维护规则版本、导入样例数据并查看异常分析结果。"],
        ],
    )
    cursor = insert_paragraph_after_table(flow5, "异常 / 备选分支：")
    insert_table_after(
        cursor,
        ["编号", "场景", "处理方式"],
        [
            ["A1", "成绩单解析失败或字段缺失", "系统提示改用标准模板、人工补录或联系管理员导入。"],
            ["A2", "课程存在复杂替代、缓修或免修规则", "系统不给出强结论，只提示人工核验路径。"],
        ],
    )


def build_33(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(heading, "分析模型从静态结构和动态交互两个视角描述核心业务对象及其关系，并确保五大模块在数据对象层与交互层均可追溯。")
    cursor = insert_picture_after(
        cursor,
        DIAG / "class-diagram.png",
        5.55,
        "图 3-3 核心业务分析类图",
        page_break_before=False,
        page_break_after=False,
    )
    entity_table = insert_table_after(
        cursor,
        ["实体类别", "代表实体", "说明"],
        [
            ["主数据", "StudentProfile", "承载学号、姓名、年级、班级和政治面貌等基础信息。"],
            ["知识治理", "KnowledgeEntry / KnowledgeSource / TemplateAsset", "支撑政策问答、来源版本和模板下载。"],
            ["党团流程", "PartyMemberStatus / PartyWorkflowNode / PartyWorkflowEvent", "描述阶段状态、节点定义和事件时间线。"],
            ["申请审批", "CommonRequest / Attachment / ApprovalTask / ApprovalAction", "承载申请单、附件、审批任务和动作留痕。"],
            ["通知管理", "NoticeMessage / NoticeTargetRule / NoticeDelivery", "支撑通知圈选、分发和接收记录。"],
            ["学业分析", "CurriculumRule / TranscriptSnapshot / AcademicGap / CourseSuggestion", "支撑培养方案比对、成绩分析、学分缺口提示和课程类型建议。"],
            ["导入导出与审计", "ImportBatch / ImportBatchRow / DocumentAuditLog", "承载批量交换、错误处理和可追溯日志。"],
        ],
    )
    cursor = insert_paragraph_after_table(entity_table, "关键动态交互采用请假申请与后台审核作为典型时序进行说明。")
    cursor = insert_picture_after(
        cursor,
        DIAG / "sequence-diagram.png",
        5.6,
        "图 3-4 学生提交请假申请至后台审核时序图",
        page_break_before=False,
        page_break_after=False,
    )
    add_bullets(
        cursor,
        [
            "后端统一维护业务状态、权限判断、事务一致性和数据持久化，不将关键规则下沉到前端。",
            "知识问答、党团流程、通知推送、审批流转和学业分析都必须以结构化对象和审计日志形成闭环。",
            "分析模型与五大模块及其 FR 保持一致，可直接作为数据库设计和接口拆分的基础。",
        ],
    )


def build_41(heading: Paragraph) -> None:
    add_bullets(
        heading,
        [
            "系统需支持并发用户数不少于 50 人，吞吐能力不低于 50 TPS，满足信息学院低并发但多任务并行的日常使用场景。",
            "在 50 并发用户基准下，95% 的学生端常见请求（政策查询、党团进度查看、申请提交）响应时间应小于 3 秒。",
            "在 50 并发用户基准下，95% 的管理端常见列表页（审批工作台、通知批次、日志查询、统计看板）响应时间应小于 5 秒。",
            "标准格式 Excel 主数据 100 条批量导入应在 60 秒内完成成功提交或整批失败回滚，并输出可下载错误报告。",
            "学生在受训条件下完成 1 个标准事务提交的操作时间应不超过 5 分钟；审批老师在材料完整场景下完成 1 条审核的主要操作时间应不超过 2 分钟。",
        ],
    )


def build_42(heading: Paragraph) -> None:
    intro = insert_paragraph_after(heading, "除第 2 章所述总体约束外，本项目在设计和实现阶段还需满足以下细化要求：")
    table = insert_table_after(
        intro,
        ["设计主题", "要求"],
        [
            ["事务一致性", "申请状态、党团节点、导入批次和学业规则更新不得出现部分生效或静默丢失。"],
            ["审计留存", "关键审批、配置、导出和敏感访问记录至少保留 1 个学期，并支持按角色、时间和动作类型检索。"],
            ["权限控制", "学生默认只看本人；班主任默认看职责范围；高敏字段默认脱敏；高敏导出需额外授权。"],
            ["内容治理", "知识条目必须携带来源、版本和更新时间；停用或过期内容不得继续作为当前有效口径对外展示。"],
            ["文件交换", "导入模板需固定字段结构；失败批次需提供记录级错误报告和重试依据。"],
            ["正式边界", "凡涉及校级正式流程边界的页面，必须显示风险提示和正式渠道引导文案。"],
            ["学业边界", "学业分析页面必须显式提示“仅供辅助参考”，不得替代学院或学校最终审核。"],
            ["易用性", "审批场景应在单一工作台内完成主要判断动作，避免跨页反复查找附件和历史。"],
        ],
    )
    insert_paragraph_after_table(table, "上述约束与 NFR-001 至 NFR-005 保持一致，是设计评审、测试和验收的底线要求。")


def build_43(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(heading, "界面设计分为学生侧小程序端和教师 / 管理员侧 PC 管理后台两类。")
    cursor = insert_paragraph_after(cursor, "小程序端总体风格：采用正式稳重的学院风格，主色建议使用深绛红，辅以深蓝灰和浅灰白底色；底部采用“首页 / 服务 / 进度 / 我的”四栏导航；页面强调卡片化、少层级和强状态反馈。")
    cursor = insert_paragraph_after(cursor, "PC 管理后台总体风格：采用左侧主导航、顶部工具栏和内容工作区布局，主色延续学院红，导航区域采用深蓝灰；适合审批、筛选、导入导出和日志查询等复杂任务。")
    table = insert_table_after(
        cursor,
        ["端别", "色彩主题", "导航结构", "设计重点"],
        [
            ["小程序端", "深绛红 + 浅灰白 + 深蓝灰", "底部 4 栏导航", "快速查询、快速提交、查看状态、低认知负担"],
            ["PC 管理后台", "深红 + 深蓝灰 + 白底内容区", "左侧一级导航 + 顶部工具栏", "高效筛选、审批、批量导出、审计追踪"],
        ],
    )
    cursor = insert_paragraph_after_table(table, "关键界面约束如下：")
    add_bullets(
        cursor,
        [
            "学生端必须显式展示“待审核、待补材料、进行中、已完成”等状态标签。",
            "PC 端所有导出和审批页面都必须在操作前展示权限范围或导出摘要。",
            "与校级正式流程边界相关的页面必须显示风险提示和引导文案。",
            "高敏字段默认脱敏展示，不得仅依赖颜色区分敏感状态。",
        ],
    )


def build_44(heading: Paragraph) -> None:
    intro = insert_paragraph_after(heading, "一期开发周期建议按 12 周执行，以知识库、流程、审批、通知、审计五个核心闭环为主线推进，同时完整纳入原始需求中的智能问答、理论自测、通知汇聚、证明预览和学业分析能力。")
    table = insert_table_after(
        intro,
        ["里程碑", "周期", "目标"],
        [
            ["M1 启动与基线冻结", "第 1-2 周", "冻结全量需求范围、角色权限、审批边界、Kingbase 和离线流转基线"],
            ["M2 知识库闭环完成", "第 3-4 周", "完成政策查询、模板下载、内容治理和来源版本化"],
            ["M3 流程与审批闭环完成", "第 5-7 周", "完成党团流程跟踪、常见事务提交、审批工作台和驳回重提"],
            ["M4 通知与审计闭环完成", "第 8-9 周", "完成标签化通知、导入导出、审计留痕和权限规则验证"],
            ["M5 联调与上线验收", "第 10-12 周", "完成 Kingbase 环境验证、部署、验收测试和交付整理"],
        ],
    )
    cursor = insert_paragraph_after_table(table, "里程碑通过要求如下：")
    add_bullets(
        cursor,
        [
            "M1 必须明确前后端分离和无校级 API 的前提。",
            "M2 必须完成学生自助查询、受控智能问答、模板下载和转人工提示闭环。",
            "M3 必须支持至少 3 类事务申请、证明 PDF 预览、理论自测和完整审批流转。",
            "M4 必须完成 Excel 导入、官方通知汇聚、多渠道发送和审计日志校验。",
            "M5 必须在 Kingbase 环境通过全量需求模块验收标准。",
        ],
    )


def build_45(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(heading, "项目最终交付应覆盖源码、数据库脚本、API 文档、部署手册和验收资料。")
    cursor = insert_paragraph_after(cursor, "源码交付：包括后端源码、Web 管理端源码、学生端源码、共享契约和配置样例。")
    cursor = insert_paragraph_after(cursor, "数据库脚本交付：包括 Kingbase 建库脚本、建表脚本、迁移脚本、初始化字典脚本、测试数据脚本和回滚脚本。")
    cursor = insert_paragraph_after(cursor, "API 文档交付：包括鉴权接口、知识库接口、流程接口、审批接口、通知接口、审计接口以及导入导出协议文档。")
    cursor = insert_paragraph_after(cursor, "部署手册交付：包括环境准备、Kingbase 部署说明、应用部署说明、安全配置说明、文件流转操作手册和运维巡检手册。")
    add_bullets(
        cursor,
        [
            "所有电子资料应以可交付的电子文件形式提供。",
            "数据库脚本应与交付版本严格对应，可重建测试环境。",
            "部署手册必须覆盖从空环境到可运行环境的完整步骤。",
            "建议同时提交验收测试报告、演示脚本、字段权限矩阵和审批矩阵最终版。",
        ],
    )


def build_46(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(heading, "项目验收必须在 Kingbase 测试环境完成，并以五大核心模块的功能和非功能结果作为判定依据。")
    cursor = insert_paragraph_after(cursor, "模块一验收：成功导入不少于 50 条政策知识条目；随机抽检 10 条内容可正确展示官方来源、版本和更新时间；模板下载可正常打开且无损坏；模糊场景必须有转人工提示；受控 AI 匹配结果不得脱离标准答案或官方链接。")
    cursor = insert_paragraph_after(cursor, "模块二验收：预置不少于 30 名学生的党团阶段数据后，学生均可正确查看当前阶段、已完成节点、下一动作和关键节点提醒；理论自测在导入官方题库后可完成答题、计分和成绩留存。")
    cursor = insert_paragraph_after(cursor, "模块三验收：管理员创建不少于 10 条带标签通知并汇聚不少于 20 条官方通知来源后，可按画像圈定目标人群；向不少于 200 名样例学生发送时，非目标人群误投率应为 0；目标学生可在通知中心查看来源、标签和接收结果，邮件和短信通道在启用时必须完整留痕。")
    cursor = insert_paragraph_after(cursor, "模块四验收：至少 3 类常见事务申请可成功提交并生成唯一编号；证明场景可调用学生字段填充标准模板并生成 PDF 预览；审批老师可查看申请详情、附件和历史动作；对不少于 20 条测试申请执行通过和驳回后，状态与审批意见保持一致；至少 5 条驳回申请可基于原表单修改并重提。")
    cursor = insert_paragraph_after(cursor, "模块五验收：导入培养方案、成绩单和开课信息后，系统可展示培养方案比对结果、缺失模块、风险提示和课程类型级建议；页面必须始终显示弱结论边界，不得输出毕业资格结论。")
    cursor = insert_paragraph_after(cursor, "共性支撑能力验收：在 Kingbase 环境下成功导入不少于 100 条标准格式 Excel 主数据且无乱码、无部分提交、无主键污染；故意包含关键错误的批次必须整批回滚并输出错误报告；对知识维护、通知发送、审批、导出和敏感访问等关键动作抽查时，审计日志完整可追溯。")
    add_bullets(
        cursor,
        [
            "所有关键流程必须可在审计日志中回查操作者、时间、对象和处理结果。",
            "学生账号不得导出学院级原始数据，高敏字段默认脱敏。",
            "若真实业务数据未完全到位，可使用甲方确认的样例数据完成专项能力验收，但不得将原始需求中的功能点从规格或交付中删除。",
        ],
    )


def build_5(heading: Paragraph) -> None:
    cursor = insert_paragraph_after(heading, "本项目当前采用线框图级文本原型说明核心界面，并按《需求文档.md》的五大模块分别给出代表页面。")
    cursor = insert_heading_after(cursor, "5.1 智能问答与政策知识库首页", 3)
    cursor = add_bullets(
        cursor,
        [
            "页面顶部提供搜索框、常见问题分类入口和“转人工咨询”按钮。",
            "主内容区按卡片展示政策问答结果，每条结果包含标准答案摘要、适用条件、官方来源链接、更新时间和可信标签。",
            "当命中模板类事项时，结果卡片下方显示“下载模板”按钮，可下载证明、请假条、预算表、简报等标准文件。",
            "当问题涉及敏感信息或规则不明确时，页面不直接给出强结论，而是显示风险提示与咨询老师入口。",
            "管理端对应页面需支持知识条目上传、停用、版本切换和来源文件维护。",
        ],
    )
    cursor = insert_heading_after(cursor, "5.2 党团事务流程进度与理论自测页面", 3)
    cursor = add_bullets(
        cursor,
        [
            "页面顶部显示“我的党团进度”，并提供“查看办理说明”“联系组织老师”和“理论自测”入口。",
            "页面上部为身份摘要卡，展示姓名、学号、当前组织类型、当前阶段和状态标签。",
            "中部以纵向时间轴展示已完成节点、当前节点、下一步要求、所需材料和预计时间，关键节点显式高亮。",
            "提醒卡展示即将到期或已超期事项，并说明需提交的思想汇报或其他材料。",
            "点击理论自测后，进入题库答题页，显示题目、选项、剩余题数、得分结果和错题回顾入口。",
        ],
    )
    cursor = insert_heading_after(cursor, "5.3 信息集成与精准推送工作台", 3)
    cursor = add_bullets(
        cursor,
        [
            "管理端页面顶部提供“新增通知”“汇聚来源”“批量发送”按钮，并展示当前待发送批次和渠道状态概览。",
            "左侧筛选区支持来源部门、标签、发布时间、目标年级、专业、毕业生身份等条件组合筛选。",
            "中部列表按卡片或表格展示通知标题、来源、标签、目标人群规模、发送渠道和当前发送状态。",
            "通知编辑区支持录入正文、附件、来源链接和标签，并可预览站内消息、邮件或短信内容。",
            "发送结果区需回显成功数、失败数、未读数和异常原因，学生端通知中心可查看来源依据与接收时间。",
        ],
    )
    cursor = insert_heading_after(cursor, "5.4 电子证明生成与审批页面", 3)
    cursor = add_bullets(
        cursor,
        [
            "学生端页面提供证明申请、请假申请、盖章申请和预算表申请等入口，统一展示模板、用途说明和附件要求。",
            "选择证明类型后，系统自动填充姓名、学号、政治面貌等基础字段，并生成 PDF 预览区。",
            "表单区支持补充用途说明、上传附件、选择提交对象和查看正式流程边界提示。",
            "审批端工作台需在单页展示申请基本信息、附件、历史动作、当前状态、审批意见输入框和通过/驳回按钮。",
            "驳回后学生可回到原表单修改并重提；涉密场景应支持转线下办理并保留线上留痕。",
        ],
    )
    cursor = insert_heading_after(cursor, "5.5 学业情况分析与预警页面", 3)
    cursor = add_bullets(
        cursor,
        [
            "页面顶部展示专业、年级、培养方案版本和“仅供辅助参考”的醒目标识。",
            "左侧或上方提供成绩单上传入口、规则版本选择器和最近一次分析时间。",
            "主内容区按模块展示已完成学分、未修满模块、风险提示和缺口说明，并高亮关键异常。",
            "建议区基于缺失模块给出课程类型级推荐，而不是直接输出选课指令或毕业资格判断。",
            "底部提供“联系老师核验”“查看培养方案明细”“下载分析快照”等操作入口，方便人工复核。",
        ],
    )
    cursor = insert_paragraph_after(cursor, "上述原型描述用于指导后续 UI 详细设计和高保真原型制作。")

    cursor = insert_paragraph_after(cursor)
    cursor.add_run().add_break(WD_BREAK.PAGE)
    appendix = insert_heading_after(cursor, "附录A 业务决策、权限矩阵与待确认事项", 1)
    cursor = insert_paragraph_after(
        appendix,
        "本附录用于收口需求补充、业务讨论和既有草稿中最有价值但不宜直接混入正文的默认方案与待确认问题，供详细设计、联调和验收阶段持续对齐。",
    )

    a1 = insert_heading_after(cursor, "A.1 一期正式范围与业务边界", 3)
    table_a1 = insert_table_after(
        a1,
        ["分类", "内容", "说明"],
        [
            [
                "一期正式范围",
                "政策知识库与 FAQ、受控 AI 匹配、模板下载、党团流程可视化、节点提醒、理论自测、常见事务在线提交、证明 PDF 预览、审批工作台、驳回撤回重提、官方通知汇聚、标签化通知、多渠道通知、文件导入导出、内容治理、字段权限、操作日志、统计看板、学业弱提示与规则维护",
                "均属于一期正式需求范围，不因数据准备困难而从规格中删除。",
            ],
            [
                "默认业务边界",
                "涉及校级正式生效的流程默认仅承担说明、预检、归档或跟踪；学业模块仅输出弱结论；短信通道按配置启用；公众号不默认采用非授权爬取。",
                "用于处理当前仍未完全拍板的跨系统边界和成本边界问题。",
            ],
            [
                "明确不纳入",
                "替代校级正式生效流程、非授权公众号抓取、全校开放、学生批量导出学院级原始数据、大模型自由生成式答复、自动毕业资格判断、动态选课竞争分析",
                "不符合当前业务边界、合规风险或责任边界要求。",
            ],
        ],
    )

    a2 = insert_heading_after(insert_paragraph_after_table(table_a1), "A.2 角色与字段权限矩阵（默认方案）", 3)
    table_a2 = insert_table_after(
        a2,
        ["数据项", "学生", "班主任", "团委老师", "默认规则"],
        [
            ["学号/姓名/年级/专业/班级", "仅看本人", "看本班或本年级", "看业务范围", "主数据默认由学院统一维护。"],
            ["联系方式", "看本人，修改建议走变更申请", "看职责范围", "看业务范围", "不建议老师直接改学生联系方式。"],
            ["成绩/学分情况", "仅看本人", "看本班汇总或个体", "看业务相关范围", "权威成绩默认通过导入更新，不允许直接改源数据。"],
            ["政治面貌/党团阶段", "仅看本人", "默认只读", "可读可维护", "默认由团委老师或受权党团老师维护。"],
            ["请假/盖章/证明申请记录", "看本人并在允许阶段修改本人申请", "看职责范围并审批写", "看业务范围并审批写", "老师写的是审批意见和状态，不改源主数据。"],
            ["处分记录等高敏字段", "默认不展示", "按需、按范围只读", "按需、按范围只读", "默认最小暴露并保留访问日志。"],
            ["导出权限", "无", "仅导出职责范围汇总", "仅导出业务范围汇总", "高敏明细导出应收敛到更高授权。"],
        ],
    )

    a3 = insert_heading_after(insert_paragraph_after_table(table_a2), "A.3 默认审批矩阵", 3)
    table_a3 = insert_table_after(
        a3,
        ["场景", "默认审批链", "说明"],
        [
            ["入党申请/阶段变更", "学生或组织角色 -> 党支部书记初审 -> 团委老师复核 -> 学院领导可选终审", "学院领导节点是否保留需业务方最终确认。"],
            ["入团申请", "学生 -> 团支书初审 -> 团委老师终审", "一期默认由团委老师作为终审角色。"],
            ["普通在读或身份类证明", "学生 -> 班主任核验 -> 团委老师审核 -> 学院授权出件人", "平台可记录结果，但不承诺全电子盖章。"],
            ["党团相关证明", "学生 -> 党支部书记或团支书核验 -> 团委老师终审 -> 学院授权出件人", "根据证明类型选择具体组织角色。"],
            ["请假/盖章/预算表等事务", "学生 -> 班主任或审批老师 -> 业务负责人或学院领导（按需）", "若正式效力属于校级系统，则平台仅承担预检、归档或跟踪。"],
        ],
    )

    a4 = insert_heading_after(insert_paragraph_after_table(table_a3), "A.4 待确认问题与默认结论", 3)
    table_a4 = insert_table_after(
        a4,
        ["编号", "待确认事项", "默认结论"],
        [
            ["Q-01", "哪些流程必须以校级系统为唯一正式生效入口？", "在正式确认前，学院平台默认只做说明、预检、归档或状态跟踪。"],
            ["Q-02", "一期形态优先级是网页优先还是小程序优先，是否双端同时交付？", "默认学生端移动优先、管理端 PC 优先，技术上采用前后端分离，入口形态可后续落地时确定。"],
            ["Q-03", "Kingbase 服务器、版本和测试部署环境由谁提供？", "规格中按 Kingbase 为唯一数据库基线编写，环境资源暂按待提供处理。"],
            ["Q-04", "是否允许接入统一认证、教务或其他校级接口？", "在未确认前不假设可用接口，默认采用文件导入导出。"],
            ["Q-05", "敏感字段是否需要完整展示，是否存在额外敏感词或脱敏规则？", "默认最小权限和脱敏展示，完整可见需额外授权。"],
            ["Q-06", "模板文件、知识库内容和通知来源由谁维护、按何频率更新？", "默认“谁上传谁维护”，需在详细设计前明确责任人和更新频率。"],
            ["Q-07", "学业模块是否允许输出强结论或选课建议的责任边界是什么？", "默认仅提供缺口提示、课程类型级建议和人工核验提示，不输出毕业结论。"],
            ["Q-08", "官方通知来源清单、邮件/短信服务和预算边界是否明确？", "默认站内通知与邮件优先，短信按预算和配置启用。"],
        ],
    )
    insert_paragraph_after_table(table_a4, "若业务方在详细设计前未给出新的明确结论，则本附录中的默认方案作为设计、联调和验收对齐的暂行基线。")


def main() -> None:
    apply_cover_content()

    if doc.tables:
        table = doc.tables[0]
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = TODAY
        table.cell(1, 2).text = "项目组 / Codex"
        table.cell(1, 3).text = "融合需求文档、需求补充、SRS 分析资料和两版现有规格说明书后的增强版需求规格说明书。"
        table.cell(1, 4).text = VERSION

    replace_section("1.1 编写目的", "1.2 读者对象", build_11)
    replace_section("1.2 读者对象", "1.3 软件项目概述", build_12)
    replace_section("1.3 软件项目概述", "1.4 文档概述", build_13)
    replace_section("1.4 文档概述", "1.5 定义", build_14)
    replace_section("1.5 定义", "1.6 参考资料", build_15)
    replace_section("1.6 参考资料", "2. 软件的一般性描述", build_16)
    replace_section("2.1软件产品与其环境之间的关系", "2.2限制与约束", build_21)
    replace_section("2.2限制与约束", "2.3假设与前提条件", build_22)
    replace_section("2.3假设与前提条件", "3. 软件功能需求描述", build_23)
    replace_section("3.1 软件功能概述", "3.2 软件需求的用例模型", build_31)
    replace_section("3.2 软件需求的用例模型", "3.3 软件需求的分析模型", build_32)
    replace_section("3.3 软件需求的分析模型", "4. 其它软件需求描述", build_33)
    replace_section("4.1 性能要求", "4.2 设计约束", build_41)
    replace_section("4.2 设计约束", "4.3 界面要求", build_42)
    replace_section("4.3 界面要求", "4.4 进度要求", build_43)
    replace_section("4.4 进度要求", "4.5 交付要求", build_44)
    replace_section("4.5 交付要求", "4.6 验收要求", build_45)
    replace_section("4.6 验收要求", "5. 软件原型", build_46)
    replace_section("5. 软件原型", None, build_5)
    cleanup_template_fragments()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"SAVED {OUT}")


if __name__ == "__main__":
    main()
