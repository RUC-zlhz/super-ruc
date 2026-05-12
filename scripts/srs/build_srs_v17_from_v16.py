from __future__ import annotations

import argparse
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt

from v1_7.common import V16_DOCX, V16_PDF, V17_DOCX, V17_PDF


S12_ADDENDUM_TITLE = "6. S12 需求缺口闭环增量说明"
S12_CHANGE_DESC = (
    "在 v1.6 基线上补充默认数据导入、成绩单 PDF 人工核验、课程推荐边界、"
    "常用模板下载、统一进度中心、受控通知抓取、短信投递治理和官方链接优先等增量要求。"
)
S12_TABLE_HEADERS = ["增量项", "关联需求", "交付接口/入口", "验收边界"]
S12_ADDENDUM_ROWS = [
    (
        "默认数据导入",
        "FR-009\nFR-015",
        "POST /admin/default-imports/students\n"
        "POST /admin/default-imports/curriculum\n"
        "POST /admin/default-imports/all",
        "默认学生不推断专业/班级；培养方案只更新 2024-default 并写入 CurriculumModule.courses。",
    ),
    (
        "成绩单 PDF 核验",
        "FR-009\nFR-014",
        "POST /report/transcript-pdf\n"
        "POST /admin/report/transcript-pdf-reviews/{batch_id}/commit",
        "学生上传只生成候选批次；教师人工核验提交后才写入正式成绩记录。",
    ),
    (
        "学业缺口与课程推荐",
        "FR-014\nFR-015",
        "GET /report/academic-gap\n"
        "GET /admin/report/academic-gap/{student_id}",
        "推荐基于缺口模块、已修课程、培养方案白名单和可用开课数据；容量、课表、先修或偏好缺失时显示“数据未配置”。",
    ),
    (
        "模板下载与官方链接",
        "FR-002\nFR-003",
        "GET /knowledge/templates\n"
        "GET /knowledge/templates/{template_id}/download",
        "学生仅可下载 active 且关联 published 知识条目的模板；同等相关度下按结构化官方标识优先来源，source_url 仅作回链。",
    ),
    (
        "统一进度中心",
        "FR-004",
        "GET /progress/my\n小程序进度中心页",
        "聚合本人事务申请与党团流程进度，并保留原详情页跳转。",
    ),
    (
        "受控通知抓取",
        "FR-010\nFR-011",
        "GET/POST/PATCH /admin/notices/sources\n"
        "POST /admin/notices/sources/{id}/run\n"
        "GET /admin/notices/ingest-runs",
        "仅支持公开 URL/RSS 和管理员手工触发；抓取结果默认生成 draft notice，不自动发布或群发。",
    ),
    (
        "短信投递治理",
        "FR-011",
        "POST /admin/notices/deliveries/{id}/retry\n"
        "POST /admin/notices/deliveries/{id}/receipt/mock",
        "一期只接 mock/local provider；记录 attempt、失败重试和 mock 回执状态。",
    ),
    (
        "Web / 小程序接入",
        "FR-003\nFR-004\nFR-007\nFR-009\nFR-010\nFR-011\nFR-014",
        "Web 导入中心、审批详情、知识模板、通知来源\n"
        "小程序学业、知识、进度中心",
        "入口与后端契约保持一致，内置证明版式 PDF、模板下载、成绩单候选明细、官方来源标识和进度聚合均可直接触达。",
    ),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Prepare the v1.7 SRS baseline by copying the frozen v1.6 docx/pdf outputs."
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="overwrite existing v1.7 baseline files",
    )
    return parser.parse_args()


def copy_output(src, dst, *, force: bool) -> None:
    if not src.exists():
        raise FileNotFoundError(src)
    if dst.exists() and not force:
        raise FileExistsError(f"{dst} already exists; rerun with --force to overwrite")
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _update_cover_metadata(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if "文档编号" in text and "SRS" in text:
            text = text.replace("V1.5", "V1.7").replace("V1.6", "V1.7")
            _set_paragraph_text(paragraph, text)
        elif text.startswith("日期："):
            _set_paragraph_text(paragraph, "日期：2026-05-11")


def _fill_change_history(document: Document) -> None:
    if not document.tables:
        return
    table = document.tables[0]
    target_row = None
    for row in table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        if values[-1] == "V1.7":
            target_row = row
            break
        if target_row is None and not any(values):
            target_row = row
    if target_row is None:
        target_row = table.add_row()
    values = ["3", "2026-05-11", "项目组", S12_CHANGE_DESC, "V1.7"]
    for cell, value in zip(target_row.cells, values, strict=True):
        cell.text = value


def _fill_row(row, values: list[str] | tuple[str, ...]) -> None:
    for cell, value in zip(row.cells, values, strict=True):
        cell.text = value


def _format_s12_tables(document: Document) -> None:
    widths = [Inches(1.15), Inches(0.95), Inches(2.25), Inches(2.25)]
    for table in document.tables:
        if not table.rows:
            continue
        header_values = [cell.text.strip() for cell in table.rows[0].cells]
        if header_values != S12_TABLE_HEADERS:
            continue
        table.style = "Table Grid"
        table.autofit = False
        _fill_row(table.rows[0], S12_TABLE_HEADERS)
        while len(table.rows) < len(S12_ADDENDUM_ROWS) + 1:
            table.add_row()
        for row, values in zip(table.rows[1:], S12_ADDENDUM_ROWS, strict=False):
            _fill_row(row, values)
        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=True):
                cell.width = width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)


def _add_s12_addendum(document: Document) -> None:
    if any(paragraph.text.strip() == S12_ADDENDUM_TITLE for paragraph in document.paragraphs):
        return

    document.add_page_break()
    document.add_heading(S12_ADDENDUM_TITLE, level=1)
    document.add_paragraph(
        "本节作为 v1.7 在 v1.6 交付基线上的增量补充，保持原有功能范围、"
        "弱结论边界和权限审计原则不变，仅记录 S12 已闭环需求缺口及验收口径。"
    )
    document.add_paragraph(
        "默认导入数据源限定为 docs/source/students/students.xlsx 与 "
        "docs/source/training program/2024_information.md。默认学生导入只写入学号、"
        "姓名、性别和预计毕业年份，不从学号或毕业年份推断专业、年级、班级。"
    )
    document.add_paragraph(
        "默认培养方案导入只维护 version_label=2024-default 的演示版本，写入课程白名单，"
        "不覆盖教师后续维护的非默认培养方案版本。"
    )
    document.add_paragraph(
        "证明类申请一期使用系统内置证明版式生成 PDF 预览；完整标准模板资产绑定、"
        "字段映射和版本留痕不在 v1.7 交付范围内。公众号通知默认手工录入，"
        "自动抓取仅限公开 URL/RSS；短信一期仅为 mock/local provider。"
    )

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _fill_row(table.rows[0], S12_TABLE_HEADERS)
    for values in S12_ADDENDUM_ROWS:
        row = table.add_row()
        _fill_row(row, values)

    document.add_paragraph(
        "上述增量以 S12 定向集成测试、Web 构建、小程序 mp-weixin 构建和 v1.7 出件门作为验收证据。"
    )


def apply_v17_increment(docx_path) -> None:
    document = Document(docx_path)
    _update_cover_metadata(document)
    _fill_change_history(document)
    _add_s12_addendum(document)
    _format_s12_tables(document)
    document.save(docx_path)


def main() -> None:
    args = parse_args()
    copy_output(V16_DOCX, V17_DOCX, force=args.force)
    copy_output(V16_PDF, V17_PDF, force=args.force)
    apply_v17_increment(V17_DOCX)
    print(f"PREPARED {V17_DOCX}")
    print(f"PREPARED {V17_PDF}")


if __name__ == "__main__":
    main()
