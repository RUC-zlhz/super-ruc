from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "docs" / "templates" / "软件需求规格说明书模板.docx"
DEFAULT_OUTPUT = ROOT / "output" / "doc" / "第12组-super-ruc-互测使用说明.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DEEP_BLUE = RGBColor(0x1F, 0x4D, 0x78)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_run_font(run, *, ascii_font: str = "Calibri", east_asia_font: str = "微软雅黑", size: int | None = None,
                 bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, *, ascii_font: str = "Calibri", east_asia_font: str = "微软雅黑", size: int | None = None,
                   bold: bool | None = None, color: RGBColor | None = None) -> None:
    style.font.name = ascii_font
    style._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_col_widths(table, widths_inch: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    total_dxa = int(sum(widths_inch) * 1440)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inch:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths_inch):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, bold=False, color=BODY)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, size=16, bold=True, color=BLUE)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.keep_with_next = True

    heading2 = doc.styles["Heading 2"]
    set_style_font(heading2, size=13, bold=True, color=BLUE)
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(7)
    heading2.paragraph_format.keep_with_next = True

    heading3 = doc.styles["Heading 3"]
    set_style_font(heading3, size=12, bold=True, color=DEEP_BLUE)
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(5)
    heading3.paragraph_format.keep_with_next = True


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_paragraph(doc: Document, text: str, *, style: str = "Normal", align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run)


def add_direct_paragraph(
    doc: Document,
    text: str,
    *,
    size: int,
    bold: bool = False,
    color: RGBColor = BODY,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_after: int = 6,
) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_col_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, ascii_font="Consolas", east_asia_font="等线", size=10, color=BODY)
        if idx < len(lines) - 1:
            run.add_break()


def fill_cell(cell, text: str, *, bold: bool = False, color: RGBColor = BODY) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold, color=color)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_col_widths(table, [1.875, 4.625])
    for label, value in rows:
        row = table.add_row()
        fill_cell(row.cells[0], label, bold=True, color=DEEP_BLUE)
        fill_cell(row.cells[1], value)
        set_cell_shading(row.cells[0], "E8EEF5")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], *,
              header_fill: str = "E8EEF5") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_col_widths(table, widths)
    set_row_cant_split(table.rows[0])
    for idx, text in enumerate(headers):
        fill_cell(table.rows[0].cells[idx], text, bold=True, color=DEEP_BLUE)
        set_cell_shading(table.rows[0].cells[idx], header_fill)
    for data in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, text in enumerate(data):
            fill_cell(row.cells[idx], text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2)
        marker = p.add_run("- ")
        set_run_font(marker, size=11, bold=True, color=DEEP_BLUE)
        content = p.add_run(item)
        set_run_font(content)


def add_numbered(doc: Document, items: list[str]) -> None:
    rows = [[str(idx), item] for idx, item in enumerate(items, start=1)]
    add_table(doc, ["步骤", "说明"], rows, [0.7, 5.8], header_fill="F4F6F9")


def build_doc(output_path: Path) -> Path:
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    configure_styles(doc)
    set_page_layout(doc)

    add_direct_paragraph(doc, "文档编号：第12组 – super-ruc – UGM – 1.0", size=11, color=MUTED)
    add_direct_paragraph(
        doc,
        "super-ruc 项目互测使用说明",
        size=22,
        bold=True,
        color=DEEP_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_direct_paragraph(
        doc,
        "第12组 | 面向其他小组的快速上手与测试说明",
        size=11,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_direct_paragraph(
        doc,
        "更新日期：2026-05-26",
        size=11,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )

    doc.add_paragraph("")

    add_kv_table(
        doc,
        [
            ("项目名称", "super-ruc（信息学院学生综合服务与党团管理平台）"),
            ("适用轮次", "测试互评实验第一轮 / 第二轮"),
            ("推荐测试入口", "优先测试 Web 管理端；小程序测试作为补充项"),
            ("当前可用环境", "内网部署：http://10.10.0.13/（2026-05-25 已核实可访问）"),
        ],
    )

    doc.add_paragraph("")
    add_paragraph(doc, "文档变更历史记录", style="Heading 2")
    add_table(
        doc,
        ["版本", "日期", "修改人", "说明"],
        [["1.0", "2026-05-25", "第12组", "首版互测使用说明，覆盖访问方式、测试账号、推荐流程、本地启动与已知限制。"]],
        [0.7, 1.0, 0.9, 3.9],
    )

    doc.add_paragraph("")
    add_paragraph(doc, "目录", style="Heading 2")
    add_bullets(
        doc,
        [
            "1. 引言",
            "2. 测试前准备",
            "3. Web 管理端测试说明",
            "4. 小程序与本地环境说明",
            "5. 已知限制与反馈建议",
        ],
    )

    add_paragraph(doc, "1. 引言", style="Heading 1")
    add_paragraph(doc, "1.1 编写目的", style="Heading 2")
    add_paragraph(
        doc,
        "本文档用于帮助其他小组快速访问、理解并测试第12组的 super-ruc 项目。文档重点提供可直接执行的访问地址、账号信息、推荐测试路径、示例输入输出和当前已知限制，减少因环境或说明不清而导致的“无法测试”问题。",
    )

    add_paragraph(doc, "1.2 读者对象", style="Heading 2")
    add_bullets(
        doc,
        [
            "参与测试互评实验、需要测试第12组项目的其他小组。",
            "负责审核 bug 报告或协助复测的助教与教师。",
            "第12组内部用于说明当前测试入口与环境状态的成员。",
        ],
    )

    add_paragraph(doc, "1.3 项目概述", style="Heading 2")
    add_paragraph(
        doc,
        "super-ruc 是一个面向信息学院师生的综合服务平台，包含 Web 管理端和微信小程序学生端，覆盖政策知识库、党团流程、通知分发、事务审批、学生画像、培养方案与学业分析等模块。互测时建议优先从 Web 管理端进入，因为当前运行态更稳定、默认数据更完整、测试路径更清晰。",
    )

    add_paragraph(doc, "1.4 文档使用方式", style="Heading 2")
    add_bullets(
        doc,
        [
            "如果只想快速体验：直接看“2. 测试前准备”和“3. Web 管理端测试说明”。",
            "如果内网环境无法访问：直接看“4. 小程序与本地环境说明”中的本地启动方式。",
            "如果要判断某个现象是不是 bug：先看“5. 已知限制与反馈建议”，避免把已明确说明的数据空态误报为缺陷。",
        ],
    )

    add_paragraph(doc, "1.5 定义", style="Heading 2")
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["Web 管理端", "老师/管理员使用的后台，推荐互测主入口。"],
            ["小程序学生端", "微信小程序学生端，适合测试学生视角的通知、请求、进度等功能。"],
            ["内网部署", "当前稳定测试环境，入口为 http://10.10.0.13/。"],
            ["本地冷启动", "在测试者自己机器上通过脚本拉起后端、Web 与默认数据的方式。"],
        ],
        [1.35, 5.15],
    )

    add_paragraph(doc, "1.6 参考资料", style="Heading 2")
    add_bullets(
        doc,
        [
            "《测试实验指导书》：明确了互测实验目标、评分规则和使用说明文档要求。",
            "《基本功能文档》：定义了 super-ruc 的五个核心功能模块与非功能目标。",
            "仓库内最新部署与计划记录：docs/notes/current-implementation-plan.md、deploy/intranet-prod/README.md。",
        ],
    )

    add_paragraph(doc, "2. 测试前准备", style="Heading 1")
    add_paragraph(doc, "2.1 推荐测试方式", style="Heading 2")
    add_table(
        doc,
        ["方式", "推荐程度", "适用场景", "说明"],
        [
            ["远端 Web 管理端", "高", "大多数互测场景", "浏览器直接访问即可，2026-05-25 已实测可用。"],
            ["本地一键启动", "中", "内网不可达、需要稳定复现或需要 mock 小程序登录", "需要 Docker、uv、pnpm，但可完全控制数据环境。"],
            ["小程序真机/开发者工具", "中", "想测试学生视角流程", "更适合第二步补充测试，不建议作为首次进入路径。"],
        ],
        [1.3, 0.8, 1.6, 2.8],
    )

    add_paragraph(doc, "2.2 访问地址与账号", style="Heading 2")
    add_table(
        doc,
        ["项目", "地址 / 账号", "当前状态与备注"],
        [
            ["Web 管理端", "http://10.10.0.13/", "推荐主入口；2026-05-26 实测 HTTP 200。"],
            ["健康检查", "http://10.10.0.13/healthz", "返回 {\"status\":\"ok\"}。"],
            ["API 前缀", "http://10.10.0.13/api/v1", "用于抓包、接口联调或复核。"],
            ["管理员账号", "admin / admin123", "2026-05-25 实测可登录；登录后会提示修改默认密码。互测时请点击“稍后处理”，不要改掉共享密码。"],
            ["旧临时公网地址", "http://123.57.54.195/", "当前返回 502，已不再作为有效测试入口。"],
        ],
        [1.15, 2.2, 3.15],
    )

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_col_widths(callout, [6.5])
    set_cell_shading(callout.cell(0, 0), "F4F6F9")
    fill_cell(
        callout.cell(0, 0),
        "共享管理员账号是为了方便互测而保留的。请测试小组不要在远端环境中修改 admin 的默认密码；如果误改，请联系第12组重置后再继续测试。",
        bold=True,
        color=DEEP_BLUE,
    )

    add_paragraph(doc, "2.3 当前默认数据状态（2026-05-26 核实）", style="Heading 2")
    add_table(
        doc,
        ["模块 / 数据", "当前状态", "互测建议"],
        [
            ["学生主档", "5 条默认学生", "可直接用于搜索、学生管理和本地 mock 小程序登录。"],
            ["党团流程模板", "2 套（党员发展、团员管理）", "适合测试流程模板查看与“发起学生流程”。"],
            ["学生流程实例", "0 条", "如果要测进度实例，请手工发起一条流程。"],
            ["通知", "1 条已发布通知", "可直接测试通知中心列表与详情查看。"],
            ["培养方案", "7 条方案记录", "可直接测试培养方案列表和查看。"],
            ["知识条目", "已有官方示例知识", "可直接测试知识搜索、标签命中和检索匹配；也可在后台新增带“[互测]”的条目。"],
            ["常用模板", "4 个示例模板", "学生端模板区可查看并下载党员证明、团员证明、教室借用和项目申报模板。"],
        ],
        [1.45, 1.45, 3.6],
    )

    add_paragraph(doc, "2.4 测试注意事项", style="Heading 2")
    add_bullets(
        doc,
        [
            "推荐先测读操作，再测写操作。写操作尽量给标题或备注加前缀“[互测]”，便于后续识别。",
            "知识库已预置官方示例知识和常用模板；党团流程实例默认仍为空，如需测试学生进度链路请主动创建一条流程实例。",
            "如果遇到地址无法访问、共享密码失效或页面整体打不开，可优先记为崩溃类问题，同时联系第12组确认环境状态。",
        ],
    )

    add_paragraph(doc, "3. Web 管理端测试说明", style="Heading 1")
    add_paragraph(doc, "3.1 快速上手步骤", style="Heading 2")
    add_numbered(
        doc,
        [
            "在浏览器中打开 http://10.10.0.13/ 。",
            "使用管理员账号 admin / admin123 登录。",
            "如果弹出“请尽快修改初始密码”，点击“稍后处理”继续进入系统。",
            "进入后优先测试：党团流程管理、用户管理、培养方案管理、通知中心。",
        ],
    )

    add_paragraph(doc, "3.2 推荐互测路径（低门槛）", style="Heading 2")
    add_table(
        doc,
        ["步骤", "操作", "预期结果（示例输出）"],
        [
            ["1", "登录 Web 管理端", "登录成功，进入后台首页或对应默认页。"],
            ["2", "打开“党团流程管理”", "可看到 2 套流程模板：党员发展完整流程、团员发展与团籍管理。"],
            ["3", "在“学生流程”中点击“发起学生流程”", "弹出发起窗口，可选择模板、搜索学生。"],
            ["4", "搜索学生：2024202721 或 曾翎一", "候选学生数量为 1，右侧预览自动显示该学生信息。"],
            ["5", "打开“用户管理 -> 学生管理”", "列表中可看到 5 名默认学生。"],
            ["6", "打开“培养方案管理”", "列表中可看到 7 条方案记录。"],
            ["7", "打开“通知中心”", "至少可看到 1 条已发布通知。"],
        ],
        [0.55, 2.65, 3.3],
    )

    add_paragraph(doc, "3.3 可选的写操作测试", style="Heading 2")
    add_table(
        doc,
        ["功能", "建议做法", "备注"],
        [
            ["党团流程实例", "在“发起学生流程”里选择模板后，搜索默认学生并发起一条流程。", "会写入新实例，适合测试流程创建与后续列表刷新。"],
            ["知识库", "先用默认官方知识测试搜索与检索匹配；如需测写操作，再新增一条标题带“[互测]”的知识条目并发布。", "默认知识用于读链路验证，写链路请加前缀便于识别。"],
            ["通知中心", "可新增一条测试通知并只用站内消息通道。", "建议标题加“[互测]”，便于识别。"],
        ],
        [1.3, 3.3, 1.9],
    )

    add_paragraph(doc, "3.4 常用直达页面", style="Heading 2")
    add_table(
        doc,
        ["页面", "路径", "用途"],
        [
            ["党团流程管理", "/workflow/party-stage", "查看流程模板、发起学生流程、查看提醒记录。"],
            ["通知中心", "/notice/list", "查看、发布和分发通知。"],
            ["知识库管理", "/knowledge/entries", "维护知识条目、来源和模板。"],
            ["培养方案管理", "/academic/curriculum", "查看和维护培养方案、模块与课程。"],
            ["用户管理", "/system/users", "查看后台账号和学生管理入口。"],
            ["审批工作台", "/approval/workbench", "测试事务审批相关列表与详情。"],
        ],
        [1.4, 1.7, 3.4],
    )

    add_paragraph(doc, "4. 小程序与本地环境说明", style="Heading 1")
    add_paragraph(doc, "4.1 小程序测试前提", style="Heading 2")
    add_bullets(
        doc,
        [
            "小程序不是本轮互测的推荐首入口，因为它对网络环境和微信登录条件更敏感。",
            "如果要测学生端，建议优先使用本地 mock 环境；这样不依赖真实微信 code，也更容易重复复现问题。",
            "内网演示包的 API 目标是 http://10.10.0.13/api/v1；旧的 123.57.54.195 已失效。",
        ],
    )

    add_paragraph(doc, "4.2 本地一键启动（推荐给需要深测的小组）", style="Heading 2")
    add_numbered(
        doc,
        [
            "确保本机已安装 Docker Desktop、uv、pnpm，并能正常打开 PowerShell。",
            "进入仓库根目录 D:\\software-engineer\\super-ruc 。",
            "执行脚本 .\\scripts\\dev\\start-dev.ps1 。",
            "脚本完成后，Web 默认地址为 http://localhost:4173 ，后端文档地址为 http://localhost:8080/docs ，管理员账号仍为 admin / admin123 。",
        ],
    )
    add_code_block(
        doc,
        [
            ".\\scripts\\dev\\start-dev.ps1",
            "Web:     http://localhost:4173",
            "Backend: http://localhost:8080/docs",
            "Admin:   admin / admin123",
        ],
    )

    add_paragraph(doc, "4.3 本地默认测试学生", style="Heading 2")
    add_table(
        doc,
        ["学号", "姓名", "用途"],
        [
            ["2024201540", "张念昊", "学生搜索、流程发起、本地 mock 登录"],
            ["2024201534", "胡晓锋", "学生搜索、流程发起、本地 mock 登录"],
            ["2024202721", "曾翎一", "学生搜索、流程发起、本地 mock 登录"],
            ["2024201517", "李明蔚", "学生搜索、画像查看、本地 mock 登录"],
            ["2024000000", "test", "边界测试或空白演示账号"],
        ],
        [1.5, 1.4, 3.6],
    )

    add_paragraph(doc, "4.4 本地小程序 mock 登录方式", style="Heading 2")
    add_numbered(
        doc,
        [
            "先按 4.2 启动本地后端，确认 http://127.0.0.1:8080/healthz 返回 ok。",
            "在微信开发者工具中导入 miniapp/dist/build/mp-weixin 。",
            "进入“我的 / 个人中心”，填写学号和姓名，例如 2024202721 / 曾翎一 。",
            "点击“微信一键登录”即可走本地 mock 登录；本地环境已经修复了重复导入开发者工具后 mock 身份变化的问题。",
        ],
    )

    add_paragraph(doc, "5. 已知限制与反馈建议", style="Heading 1")
    add_paragraph(doc, "5.1 当前已知限制", style="Heading 2")
    add_bullets(
        doc,
        [
            "知识库模块已预置官方示例知识和常用模板，但真实学院资料仍可能不完整；遇到特殊情形时应以学院正式通知和人工确认为准。",
            "党团流程模板已存在，但默认没有学生流程实例；需要老师侧手工发起后，学生进度链路才会出现实例数据。",
            "小程序真实环境依赖微信登录与网络条件，本轮互测建议将其作为补充测试项，而不是唯一入口。",
            "共享管理员账号会提示修改初始密码。互测时请不要真的修改，否则会影响其他小组继续测试。",
        ],
    )

    add_paragraph(doc, "5.2 建议优先关注的 bug 类型", style="Heading 2")
    add_bullets(
        doc,
        [
            "崩溃类：地址无法访问、登录直接失败、页面白屏、按钮导致页面卡死、后端返回 5xx。",
            "逻辑类：学生搜索结果不对、流程发起后列表不刷新、通知状态与操作不一致、培养方案统计不正确、知识条目发布后仍不可检索。",
            "文档类：如果你严格按本文档操作仍无法访问或无法完成关键路径，也可以视为有效问题记录。",
        ],
    )

    add_paragraph(doc, "5.3 问题反馈建议", style="Heading 2")
    add_bullets(
        doc,
        [
            "记录问题时，建议至少写明：访问地址、账号、操作路径、输入内容、实际结果、期望结果。",
            "如果是页面问题，请附上截图；如果是接口问题，请附上状态码或返回内容；如果需要在课程互测平台登记结果，可按《测试实验指导书》中的统一平台和账号规则提交。",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build peer testing usage guide DOCX.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    output = build_doc(args.output)
    print(output)


if __name__ == "__main__":
    main()
