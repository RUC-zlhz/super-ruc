from __future__ import annotations

import html
import math
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from lxml import etree
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
SOURCE_DOCX = ROOT / "软件设计规格说明书0513.docx"
OUTPUT_DOCX = ROOT / "output" / "doc" / "软件设计规格说明书-信息学院学生综合服务与党团管理平台-v1.1.docx"
MMD_DIR = ROOT / "docs" / "source" / "diagrams" / "mermaid" / "software-design-spec"
RENDER_DIR = ROOT / "docs" / "source" / "diagrams" / "rendered" / "software-design-spec-v1_1"

FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\Deng.ttf"),
]


INIT_BLOCK = """%%{
  init: {
    "theme": "base",
    "themeVariables": {
      "fontFamily": "Times New Roman, SimSun, serif",
      "primaryColor": "#ffffff",
      "primaryBorderColor": "#222222",
      "primaryTextColor": "#111111",
      "lineColor": "#666666",
      "clusterBkg": "#f8f3f4",
      "clusterBorder": "#b70f24",
      "edgeLabelBackground": "#ffffff"
    },
    "flowchart": {
      "curve": "linear",
      "nodeSpacing": 56,
      "rankSpacing": 64,
      "htmlLabels": false
    },
    "sequence": {
      "mirrorActors": false,
      "wrap": true,
      "useMaxWidth": true
    }
  }
}%%"""


MMD_SOURCES: dict[str, str] = {
    "architecture": INIT_BLOCK
    + """
flowchart TB
  subgraph Client[用户接入层]
    direction TB
    Mini[学生端微信小程序\nminiapp/src/pages]
    Web[PC 管理端 Web\nweb/src/views]
  end
  subgraph Gateway[统一接口与鉴权层]
    direction TB
    Request[utils/request.ts\n统一 API 基址、Token、错误处理]
    Auth[backend/app/auth\n微信登录、账号角色、JWT]
  end
  subgraph Domain[后端领域服务层 FastAPI]
    direction TB
    Knowledge[knowledge\n政策知识、模板、官方来源优先]
    Workflow[workflow/progress\n党团流程、事务申请、统一进度]
    Notice[notice\n站内通知、邮件/短信治理]
    Report[report/exchange\n学业缺口、成绩单候选解析、导入导出]
    ProfileHonor[profile/honor\n画像、敏感查看、荣誉公示]
    Audit[audit\n操作审计、归档与权限策略]
  end
  subgraph Data[数据与外部支撑层]
    direction TB
    DB[(Kingbase ES\n关系数据与迁移)]
    Redis[(Redis\n锁、缓存、调度辅助)]
    ObjectStore[(MinIO / 本地对象存储\n附件、模板、成绩单)]
    Message[SMTP / Mock SMS\n通知触达凭据]
    Wechat[微信 code2session\n学生身份绑定]
  end
  Mini --> Request
  Web --> Request
  Request --> Auth
  Auth --> Domain
  Domain --> DB
  Domain --> Redis
  Knowledge --> ObjectStore
  Workflow --> ObjectStore
  Report --> ObjectStore
  Notice --> Message
  Auth --> Wechat
  Audit --> DB
""".strip()
    + "\n",
    "ui-flow": INIT_BLOCK
    + """
flowchart TB
  Start([进入平台]) --> Role{访问端与角色}
  Role --> MiniHome[学生端：首页]
  Role --> WebLogin[管理端：登录]
  subgraph Mini[学生端主流程]
    direction TB
    MiniHome --> Knowledge[知识查询\n搜索、AI 匹配、模板下载]
    MiniHome --> RequestList[服务/事务申请\n列表、发起、详情]
    MiniHome --> NoticeBox[消息/通知中心\n列表、详情、已读]
    MiniHome --> Profile[我的画像\n登录绑定、纠错、敏感查看]
    MiniHome --> Progress[统一进度\n党团流程 + 事务申请聚合]
    Progress --> Workflow[党团进度详情]
    Workflow --> Quiz[理论自测]
    MiniHome --> Academic[学业查看\n缺口、建议课程、成绩单上传]
    MiniHome --> Honor[荣誉榜\n筛选、详情]
  end
  subgraph Web[PC 管理端主流程]
    direction TB
    WebLogin --> Dashboard[运营看板]
    Dashboard --> Approval[审批工作台]
    Dashboard --> KnowledgeAdmin[知识条目管理]
    Dashboard --> NoticeAdmin[通知中心]
    Dashboard --> Import[导入中心]
    Dashboard --> Curriculum[培养方案]
    Dashboard --> StudentProfile[学生画像]
    Dashboard --> AuditLog[审计日志]
    Dashboard --> UserManage[用户管理]
  end
""".strip()
    + "\n",
    "request-sequence": INIT_BLOCK
    + """
sequenceDiagram
  autonumber
  participant Student as 学生小程序
  participant API as workflow/request API
  participant Storage as 对象存储
  participant Progress as 统一进度服务
  participant Teacher as 管理端审批人
  participant Audit as 审计服务
  Student->>API: 获取申请类型与 form_schema
  API-->>Student: 返回动态表单、附件规则、证明预览能力
  Student->>API: 创建/保存草稿
  Student->>API: 上传附件
  API->>Storage: 保存附件对象和元数据
  Student->>API: 提交申请
  API->>Audit: 记录提交动作
  Progress->>API: 聚合申请状态到 /progress/my
  Teacher->>API: 查看申请详情与审批流水
  Teacher->>API: 审批通过/驳回/转线下
  API->>Audit: 记录审批动作
  API-->>Student: 返回最新状态、意见与证明预览入口
""".strip()
    + "\n",
    "package-class": INIT_BLOCK
    + """
classDiagram
  class AuthModule {
    User
    Role
    UserRole
    Student
    +issueToken()
    +checkScope()
  }
  class KnowledgeModule {
    KnowledgeEntry
    KnowledgeSource
    TemplateAsset
    +search()
    +aiMatch()
  }
  class WorkflowModule {
    Request
    RequestType
    RequestAttachment
    RequestApprovalRecord
    +submitRequest()
    +previewProof()
  }
  class ProgressModule {
    ProgressItem
    +buildMyProgress()
  }
  class NoticeModule {
    Notice
    NoticeDelivery
    SmsAttempt
    +markRead()
    +dispatch()
  }
  class ReportExchangeModule {
    CurriculumPlan
    StudentCourseRecord
    TranscriptUpload
    ImportBatch
    +analyzeGap()
    +parseTranscriptCandidate()
  }
  class ProfileHonorModule {
    ProfileFact
    ProfileCorrection
    HonorRecord
    HonorRecipient
    +requestFullView()
  }
  class AuditModule {
    AuditLog
    AuditLogHistory
    RoleFieldPolicy
    +recordEvent()
    +archiveExpiredLogs()
  }
  AuthModule --> WorkflowModule : actor/student binding
  AuthModule --> NoticeModule : audience scope
  WorkflowModule --> ProgressModule : request state aggregation
  WorkflowModule --> AuditModule : approval events
  KnowledgeModule --> AuditModule : source and template maintenance
  ReportExchangeModule --> ProfileHonorModule : academic/profile facts
  ProfileHonorModule --> AuditModule : sensitive read and export
""".strip()
    + "\n",
    "data-er": INIT_BLOCK
    + """
erDiagram
  USERS ||--o{ USER_ROLES : has
  ROLES ||--o{ USER_ROLES : grants
  USERS ||--o| STUDENTS : binds
  STUDENTS ||--o{ REQUESTS : submits
  REQUEST_TYPES ||--o{ REQUESTS : defines
  REQUESTS ||--o{ REQUEST_ATTACHMENTS : owns
  REQUESTS ||--o{ REQUEST_APPROVAL_RECORDS : records
  STUDENTS ||--o{ NOTICE_DELIVERIES : receives
  NOTICES ||--o{ NOTICE_DELIVERY_BATCHES : dispatches
  NOTICE_DELIVERY_BATCHES ||--o{ NOTICE_DELIVERIES : contains
  STUDENTS ||--o{ STUDENT_WORKFLOW_NODES : tracks
  STUDENTS ||--o{ STUDENT_COURSE_RECORDS : has
  CURRICULUM_PLANS ||--o{ CURRICULUM_MODULES : defines
  HONOR_RECORDS ||--o{ HONOR_RECIPIENTS : includes
  STUDENTS ||--o{ HONOR_RECIPIENTS : wins
  STUDENTS ||--o{ PROFILE_FACTS : aggregates
  KNOWLEDGE_ENTRIES ||--o{ KNOWLEDGE_SOURCES : cites
  KNOWLEDGE_ENTRIES ||--o{ TEMPLATE_ASSETS : provides
  AUDIT_LOGS ||--o{ AUDIT_LOG_HISTORY : archives
  ROLES ||--o{ ROLE_FIELD_POLICIES : controls
""".strip()
    + "\n",
    "deployment": INIT_BLOCK
    + """
flowchart TB
  subgraph Client[访问终端]
    direction TB
    Browser[PC 浏览器\n管理端 Web]
    WechatClient[微信客户端\n学生端小程序]
  end
  subgraph Edge[接入与静态资源]
    direction TB
    Nginx[Nginx\nWeb 静态资源与 /api 反向代理]
    WechatGateway[HTTPS API 网关\n小程序合法域名]
  end
  subgraph App[后端应用]
    direction TB
    FastAPI[FastAPI / Uvicorn\nbackend/app/main.py]
    Scheduler[审计归档与异步治理任务]
  end
  subgraph Infra[基础设施]
    direction TB
    Kingbase[(Kingbase ES\n业务表、迁移、索引)]
    Redis[(Redis\n缓存、锁、任务辅助)]
    MinIO[(MinIO / 本地对象存储\n附件、模板、PDF)]
    MailSms[SMTP / MailHog / Mock SMS\n触达凭据与 attempt]
    WechatAPI[微信 code2session]
  end
  Browser --> Nginx --> FastAPI
  WechatClient --> WechatGateway --> FastAPI
  FastAPI --> Kingbase
  FastAPI --> Redis
  FastAPI --> MinIO
  FastAPI --> MailSms
  FastAPI --> WechatAPI
  Scheduler --> Kingbase
""".strip()
    + "\n",
}


@dataclass(frozen=True)
class DiagramAsset:
    key: str
    title: str
    caption: str
    kind: str


DIAGRAMS = [
    DiagramAsset("architecture", "系统总体架构", "图 3-1 系统总体架构图", "architecture"),
    DiagramAsset("ui-flow", "主要界面跳转", "图 3-2 主要界面跳转流程图", "ui_flow"),
    DiagramAsset("request-sequence", "事务申请顺序", "图 3-3 事务申请状态流转与审批顺序图", "request"),
    DiagramAsset("package-class", "核心类关系", "图 3-4 核心类关系图", "class"),
    DiagramAsset("data-er", "数据实体关系", "图 3-5 数据实体关系图", "data"),
    DiagramAsset("deployment", "系统部署结构", "图 3-6 系统部署结构图", "deployment"),
]


def font_path() -> Path | None:
    return next((path for path in FONT_CANDIDATES if path.exists()), None)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidate = Path(r"C:\Windows\Fonts\simhei.ttf") if bold else font_path()
    if candidate and candidate.exists():
        return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def wrap_label(text: str, width: int = 15) -> list[str]:
    raw_lines = text.split("\n")
    lines: list[str] = []
    for raw_line in raw_lines:
        line = raw_line.strip()
        while len(line) > width:
            lines.append(line[:width])
            line = line[width:]
        if line:
            lines.append(line)
    return lines or [""]


def svg_text(x: int, y: int, lines: Iterable[str], *, size: int = 24, weight: str = "400", anchor: str = "middle") -> str:
    tspans = []
    for index, line in enumerate(lines):
        dy = 0 if index == 0 else size + 6
        tspans.append(
            f'<tspan x="{x}" dy="{dy}">{html.escape(line)}</tspan>'
        )
    return (
        f'<text x="{x}" y="{y}" text-anchor="{anchor}" '
        f'font-family="SimSun, Microsoft YaHei, Times New Roman, serif" '
        f'font-size="{size}" font-weight="{weight}" fill="#222">'
        + "".join(tspans)
        + "</text>"
    )


class DiagramPainter:
    def __init__(self, width: int = 1200, height: int = 1600):
        self.width = width
        self.height = height
        self.svg: list[str] = []
        self.image = Image.new("RGB", (width, height), "white")
        self.draw = ImageDraw.Draw(self.image)
        self.font_title = load_font(34, bold=True)
        self.font_box = load_font(23)
        self.font_small = load_font(20)
        self.svg.append(
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">'
        )
        self.svg.append("<defs><marker id='arrow' viewBox='0 0 10 10' refX='9' refY='5' markerWidth='8' markerHeight='8' orient='auto-start-reverse'><path d='M 0 0 L 10 5 L 0 10 z' fill='#666'/></marker></defs>")
        self.svg.append("<rect width='100%' height='100%' fill='#ffffff'/>")

    def title(self, text: str) -> None:
        self.svg.append(svg_text(self.width // 2, 54, [text], size=34, weight="700"))
        self.draw.text((self.width // 2, 28), text, anchor="ma", font=self.font_title, fill="#222")

    def box(self, x: int, y: int, w: int, h: int, label: str, *, fill: str = "#ffffff", stroke: str = "#b70f24", radius: int = 18) -> None:
        self.svg.append(
            f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{radius}" fill="{fill}" stroke="{stroke}" stroke-width="3"/>'
        )
        self.draw.rounded_rectangle((x, y, x + w, y + h), radius=radius, fill=fill, outline=stroke, width=3)
        lines = wrap_label(label, max(8, int(w / 34)))
        start_y = y + h // 2 - ((len(lines) - 1) * 15)
        self.svg.append(svg_text(x + w // 2, start_y, lines, size=23))
        for idx, line in enumerate(lines):
            self.draw.text((x + w // 2, start_y - 21 + idx * 30), line, anchor="ma", font=self.font_box, fill="#222")

    def lane(self, x: int, y: int, w: int, h: int, title: str, fill: str = "#fbf7f7") -> None:
        self.svg.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="24" fill="{fill}" stroke="#d8a3ab" stroke-width="2" stroke-dasharray="8 6"/>')
        self.draw.rounded_rectangle((x, y, x + w, y + h), radius=24, fill=fill, outline="#d8a3ab", width=2)
        self.svg.append(svg_text(x + w // 2, y + 36, [title], size=24, weight="700"))
        self.draw.text((x + w // 2, y + 14), title, anchor="ma", font=self.font_box, fill="#222")

    def arrow(self, x1: int, y1: int, x2: int, y2: int, label: str | None = None) -> None:
        self.svg.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="#666" stroke-width="3" marker-end="url(#arrow)"/>')
        self.draw.line((x1, y1, x2, y2), fill="#666", width=3)
        angle = math.atan2(y2 - y1, x2 - x1)
        arrow_len = 16
        for delta in (math.pi * 0.82, -math.pi * 0.82):
            ax = x2 + arrow_len * math.cos(angle + delta)
            ay = y2 + arrow_len * math.sin(angle + delta)
            self.draw.line((x2, y2, ax, ay), fill="#666", width=3)
        if label:
            mx, my = (x1 + x2) // 2, (y1 + y2) // 2 - 10
            self.svg.append(svg_text(mx, my, wrap_label(label, 10), size=18))
            self.draw.text((mx, my - 16), label, anchor="ma", font=self.font_small, fill="#555")

    def save(self, svg_path: Path, png_path: Path) -> None:
        self.svg.append("</svg>")
        svg_path.write_text("\n".join(self.svg), encoding="utf-8")
        self.image.save(png_path, "PNG")


def render_column_diagram(asset: DiagramAsset, groups: list[tuple[str, list[str]]], svg_path: Path, png_path: Path) -> None:
    painter = DiagramPainter()
    painter.title(asset.title)
    y = 115
    group_h = 285 if len(groups) <= 4 else 235
    prev_center: tuple[int, int] | None = None
    colors = ["#fff7f8", "#f7fbff", "#f9fff7", "#fffaf2", "#f8f7ff"]
    for idx, (group_title, items) in enumerate(groups):
        painter.lane(70, y, 1060, group_h, group_title, fill=colors[idx % len(colors)])
        columns = min(3, max(1, len(items)))
        box_w = 300 if columns == 3 else 430
        gap = (1060 - columns * box_w) // (columns + 1)
        box_h = 78
        row_gap = 28
        for item_idx, item in enumerate(items):
            row = item_idx // columns
            col = item_idx % columns
            x = 70 + gap + col * (box_w + gap)
            by = y + 74 + row * (box_h + row_gap)
            painter.box(x, by, box_w, box_h, item, fill="#ffffff")
        center = (600, y + group_h)
        if prev_center:
            painter.arrow(prev_center[0], prev_center[1] + 8, center[0], y - 10)
        prev_center = center
        y += group_h + 65
    painter.save(svg_path, png_path)


def render_ui_flow(asset: DiagramAsset, svg_path: Path, png_path: Path) -> None:
    painter = DiagramPainter()
    painter.title(asset.title)
    painter.box(445, 105, 310, 70, "进入平台", fill="#fff7f8")
    painter.box(445, 220, 310, 70, "角色与访问端判断", fill="#fff7f8")
    painter.arrow(600, 175, 600, 220)
    painter.lane(70, 340, 500, 1120, "学生端微信小程序", fill="#f9fff7")
    painter.lane(630, 340, 500, 1120, "PC 管理端 Web", fill="#f7fbff")
    mini = ["首页", "知识查询", "事务申请", "通知中心", "我的画像", "统一进度", "学业查看", "荣誉榜"]
    web = ["登录", "运营看板", "审批工作台", "知识条目", "通知中心", "导入中心", "学生画像", "审计/用户"]
    for idx, text in enumerate(mini):
        y = 430 + idx * 118
        painter.box(145, y, 350, 70, text)
        if idx > 0:
            painter.arrow(320, y - 48, 320, y)
    for idx, text in enumerate(web):
        y = 430 + idx * 118
        painter.box(705, y, 350, 70, text, stroke="#2166ac")
        if idx > 0:
            painter.arrow(880, y - 48, 880, y)
    painter.arrow(510, 290, 320, 340, "学生")
    painter.arrow(690, 290, 880, 340, "教师/管理员")
    painter.save(svg_path, png_path)


def render_request(asset: DiagramAsset, svg_path: Path, png_path: Path) -> None:
    painter = DiagramPainter()
    painter.title(asset.title)
    steps = [
        "1. 获取申请类型与动态表单",
        "2. 创建/保存草稿",
        "3. 上传附件到对象存储",
        "4. 提交申请并写审计",
        "5. /progress/my 聚合进度",
        "6. 管理端审批通过/驳回/转线下",
        "7. 学生查看状态、意见与证明预览",
    ]
    x = 210
    y = 130
    for idx, step in enumerate(steps):
        painter.box(x, y, 780, 82, step, fill="#fffaf2" if idx in (2, 5) else "#ffffff")
        if idx < len(steps) - 1:
            painter.arrow(x + 390, y + 82, x + 390, y + 122)
        y += 122
    painter.lane(70, 1085, 1060, 300, "关键边界", fill="#f8f7ff")
    bounds = ["学生端只采集表单与附件", "后端控制状态机和权限", "对象存储保存文件本体", "审计记录提交/审批动作"]
    for idx, item in enumerate(bounds):
        painter.box(135 + idx * 250, 1195, 220, 78, item, stroke="#6a51a3")
    painter.save(svg_path, png_path)


def render_diagrams() -> None:
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    MMD_DIR.mkdir(parents=True, exist_ok=True)
    for key, source in MMD_SOURCES.items():
        (MMD_DIR / f"{key}.mmd").write_text(source, encoding="utf-8")
    for asset in DIAGRAMS:
        svg_path = RENDER_DIR / f"{asset.key}.svg"
        png_path = RENDER_DIR / f"{asset.key}.png"
        if asset.kind == "architecture":
            render_column_diagram(
                asset,
                [
                    ("用户接入层", ["学生端微信小程序", "PC 管理端 Web"]),
                    ("统一接口与鉴权层", ["统一请求封装", "微信登录/账号角色", "JWT 与权限校验"]),
                    ("后端领域服务层", ["知识库", "流程/申请/进度", "通知触达", "学业/导入", "画像/荣誉", "审计"]),
                    ("数据与外部支撑层", ["Kingbase", "Redis", "对象存储", "SMTP/Mock SMS", "微信 code2session"]),
                ],
                svg_path,
                png_path,
            )
        elif asset.kind == "ui_flow":
            render_ui_flow(asset, svg_path, png_path)
        elif asset.kind == "request":
            render_request(asset, svg_path, png_path)
        elif asset.kind == "class":
            render_column_diagram(
                asset,
                [
                    ("身份与权限", ["AuthModule\nUser/Role/Student", "AuditModule\nAuditLog/Policy"]),
                    ("学生服务核心", ["WorkflowModule\nRequest/Approval", "ProgressModule\nProgressItem", "NoticeModule\nNotice/Delivery"]),
                    ("知识与数据", ["KnowledgeModule\nEntry/Source/Template", "ReportExchangeModule\nPlan/Course/Upload", "ProfileHonorModule\nFact/Honor"]),
                    ("依赖方向", ["Auth 控制访问", "Workflow 写审计", "Progress 聚合状态", "Profile 敏感查看留痕"]),
                ],
                svg_path,
                png_path,
            )
        elif asset.kind == "data":
            render_column_diagram(
                asset,
                [
                    ("基础身份", ["users", "roles", "user_roles", "students"]),
                    ("事务与进度", ["request_types", "requests", "request_attachments", "approval_records", "workflow_nodes"]),
                    ("知识/通知/学业", ["knowledge_entries", "knowledge_sources", "template_assets", "notices", "notice_deliveries", "curriculum_plans", "student_course_records"]),
                    ("画像/荣誉/审计", ["profile_facts", "profile_corrections", "honor_records", "honor_recipients", "audit_logs", "audit_log_history"]),
                ],
                svg_path,
                png_path,
            )
        elif asset.kind == "deployment":
            render_column_diagram(
                asset,
                [
                    ("访问终端", ["PC 浏览器\n管理端", "微信客户端\n学生端"]),
                    ("接入层", ["Nginx\n静态资源/反向代理", "HTTPS API 网关\n小程序合法域名"]),
                    ("应用层", ["FastAPI/Uvicorn", "审计归档任务", "路由异常统一封装"]),
                    ("基础设施", ["Kingbase ES", "Redis", "MinIO/本地存储", "SMTP/MailHog/Mock SMS", "微信 API"]),
                ],
                svg_path,
                png_path,
            )


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        set_paragraph_text(new_paragraph, text)
    return new_paragraph


def add_caption(paragraph: Paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, bold=True)


def add_body_note(paragraph: Paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def insert_picture(paragraph: Paragraph, asset_key: str) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(RENDER_DIR / f"{asset_key}.png"), width=Inches(6.25))


def find_paragraph(doc: Document, text: str, style: str | None = None) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() != text:
            continue
        if style is not None and paragraph.style.name != style:
            continue
        return paragraph
    raise ValueError(f"paragraph not found: {text}")


def replace_manual_toc(doc: Document) -> None:
    toc_heading = find_paragraph(doc, "目录")
    intro_heading = find_paragraph(doc, "1、引言")
    between: list[Paragraph] = []
    collect = False
    for paragraph in doc.paragraphs:
        if paragraph._p is toc_heading._p:
            collect = True
            continue
        if paragraph._p is intro_heading._p:
            break
        if collect:
            between.append(paragraph)
    for paragraph in between:
        delete_paragraph(paragraph)
    toc_paragraph = insert_paragraph_after(toc_heading)
    toc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = toc_paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在 Word 中自动更新"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def update_tables(doc: Document) -> None:
    tables = doc.tables
    # 文档变更历史记录
    history = tables[0]
    row = history.add_row().cells
    values = ["2", "2026-05-13", "项目组", "在 0513 版基础上对齐 SRS v1.7 与 S13 实现状态，补齐 6 类 Mermaid 图、自动目录、图题/表题、Word/PDF 导出与渲染检查。", "V1.1"]
    for cell, value in zip(row, values, strict=True):
        cell.text = value
    # 项目概述：产品形态
    for row in tables[1].rows:
        if row.cells[0].text.strip() == "产品形态":
            row.cells[1].text = "学生端微信小程序 + PC 管理端 Web + 后端 REST API + Kingbase/Redis/对象存储支撑服务（当前仓库均已有实现或配置资产）。"
    # 前端模块划分：pages.json 页面数
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("声明 13 个页面", "声明 14 个页面")
                cell.text = cell.text.replace("operation_logs", "audit_logs")
                cell.text = cell.text.replace("PC 管理端、政策文件批量导入、邮件/微信模板消息、短信真实通道、AI 纠错与知识库更新审核可作为后续迭代。", "真实短信运营商、外部教务系统/微人大直连、完整证明模板字段映射、AI 纠错闭环与知识库更新审核可作为后续迭代。")
    # 主要接口契约
    api_table = tables[8]
    row = api_table.add_row().cells
    values = ["统一进度", "GET /progress/my", "聚合本人党团流程节点和事务申请状态，返回统一进度中心列表。"]
    for cell, value in zip(row, values, strict=True):
        cell.text = value
    # 页面结构与导航
    page_table = tables[10]
    row = page_table.add_row().cells
    values = ["统一进度", "pages/progress/index", "聚合党团流程和事务申请，展示状态、当前步骤、截止时间和详情入口。", "可跳转党团流程详情或申请详情。"]
    for cell, value in zip(row, values, strict=True):
        cell.text = value
    # 用例总览
    usecase_table = tables[13]
    row = usecase_table.add_row().cells
    values = ["UC-12", "统一进度中心查看", "普通学生", "P1", "progress/index.vue；api/progress.ts；/progress/my"]
    for cell, value in zip(row, values, strict=True):
        cell.text = value
    # 当前代码实现与需求对应关系
    mapping_table = tables[24]
    row = mapping_table.add_row().cells
    values = ["统一进度中心", "miniapp/src/pages/progress/index.vue；miniapp/src/api/progress.ts；backend/app/progress", "已实现本人党团流程与事务申请状态聚合，作为 S13 后的一期进度中心口径。"]
    for cell, value in zip(row, values, strict=True):
        cell.text = value
    # 管理端和后端代码位置补强
    row = mapping_table.add_row().cells
    values = ["PC 管理端与后端服务", "web/src/router/index.ts；web/src/views/**；backend/app/**/router.py", "管理端路由、角色访问控制与后端 /api/v1 业务路由均已纳入当前实现状态，不再作为纯规划内容描述。"]
    for cell, value in zip(row, values, strict=True):
        cell.text = value


def polish_docx() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(SOURCE_DOCX))
    replacements = {
        "文档编号：SIP – SDS – 1.0": "文档编号：SIP – SDS – 1.1",
        "版本：V1.0": "版本：V1.1",
        "日期：2026 年 5 月 9 日": "日期：2026 年 5 月 13 日",
        "说明：当前代码包主要覆盖“学生端微信小程序”。本文对后端服务与 PC 管理端采用“接口契约 + 需求约束”的方式进行设计描述，供后续实现和联调使用。": "说明：当前版本基于 S13 后的完整实现状态完善，覆盖后端 FastAPI 服务、PC 管理端 Web 与学生端微信小程序，并保留接口契约、数据表、部署结构和验收证据的一致性。",
        "当前补全文档重点结合已提交的 miniapp 代码包。该代码包以 uni-app/Vue 3/TypeScript 实现微信小程序学生端，并通过 REST API 与后端服务交互。对于代码中尚未包含的 PC 管理端与后端实现，本文基于前端接口契约和需求文件进行设计约束描述。": "当前完善版以 2026 年 5 月 13 日仓库实现为依据：backend 已提供 FastAPI 业务路由与 Kingbase 迁移，web 已提供 PC 管理端路由与权限控制，miniapp 已提供微信学生端页面与 API 封装。本文不再将后端/管理端作为纯规划内容，而按已实现模块与仍需外部联调的能力分别描述。",
        "本文档共分三章：第一章说明文档目的、读者、项目背景和术语；第二章说明设计目标、原则与约束；第三章从体系结构、用户界面、用例、类、数据和部署六个方面给出软件设计。本文档中的接口路径和数据对象以 miniapp/src/api 目录和页面调用逻辑为主要依据。": "本文档共分三章：第一章说明文档目的、读者、项目背景和术语；第二章说明设计目标、原则与约束；第三章从体系结构、用户界面、用例、类、数据和部署六个方面给出软件设计。本文档中的接口路径和数据对象以 backend/app、web/src/router、miniapp/src/api 与 miniapp/src/pages 的当前实现为主要依据。",
        "系统采用“微信小程序学生端 + 后端业务服务 + Kingbase 数据库 + 文件存储 + PC 管理端”的分层结构。当前代码包实现的是学生端小程序，后端与管理端通过接口契约抽象。": "系统采用“微信小程序学生端 + PC 管理端 Web + 后端业务服务 + Kingbase 数据库 + Redis + 文件存储”的分层结构。当前仓库已包含后端 FastAPI、PC 管理端 Web 和学生端微信小程序三端实现，本文按实际模块边界描述接口、类、数据和部署设计。",
        "该结构将高频学生操作放在小程序端，将批量导入、权限审批、日志审计等复杂任务放在后端和 PC 管理端。学生端不直接访问数据库，也不直接处理敏感字段授权。": "该结构将高频学生操作放在小程序端，将批量导入、权限审批、日志审计等复杂任务放在后端和 PC 管理端。Web 管理端通过路由权限控制访问范围，学生端不直接访问数据库，也不直接处理敏感字段授权。",
        "图 3-2 学生端主要界面跳转流程图": "图 3-2 主要界面跳转流程图",
        "图 3-3 事务申请状态流转图": "图 3-3 事务申请状态流转与审批顺序图",
        "图 3-4 系统部署结构图": "图 3-6 系统部署结构图",
        "本文档已经按照软件设计规格说明书模板补全了项目的设计内容。由于当前提交代码主要是学生端小程序，本文对后端和管理端的设计采用了接口契约、数据表和部署结构的方式进行约束；后续若补充后端或 PC 管理端代码，应继续更新 3.1、3.4、3.5 和 3.6 中的接口、类和部署细节。": "本文档已经按照软件设计规格说明书模板在 0513 版基础上完成完善，当前版本与 SRS v1.7、S13 实现状态、后端/管理端/小程序代码结构保持一致。后续若新增真实短信运营商、完整证明模板字段映射或外部教务系统直连，应继续更新 3.1、3.4、3.5 和 3.6 中的接口、类、数据表与部署细节。",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                if "SDS – 1.0" in paragraph.text:
                    set_paragraph_text(paragraph, paragraph.text.replace("SDS – 1.0", "SDS – 1.1"))
    replace_manual_toc(doc)
    update_tables(doc)
    drawing_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph._p.xpath(".//w:drawing")]
    for paragraph, asset in zip(drawing_paragraphs, ["architecture", "ui-flow", "request-sequence", "deployment"], strict=True):
        insert_picture(paragraph, asset)
    class_heading = find_paragraph(doc, "3.4 类设计", "Heading 2")
    class_image = insert_paragraph_after(class_heading)
    insert_picture(class_image, "package-class")
    class_caption = insert_paragraph_after(class_image)
    add_caption(class_caption, "图 3-4 核心类关系图")
    class_note = insert_paragraph_after(class_caption)
    add_body_note(class_note, "核心类关系图以后端领域模块为中心，补充展示身份权限、事务申请、统一进度、知识库、通知、学业、画像荣誉和审计之间的依赖边界。")
    data_heading = find_paragraph(doc, "3.5 数据设计", "Heading 2")
    data_image = insert_paragraph_after(data_heading)
    insert_picture(data_image, "data-er")
    data_caption = insert_paragraph_after(data_image)
    add_caption(data_caption, "图 3-5 数据实体关系图")
    data_note = insert_paragraph_after(data_caption)
    add_body_note(data_note, "数据实体关系图按身份、事务、通知、知识、学业、画像荣誉和审计归档分组，强调学生主数据、申请状态流转、通知投递和敏感访问审计之间的关联。")
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("图 3-"):
            add_caption(paragraph, paragraph.text.strip())
    doc.save(str(OUTPUT_DOCX))


def embed_svg_extensions(docx_path: Path) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    shutil.copy2(docx_path, temp_path)
    svg_ns = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    etree.register_namespace("asvg", svg_ns)
    with zipfile.ZipFile(temp_path, "r") as zin, zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        document_xml = etree.fromstring(zin.read("word/document.xml"))
        rels_xml = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
        content_types_xml = etree.fromstring(zin.read("[Content_Types].xml"))
        existing_ids = {rel.get("Id") for rel in rels_xml.findall(f"{{{rel_ns}}}Relationship")}
        def new_rel_id(index: int) -> str:
            base = f"rIdSdsSvg{index}"
            if base not in existing_ids:
                existing_ids.add(base)
                return base
            suffix = 1
            while f"{base}_{suffix}" in existing_ids:
                suffix += 1
            rel_id = f"{base}_{suffix}"
            existing_ids.add(rel_id)
            return rel_id
        defaults = content_types_xml.findall(f"{{{ct_ns}}}Default")
        if not any(default.get("Extension") == "svg" for default in defaults):
            default = etree.Element(f"{{{ct_ns}}}Default")
            default.set("Extension", "svg")
            default.set("ContentType", "image/svg+xml")
            content_types_xml.append(default)
        blips = document_xml.xpath(".//a:blip[@r:embed]", namespaces=ns)
        for index, (blip, asset) in enumerate(zip(blips, DIAGRAMS, strict=True), 1):
            rel_id = new_rel_id(index)
            rel = etree.Element(f"{{{rel_ns}}}Relationship")
            rel.set("Id", rel_id)
            rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
            rel.set("Target", f"media/sds-v11-{asset.key}.svg")
            rels_xml.append(rel)
            ext_lst = blip.find("a:extLst", namespaces=ns)
            if ext_lst is None:
                ext_lst = etree.SubElement(blip, f"{{{ns['a']}}}extLst")
            ext = etree.SubElement(ext_lst, f"{{{ns['a']}}}ext")
            ext.set("uri", "{96DAC541-7B7A-43D3-8B79-37D633B846F1}")
            svg_blip = etree.SubElement(ext, f"{{{svg_ns}}}svgBlip")
            svg_blip.set(f"{{{ns['r']}}}embed", rel_id)
        for item in zin.infolist():
            if item.filename in {"word/document.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"}:
                continue
            zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/document.xml", etree.tostring(document_xml, xml_declaration=True, encoding="UTF-8", standalone="yes"))
        zout.writestr("word/_rels/document.xml.rels", etree.tostring(rels_xml, xml_declaration=True, encoding="UTF-8", standalone="yes"))
        zout.writestr("[Content_Types].xml", etree.tostring(content_types_xml, xml_declaration=True, encoding="UTF-8", standalone="yes"))
        for asset in DIAGRAMS:
            zout.writestr(f"word/media/sds-v11-{asset.key}.svg", (RENDER_DIR / f"{asset.key}.svg").read_bytes())
    temp_path.unlink(missing_ok=True)


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    render_diagrams()
    polish_docx()
    embed_svg_extensions(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
