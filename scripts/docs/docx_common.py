from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DEEP_BLUE = RGBColor(0x1F, 0x4D, 0x78)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"


def today_label() -> str:
    return os.environ.get("CODEX_CURRENT_DATE") or date.today().isoformat()


def set_run_font(
    run,
    *,
    ascii_font: str = "Calibri",
    east_asia_font: str = "微软雅黑",
    size: int | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(
    style,
    *,
    ascii_font: str = "Calibri",
    east_asia_font: str = "微软雅黑",
    size: int | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    if style is None:
        return
    style.font.name = ascii_font
    style._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_col_widths(table, widths_inch: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    total_dxa = int(sum(widths_inch) * 1440)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inch:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths_inch):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def apply_table_style(table) -> None:
    for style_name in ("Table Grid", "Table", "Normal Table"):
        try:
            table.style = style_name
            return
        except KeyError:
            continue


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def configure_base_styles(doc: Document) -> None:
    for style_name, size, bold, color in (
        ("Normal", 11, False, BODY),
        ("Body Text", 11, False, BODY),
        ("First Paragraph", 11, False, BODY),
        ("Compact", 11, False, BODY),
        ("Heading 1", 16, True, BLUE),
        ("Heading 2", 13, True, BLUE),
        ("Heading 3", 12, True, DEEP_BLUE),
    ):
        style = doc.styles[style_name] if style_name in doc.styles else None
        set_style_font(style, size=size, bold=bold, color=color)
        if style is None:
            continue
        pf = style.paragraph_format
        if style_name.startswith("Heading"):
            pf.keep_with_next = True
        if style_name == "Heading 1":
            pf.space_before = Pt(18)
            pf.space_after = Pt(10)
        elif style_name == "Heading 2":
            pf.space_before = Pt(14)
            pf.space_after = Pt(7)
        elif style_name == "Heading 3":
            pf.space_before = Pt(10)
            pf.space_after = Pt(5)
        else:
            pf.space_after = Pt(4)
            pf.line_spacing = 1.25


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def prepare_document(template_path: Path) -> Document:
    doc = Document(str(template_path))
    clear_document(doc)
    configure_base_styles(doc)
    set_page_layout(doc)
    return doc


def add_direct_paragraph(
    doc: Document,
    text: str,
    *,
    size: int,
    bold: bool = False,
    color: RGBColor = BODY,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_after: int = 6,
) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_paragraph(doc: Document, text: str, *, style: str = "Normal", align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run)


def add_spacer(doc: Document) -> None:
    doc.add_paragraph("")


def add_bullets(doc: Document, items: list[str], *, left_indent: float = 0.22) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(left_indent)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2)
        marker = p.add_run("- ")
        set_run_font(marker, size=11, bold=True, color=DEEP_BLUE)
        content = p.add_run(item)
        set_run_font(content)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    apply_table_style(table)
    set_table_col_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, ascii_font="Consolas", east_asia_font="等线", size=10, color=BODY)
        if idx < len(lines) - 1:
            run.add_break()


def fill_cell(cell, text: str, *, bold: bool = False, color: RGBColor = BODY, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold, color=color)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    header_fill: str = HEADER_FILL,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    apply_table_style(table)
    set_table_col_widths(table, widths)
    set_row_cant_split(table.rows[0])
    for idx, text in enumerate(headers):
        fill_cell(table.rows[0].cells[idx], text, bold=True, color=DEEP_BLUE)
        set_cell_shading(table.rows[0].cells[idx], header_fill)
    for data in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, text in enumerate(data):
            fill_cell(row.cells[idx], text)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    apply_table_style(table)
    set_table_col_widths(table, [1.875, 4.625])
    for label, value in rows:
        row = table.add_row()
        fill_cell(row.cells[0], label, bold=True, color=DEEP_BLUE)
        fill_cell(row.cells[1], value)
        set_cell_shading(row.cells[0], HEADER_FILL)


def add_steps_table(doc: Document, steps: list[str]) -> None:
    add_table(
        doc,
        ["步骤", "说明"],
        [[str(idx), step] for idx, step in enumerate(steps, start=1)],
        [0.7, 5.8],
        header_fill=LIGHT_FILL,
    )


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    apply_table_style(table)
    set_table_col_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    fill_cell(cell, text, bold=True, color=DEEP_BLUE)


def finalize(doc: Document, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
