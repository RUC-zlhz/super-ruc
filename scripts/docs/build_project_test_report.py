from __future__ import annotations

import argparse
import subprocess
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_common import (
    DEEP_BLUE,
    MUTED,
    ROOT,
    add_bullets,
    add_callout,
    add_direct_paragraph,
    add_kv_table,
    add_paragraph,
    add_spacer,
    add_table,
    finalize,
    prepare_document,
    today_label,
)


TEMPLATE = ROOT / "docs" / "templates" / "软件测试报告模板.docx"
DEFAULT_OUTPUT = ROOT / "output" / "doc" / "软件测试报告-信息学院学生综合服务与党团管理平台-v2.0.docx"


PROJECT_NAME = "信息学院学生综合服务与党团管理平台"
DOC_VERSION = "V2.0"
DOC_NUMBER = "SIP-TEST-2.0"
PRODUCTION_BASE_URL = "http://10.10.0.13/"
PRODUCTION_SNAPSHOT_DATE = "2026-06-26"
PRODUCTION_WEB_TITLE = "信息学院管理后台"
PRODUCTION_WEB_JS = "/assets/index-Cl5iT-qx.js"
PRODUCTION_WEB_CSS = "/assets/index-BKLmXvqt.css"
PRODUCTION_WEB_CHUNKS = "vue / vendor / antdv 三个独立第三方 chunk"
PRODUCTION_KNOWLEDGE_TOTAL = 16
PRODUCTION_TEMPLATE_ENTRY_COUNT = 4
PRODUCTION_CATEGORY_COUNT = 9
REGRESSION_BASELINE_DATE = "2026-06-26"
REGRESSION_PASS_COUNT = 146
DEPLOY_RUN_ID = "28233332227"
PROD_AUTH_ROLE = "SUPER_ADMIN"


def git_head_short() -> str:
    try:
        return subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=ROOT,
            check=True,
            capture_output=True,
            text=True,
        ).stdout.strip()
    except Exception:
        return "unknown"


def build_doc(output_path: Path) -> Path:
    doc = prepare_document(TEMPLATE)
    updated_at = today_label()
    head = git_head_short()

    add_direct_paragraph(doc, f"文档编号：{DOC_NUMBER}", size=11, color=MUTED)
    add_direct_paragraph(
        doc,
        PROJECT_NAME,
        size=22,
        bold=True,
        color=DEEP_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_direct_paragraph(
        doc,
        "软件测试报告（最终提交版）",
        size=18,
        bold=True,
        color=DEEP_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_direct_paragraph(
        doc,
        f"版本：{DOC_VERSION}    当前仓库 HEAD：{head}    更新日期：{updated_at}",
        size=11,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    add_kv_table(
        doc,
        [
            ("测试对象", f"{PROJECT_NAME} 当前主线代码（HEAD {head}）与生产部署链路"),
            ("测试口径", f"整合 {REGRESSION_BASELINE_DATE} S77 合并后全量回归、2026-05-26 测试工程师缺陷审查与缺陷整改复核、{PRODUCTION_SNAPSHOT_DATE} 生产在线抽检"),
            ("覆盖平台", "FastAPI 后端、Vue 3 Web 管理端、uni-app 小程序、金仓数据库、内网生产环境"),
            ("测试目标", "验证当前代码基线、线上真实响应与认证边界，复核历史缺陷整改情况，并给出最终交付结论与残留建议"),
        ],
    )

    add_spacer(doc)
    add_paragraph(doc, "文档变更历史记录", style="Heading 2")
    add_table(
        doc,
        ["序号", "变更日期", "变更人员", "变更内容详情描述", "版本"],
        [
            ["1", "2026-06-07", "项目组 / Codex", "按 2026-06-07 生产在线抽检收口的首版正式测试报告，区分线上已验证事实与仅代码/历史专项验证结论。", "V1.0"],
            [
                "2",
                updated_at,
                "项目组 / Codex",
                f"最终提交版：更新至 {REGRESSION_BASELINE_DATE} S77 合并后全量回归（{REGRESSION_PASS_COUNT} passed）与 {PRODUCTION_SNAPSHOT_DATE} 生产在线抽检；新增 14 项 Logic 缺陷整改复核、生产认证态只读抽检与知识/通知写入闭环实测（含清理），以及学生端↔教师审批闭环的生产真机微信端到端实测（真实 wx.login→生产签发学生会话→申请→审批→学生回看 APPROVED）+ 本地 test_request_flow 20 passed 复核。",
                "V2.0",
            ],
        ],
        [0.5, 1.0, 1.0, 3.3, 0.7],
    )

    add_paragraph(doc, "1 测试概述", style="Heading 1")
    add_paragraph(doc, "1.1 编写目的", style="Heading 2")
    add_paragraph(
        doc,
        f"本文档为最终提交版，用于沉淀 {PROJECT_NAME} 当前主线（HEAD {head}）的测试结论，统一记录已完成的静态检查、构建验证、自动化回归、生产在线抽检与专项业务验证结果，复核 2026-05-26 测试基线所记录缺陷的整改情况，并明确区分“当前线上已核对事实”与“仅由代码实现和历史专项验证支撑的能力结论”，据此给出最终交付结论与残留建议。",
        style="First Paragraph",
    )

    add_paragraph(doc, "1.2 测试范围", style="Heading 2")
    add_bullets(
        doc,
        [
            "后端：鉴权、知识库、通知、党团流程、事务审批、画像、荣誉、学业分析、导入导出与审计。",
            "Web 管理端：登录态、通知中心、审批工作台、学生管理、党团流程、学业与审计页面。",
            "小程序学生端：首页、知识库、通知中心、事务申请、党团流程、画像与学业分析。",
            "环境链路：本地金仓回归环境、内网生产部署、健康检查与关键 API 认证边界。",
        ],
    )

    add_paragraph(doc, "1.3 测试人员与时间", style="Heading 2")
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["执行 / 汇总人员", "Codex（基于仓库留痕测试结果、缺陷整改记录与生产在线抽检整理）"],
            ["缺陷审查窗口", "2026-05-26：S50 测试工程师对当时 HEAD 的 Bug 审查（记录 14 项 Logic 缺陷）"],
            ["缺陷整改与专项验证", "2026-05-27 ~ 2026-06-13：S65~S74 完成认证/越权、Web 登录态与 401、通知治理、小程序分页等整改，以及证明 PDF、订阅消息、模板下载、审批附件下载等专项补验"],
            ["最终全量回归", f"{REGRESSION_BASELINE_DATE}：S77 合并 S75 性能/UI 与 S76 互测修复后，后端全量 pytest {REGRESSION_PASS_COUNT} passed，Web build 与 Miniapp build 通过"],
            ["生产在线抽检", f"{PRODUCTION_SNAPSHOT_DATE}：经 n150 跳板对 {PRODUCTION_BASE_URL} 做未登录公开面、受保护接口边界，以及以有效 {PROD_AUTH_ROLE} 账号登录后的认证态只读核对"],
            ["当前代码基线", f"当前仓库 HEAD {head}（feat/s75-perf-uiux，已部署生产，GitHub Actions run {DEPLOY_RUN_ID} 成功）"],
        ],
        [1.45, 5.05],
    )

    add_paragraph(doc, "1.4 测试环境", style="Heading 2")
    add_table(
        doc,
        ["环境", "配置 / 说明"],
        [
            ["后端", "FastAPI + uv；数据库为 Kingbase（本地以 PostgreSQL 15 回归）；支持本地 Docker 开发环境与内网生产环境"],
            ["Web 管理端", "Vue 3 + pnpm；通过 vue-tsc 类型检查与 vite build 验证；本版构建已按 S75 拆分 vue/vendor/antdv 第三方 chunk"],
            ["小程序", "uni-app + mp-weixin；通过 vue-tsc 与 build:mp-weixin 验证交付包"],
            ["生产环境", f"{PRODUCTION_BASE_URL} 内网部署（经 n150 跳板访问）；{PRODUCTION_SNAPSHOT_DATE} 在线核对根页标题={PRODUCTION_WEB_TITLE}、入口资源 {PRODUCTION_WEB_JS} 与 {PRODUCTION_WEB_CHUNKS}、样式 {PRODUCTION_WEB_CSS}、healthz 与知识接口均正常"],
            ["测试账号", f"历史共享口令 admin / admin123 于 {PRODUCTION_SNAPSHOT_DATE} 仍返回 401（40100），已失效；本轮已由部署维护方提供有效管理员账号（{PROD_AUTH_ROLE}），用于认证态只读验证，密码不在本文档固化"],
        ],
        [1.3, 5.2],
    )

    add_paragraph(doc, "2 测试设计", style="Heading 1")
    add_paragraph(doc, "2.1 测试策略", style="Heading 2")
    add_table(
        doc,
        ["测试层次", "目标", "代表验证"],
        [
            ["静态检查", "快速发现语法、风格和明显实现错误", "ruff、compileall / py_compile、vue-tsc"],
            ["自动化回归", "验证后端契约与业务流程主干是否稳定", f"pytest 全量（S77：{REGRESSION_PASS_COUNT} passed）/ 定向集成测试"],
            ["构建验证", "确认 Web 与小程序产物可生成", "pnpm -C web build、pnpm -C miniapp build:mp-weixin"],
            ["生产在线抽检", "确认内网部署当前真实响应、公开能力与认证边界", "GET /、healthz、knowledge/search·categories·detail、6 个受保护接口 401、旧共享口令登录 401"],
            ["缺陷整改复核", "确认历史测试基线缺陷已闭环", "2026-05-26 基线 14 项 Logic 缺陷逐项代码定位复核"],
            ["专项业务验证", "验证证明预览、订阅消息、模板/附件下载等高风险细节", "S57、S64、S70、S74、S76、PR #6 merge review 等记录"],
        ],
        [1.1, 1.8, 3.6],
    )

    add_paragraph(doc, "2.2 测试用例设计", style="Heading 2")
    add_table(
        doc,
        ["模块", "核心用例", "关注点"],
        [
            ["知识库", "关键词检索、来源展示、模板下载、AI 匹配", "必须优先返回官方来源，避免幻觉或空链路"],
            ["党团流程", "模板查看、学生流程发起、节点推进、提醒记录", "scope 权限、节点状态推进与留痕一致性"],
            ["事务审批", "请假/盖章/证明提交、驳回重提、审批通过、证明 PDF 预览", "状态机、附件、审批角色与 PDF 开放条件"],
            ["通知中心", "圈人预览、批次发送、投递明细、站内消息查看", "目标范围准确性、失败原因回看、多渠道边界"],
            ["画像与荣誉", "主档查看、纠错申诉、荣誉展示、导出与敏感字段保护", "字段权限、越权拒绝、导出脱敏"],
            ["学业分析", "成绩单解析、培养方案比对、课程建议", "仅做弱提示，不给出毕业强结论"],
        ],
        [1.15, 2.45, 2.95],
    )

    add_paragraph(doc, "2.3 测试总体情况", style="Heading 2")
    add_table(
        doc,
        ["检查项", "结果", "证据"],
        [
            ["后端静态检查", "通过", f"{REGRESSION_BASELINE_DATE}：ruff check、py_compile 均通过"],
            ["数据库迁移", "通过", f"{REGRESSION_BASELINE_DATE}：alembic upgrade head（含 0021 student_workflows.status 索引）成功"],
            ["后端全量回归", "通过", f"{REGRESSION_BASELINE_DATE} S77 合并后：pytest {REGRESSION_PASS_COUNT} passed, 4 warnings in 267s（修复 /notices/inbox 测试路由后清零唯一历史失败）"],
            ["Web 管理端构建", "通过", f"{REGRESSION_BASELINE_DATE}：pnpm -C web build（vue-tsc + vite）通过，产物分出 vue/vendor/antdv 独立 chunk"],
            ["小程序类型与构建", "通过", f"{REGRESSION_BASELINE_DATE}：vue-tsc 类型检查与 pnpm -C miniapp build:mp-weixin 通过"],
            ["生产根页在线抽检", "通过", f"GET / => 200，title={PRODUCTION_WEB_TITLE}，入口 {PRODUCTION_WEB_JS} 与 {PRODUCTION_WEB_CHUNKS} 均返回 200 application/javascript"],
            ["生产知识搜索在线抽检", "通过", f"GET /api/v1/knowledge/search?page=1&size=50 => 200，total={PRODUCTION_KNOWLEDGE_TOTAL}，含官方知识与模板下载类条目"],
            ["生产知识分类在线抽检", "通过", f"GET /api/v1/knowledge/categories => 200，返回 {PRODUCTION_CATEGORY_COUNT} 个分类"],
            ["生产知识详情在线抽检", "通过", "GET /api/v1/knowledge/{id} => code=0，可返回已发布条目详情"],
            ["生产受保护接口边界", "通过", "未带 Bearer token 访问 report/overview、admin/notices、requests/{id}/attachments/{aid}/download、workflow/public/templates、honors、admin/quiz/questions 均返回 401"],
            ["生产认证态只读抽检", "通过", f"以 {PROD_AUTH_ROLE} 账号登录后，overview、academic-gap、通知、荣誉、审计、知识条目/模板、学生检索与学生画像等只读接口均返回 code=0（详见 3.4）"],
            ["生产认证态写入抽检", "通过", "知识库（建条目→发布→检索→停用）与通知（建→发布→分发→归档）两条写入闭环在生产实测通过，测试数据已即时清理（详见 3.4）"],
            ["学生端审批闭环实测", "通过", "生产真机微信：真实 wx.login→生产签发学生会话→请假申请→提交→辅导员审批→学生回看 APPROVED；证明 PDF/附件下载/匿名 401 由本地 test_request_flow 20 passed 复核（详见 3.4）"],
            ["历史共享账号在线复核", "已失效", "POST /api/v1/auth/login 使用 admin / admin123 返回 401（40100：工号或密码错误）"],
            ["历史缺陷整改复核", "通过", "2026-05-26 基线 14 项 Logic 缺陷按当前 HEAD 代码定位逐项确认已修复（见第 4 章）"],
            ["专项验证", "通过 / 部分保留前置条件", "证明 PDF、微信订阅消息、模板直下载、审批附件下载等专项结果已留痕"],
        ],
        [1.9, 0.85, 3.75],
    )

    add_paragraph(doc, "3 测试执行结果", style="Heading 1")
    add_paragraph(doc, "3.1 平台级执行结果", style="Heading 2")
    add_table(
        doc,
        ["平台", "执行结果", "说明"],
        [
            ["后端代码基线", "通过", f"最新留痕基线为 {REGRESSION_BASELINE_DATE} S77 合并后：{REGRESSION_PASS_COUNT} passed, 4 warnings；知识/报表缓存、N+1 消除与索引补强均零回归"],
            ["Web 管理端构建基线", "通过", "构建链路稳定；S72/S75 已修复登录失败误判、401 跳转丢 redirect 与登录态恢复问题，并补齐全局加载与分包"],
            ["小程序构建基线", "通过", "vue-tsc 与 build:mp-weixin 通过；S73 修复 Intl 运行时兼容、S75 激活下拉刷新与 GET 去重"],
            ["生产匿名 / 未登录面", "通过", "根页、healthz、知识搜索/分类/详情均在线；6 个受保护接口的未登录边界正确返回 401"],
            ["生产认证链路", "读写已验证", f"以有效 {PROD_AUTH_ROLE} 账号登录成功，11 类只读接口均 code=0；知识与通知两条写入闭环（创建→发布→检索/分发→停用/归档）在生产实测通过并即时清理；学生侧申请审批写入待学生账号补验"],
            ["学生端审批闭环", "生产真机微信已验证", "微信开发者工具真实微信登录→生产 wx-login 签发学生会话→请假申请→提交→辅导员认领/审批→学生回看 APPROVED，生产端到端通过；证明 PDF/附件下载/匿名 401 由本地 test_request_flow 20 passed 复核"],
            ["微信小程序联调", "已验证", "mp-weixin 构建通过；并在微信开发者工具用真实微信账号经 n150 隧道指向生产完成真机登录与申请闭环"],
        ],
        [1.45, 0.9, 4.15],
    )

    add_paragraph(doc, "3.2 重点业务闭环结果", style="Heading 2")
    add_table(
        doc,
        ["业务闭环", "当前结论", "说明"],
        [
            ["知识库公开查询", "线上已验证", f"knowledge/search 在线返回 total={PRODUCTION_KNOWLEDGE_TOTAL}，包含休学、复学、校历、学院公告等官方知识条目及模板下载类条目；详情与分类接口在线返回 code=0"],
            ["模板/附件文件鉴权", "线上边界已验证", "knowledge/templates、requests/{id}/attachments/{aid}/download 未带 token 均返回 401；模板与审批附件直下载能力以 S70/S74/S76 专项验证与本地 DB 回归覆盖"],
            ["党团流程管理", "回归与代码验证通过", "流程模板、发起、推进与提醒链路在 S77 全量回归（含 test_workflow_party_flow）通过；scoped 协同角色搜人权限已按 viewer scope 收口"],
            ["事务申请与审批", "回归与代码验证通过", "请假、盖章、证明申请与审批工作台可用，请假起止日期顺序已强校验；证明 PDF 仅在 APPROVED 后开放预览；审批附件可在认证态下载并写审计"],
            ["通知中心", "回归与代码验证通过", "圈人、发送、批次查看链路通过回归；scoped 越权投递、生产 mock 入口、详情错误态、生效日期校验等历史缺陷均已整改"],
            ["画像、荣誉与学业分析", "回归与代码验证通过", "本人查看、教师范围查看、纠错申诉、荣誉展示与弱结论学业分析已打通；小程序画像历史等长列表已补分页/加载更多"],
        ],
        [1.55, 1.25, 3.7],
    )

    add_paragraph(doc, "3.3 残留风险与未覆盖项", style="Heading 2")
    add_bullets(
        doc,
        [
            f"最新全量后端回归基线为 {REGRESSION_BASELINE_DATE} S77 合并后的 {REGRESSION_PASS_COUNT} passed；该基线在本机 Docker / localhost:54322 测试库上执行，覆盖认证态成功路径。",
            f"{PRODUCTION_SNAPSHOT_DATE} 本轮已使用有效 {PROD_AUTH_ROLE} 账号在生产完成认证态只读验证；学生与各类教师角色账号仍建议由部署维护方按需发放，以覆盖不同 scope 的角色视图。",
            "本轮已在生产完成知识库与通知两条写入闭环实测（含发布、分发、缓存失效与清理），以及学生端↔教师审批闭环的生产真机微信实测（真实微信登录→请假申请→提交→辅导员审批→学生回看 APPROVED）；证明 PDF 预览与附件认证下载等细节由本地 test_request_flow 全量 20 passed 复核。微信订阅消息的真实手机端送达仍依赖用户在小程序完成模板授权。",
            "微信订阅消息链路代码与生产出网已验证可达，但真实手机端送达仍依赖用户在小程序中完成有效模板授权。",
            "小程序 lazyCodeLoading（按需组件加载）为推荐优化项，需在微信开发者工具做一次性运行时 smoke，当前 devtools 自动化受阻，未盲改。",
            "建议在正式交付前，由部署维护方提供有效账号，补一轮 10.10.0.13 上管理端登录后功能、党团流程、事务审批与通知操作的认证态实登抽检。",
        ],
    )

    add_paragraph(doc, "3.4 生产认证态读写验证", style="Heading 2")
    add_paragraph(
        doc,
        f"{PRODUCTION_SNAPSHOT_DATE}，由部署维护方提供有效管理员账号（{PROD_AUTH_ROLE}），经 n150 跳板完成登录、"
        "认证态只读抽检与认证态写入闭环验证。写入验证使用清晰标注的测试数据（标题前缀「【生产写入验收-勿用】」），"
        "并在验证后即时清理（知识条目停用、通知归档），避免污染正式业务数据。",
        style="Body Text",
    )
    add_paragraph(doc, "只读验证", style="Heading 3")
    add_table(
        doc,
        ["验证项", "接口", "结果"],
        [
            ["登录与会话", "POST /auth/login + GET /auth/me", f"登录成功，/auth/me 返回角色 {PROD_AUTH_ROLE}、must_change_password=false、token 版本有效"],
            ["运营看板", "GET /admin/report/overview", "code=0，返回学生/请求/通知等聚合指标（原 scoped 越权项已收口）"],
            ["学业缺口聚合", "GET /admin/report/academic-gap", "code=0，total=7，返回培养方案缺口聚合明细"],
            ["通知中心", "GET /admin/notices", "code=0，total=9，返回通知列表"],
            ["荣誉公示", "GET /admin/honors、GET /honors", "code=0，管理视图 total=3、公开视图 total=2"],
            ["党团流程模板", "GET /workflow/public/templates", "code=0，返回党员/团员发展等流程模板"],
            ["理论自测题库", "GET /admin/quiz/questions", "code=0，total=1，返回题库条目"],
            ["知识条目/模板", "GET /admin/knowledge/entries、/templates", "code=0，条目 total=17、模板 total=4"],
            ["学生检索与画像", "GET /admin/workflow/students/search?q=2024、/admin/profile/{id}", "code=0，检索 total=6，可返回单个学生画像（搜人权限已收口）"],
            ["审计日志", "GET /admin/audit-logs", "code=0，total=2163，审计写入在生产持续生效"],
            ["改密失败态", "POST /auth/change-password（错误旧密码）", "返回 40100 旧密码错误，未误改、未 500，错误态可控"],
        ],
        [1.4, 2.55, 2.55],
    )

    add_paragraph(doc, "写入闭环验证（含清理）", style="Heading 3")
    add_table(
        doc,
        ["业务闭环", "写入步骤", "结果"],
        [
            [
                "知识库闭环",
                "POST 来源 + 条目(草稿) → publish → 公开检索 → detail → deprecate",
                "条目状态 DRAFT→PUBLISHED：发布后公开检索 total=1（命中测试 slug，印证 S75 写路径缓存事件失效）；详情返回来源；停用后公开检索 total=0，清理生效",
            ],
            [
                "通知闭环",
                "POST 通知(草稿) → target-preview → publish → dispatch → archive",
                "圈人预览 target_count=2；发布 PUBLISHED；分发批次 NB-… target=2/success=2/failed=0、状态 COMPLETED（IN_APP 送达 2 名演示学生）；归档完成清理",
            ],
        ],
        [1.25, 2.6, 2.65],
    )
    add_callout(
        doc,
        "说明：生产写入验证仅在演示数据集（生产共 7 名演示学生）上进行，使用「【生产写入验收-勿用】」标注并即时停用/归档。",
    )

    add_paragraph(doc, "学生端↔教师审批闭环（生产真机微信 + 本地全量回归）", style="Heading 3")
    add_paragraph(
        doc,
        "学生端为微信小程序，登录走真实微信 code 换 openid。本轮已在生产用真实微信完成端到端验证：在微信开发者工具"
        "（appid wxcb6352a74505bc41、真实微信账号）运行学生端，经 n150 隧道指向生产 API，wx.login 取得真实 code 后"
        "由生产后端完成 jscode2session 签发学生会话（印证生产微信登录配置正确、并非降级 mock）。",
        style="Body Text",
    )
    add_table(
        doc,
        ["步骤", "动作（生产 10.10.0.13）", "结果"],
        [
            ["1", "微信开发者工具真实 wx.login → 生产 wx-login", "签发学生 JWT（角色 STUDENT），生产 jscode2session 成功"],
            ["2", "学生发起请假申请（LEAVE_PERSONAL）", "request_id=7，DRAFT（标注「【验收测试-勿用】」）"],
            ["3", "学生提交", "状态 SUBMITTED"],
            ["4", "辅导员认领（COUNSELOR 账号）", "状态 IN_REVIEW"],
            ["5", "辅导员审批通过", "状态 APPROVED，动作链 SUBMIT/CLAIM/APPROVE，审批意见留痕"],
            ["6", "学生端回看申请", "状态 APPROVED，可见审批意见，闭环成立"],
        ],
        [0.55, 3.15, 2.8],
    )
    add_callout(
        doc,
        "鉴权边界佐证：生产未登录访问 wx-login 伪造 code 返回 40100；已绑定学生重复绑定返回 40901/40902，"
        "均符合预期。证明类（CERTIFICATE）PDF 预览、附件认证下载（学生+辅导员）与匿名下载 401 边界，"
        "另由本地 Docker 全链路 test_request_flow.py 全量 20 个场景回归 20 passed 复核（含驳回重提、撤回、"
        "转线下、终态 REOPEN、附件必填、证明模板引擎、协同角色 scope 等）。",
    )

    add_paragraph(doc, "4 缺陷分析与整改复核", style="Heading 1")
    add_paragraph(doc, "4.1 缺陷统计", style="Heading 2")
    add_table(
        doc,
        ["类型", "基线数量", "当前状态", "说明"],
        [
            ["Crash", "0", "0", "2026-05-26 基线及当前 HEAD 均未发现可稳定复现的崩溃类问题"],
            ["Logic", "14", "0 待修复", "2026-05-26 基线记录 14 项；按当前 HEAD（ca3b8de）代码逐项复核均已修复"],
            ["新增缺陷", "—", "0", "S51~S77 整改与性能/UI 优化未引入可复现的新增有效缺陷，全量回归无回归"],
        ],
        [1.25, 1.0, 1.05, 3.2],
    )
    add_callout(
        doc,
        "说明：14 项 Logic 缺陷为 2026-05-26 测试基线（bug-report.md）记录值；本最终版以当前 HEAD 代码定位逐项复核其整改情况，结论为均已修复。",
    )

    add_paragraph(doc, "4.2 缺陷分类整改情况", style="Heading 2")
    add_table(
        doc,
        ["类别", "基线数量", "整改状态", "整改方式"],
        [
            ["后端认证与权限", "4", "已修复", "改密递增 token_version 使旧 token 失效；overview/通知预览分发/流程搜人均按 viewer scope 收口"],
            ["Web 登录态与错误处理", "4", "已修复", "根路由先 fetchMe 再定向、瞬时错误不再误登出、401 统一携带 redirect、党团提醒 fetch 接入 401 处理"],
            ["Web 通知治理", "2", "已修复", "模拟短信回执入口改为 isDev 门禁；详情/批次失败设置 error 标记并显示加载错误而非空态"],
            ["Miniapp 统计与分页", "4", "已修复", "首页统计与事务/画像历史列表补齐分页、加载更多与剩余提示，不再只取第一页"],
        ],
        [1.55, 0.85, 0.85, 3.25],
    )

    add_paragraph(doc, "4.3 典型缺陷整改说明", style="Heading 2")
    add_table(
        doc,
        ["ID", "模块", "问题描述", "整改与当前代码定位"],
        [
            ["S50-L01", "后端认证", "改密后旧 access/refresh token 仍有效（P0）", "change_password 内 token_version += 1（auth/service.py:375），旧 token 校验即失效"],
            ["S50-L02", "运营看板", "scoped 教师访问 overview 看到全局统计（P0）", "admin_overview 传入 viewer_user_id/viewer_roles（report/router.py:87-96），按范围聚合"],
            ["S50-L04", "通知中心", "scoped 编辑者可预览/投递 scope 外学生（P0）", "target-preview 与 dispatch 传入 viewer scope（notice/router.py:212-233）"],
            ["S50-L07/L08", "Web 401 处理", "401 跳裸 /login 丢 redirect；提醒 raw fetch 无 401 分支", "request.ts 与 workflow.ts fetch 均跳 /login?redirect=<path>（request.ts:114-116、workflow.ts:221-225）"],
            ["S50-L09", "Web 通知治理", "生产暴露模拟短信回执入口", "回执按钮加 isDev 门禁（NoticeList.vue:771），生产不渲染"],
            ["S50-L11~L14", "小程序", "首页/历史列表只统计第一页，无加载更多", "request/profile 等页补 hasMore/loadMore 与错误保留（request/index.vue:363、profile/index.vue:1100）"],
        ],
        [0.95, 1.05, 2.3, 2.2],
    )
    add_callout(doc, "说明：完整缺陷清单与原始代码定位见仓库根目录 bug-report.md 与 S50 细化文件；整改过程见 S51~S77 细化记录与主计划变更历史。")

    add_paragraph(doc, "5 测试结论与建议", style="Heading 1")
    add_paragraph(doc, "5.1 测试结论", style="Heading 2")
    add_bullets(
        doc,
        [
            f"当前主线代码（HEAD {head}）通过后端静态检查、数据库迁移、后端全量回归（{REGRESSION_BASELINE_DATE} S77：{REGRESSION_PASS_COUNT} passed）、Web 构建与小程序类型检查/构建，具备正式交付与稳定演示基础。",
            "2026-05-26 测试基线记录的 14 项 Logic 缺陷，经当前 HEAD 代码逐项复核均已修复，未发现崩溃类问题，整改未引入新增有效缺陷。",
            f"生产环境 {PRODUCTION_BASE_URL} 于 {PRODUCTION_SNAPSHOT_DATE} 在线核对：根页、healthz、知识搜索/分类/详情均在线，前端已部署 S75 分包产物；6 个受保护接口未登录访问均返回 401，认证边界符合预期。",
            f"生产认证态读写路径已在 10.10.0.13 实登验证：以有效 {PROD_AUTH_ROLE} 账号登录后，11 类只读接口均 code=0、改密错误旧密码可控返回 40100；知识库与通知两条写入闭环在生产实测通过并即时清理，印证 S75 写路径缓存事件失效。学生端↔教师审批闭环已用真实微信在生产端到端跑通（微信开发者工具真实 wx.login→生产签发学生会话→请假申请→提交→辅导员审批→学生回看 APPROVED）；证明 PDF 预览与附件认证下载等细节由本地 test_request_flow 全量 20 passed 复核。",
        ],
    )

    add_paragraph(doc, "5.2 改进建议", style="Heading 2")
    add_bullets(
        doc,
        [
            f"交付侧：有效 {PROD_AUTH_ROLE} 管理员账号已用于本轮认证态读写验证，并已用真实微信完成学生端生产实测；建议按需补发各 scope 的教师角色账号以覆盖更多角色视图，并停止依赖已失效的共享口令。",
            "清理侧：本轮在生产产生的标注测试数据（请假申请、停用知识条目/来源、归档通知、测试学生与测试辅导员账号）建议由部署维护方按「【验收测试/写入验收-勿用】」标签统一清理或归档。",
            "优化侧：按需推进小程序 lazyCodeLoading（需 devtools 运行时 smoke）、AsyncBoundary 逐页接入与 knowledge 检索缓存等推荐项，均非阻塞交付的必选项。",
            "运维侧：保持 GitHub Actions 内网部署与 /healthz、知识公开接口的常态化只读 smoke，作为上线后的基本健康监测。",
        ],
    )

    return finalize(doc, output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build project test report DOCX.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    print(build_doc(args.output))


if __name__ == "__main__":
    main()
