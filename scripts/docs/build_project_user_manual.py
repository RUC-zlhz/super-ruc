from __future__ import annotations

import argparse
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_common import (
    DEEP_BLUE,
    MUTED,
    ROOT,
    add_bullets,
    add_callout,
    add_code_block,
    add_direct_paragraph,
    add_kv_table,
    add_paragraph,
    add_spacer,
    add_steps_table,
    add_table,
    finalize,
    prepare_document,
    today_label,
)


TEMPLATE = ROOT / "docs" / "templates" / "用户使用说明书模板.docx"
DEFAULT_OUTPUT = ROOT / "output" / "doc" / "用户使用说明书-信息学院学生综合服务与党团管理平台-v1.0.docx"


PROJECT_NAME = "信息学院学生综合服务与党团管理平台"
PRODUCTION_BASE_URL = "http://10.10.0.13/"
PRODUCTION_SNAPSHOT_DATE = "2026-06-07"
MINIAPP_BUILD_ENTRY = "miniapp/dist/build/mp-weixin"


def build_doc(output_path: Path) -> Path:
    doc = prepare_document(TEMPLATE)
    updated_at = today_label()

    add_direct_paragraph(doc, "文档编号：SIP-UM-1.0", size=11, color=MUTED)
    add_direct_paragraph(
        doc,
        PROJECT_NAME,
        size=22,
        bold=True,
        color=DEEP_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_direct_paragraph(
        doc,
        "用户使用说明书",
        size=18,
        bold=True,
        color=DEEP_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_direct_paragraph(
        doc,
        f"版本：V1.0    更新日期：{updated_at}",
        size=11,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    add_kv_table(
        doc,
        [
            ("适用对象", "学院管理员、辅导员/班主任/党团教师、班团骨干、学生"),
            ("适用范围", "Web 管理端、微信小程序学生端、部署维护与课程验收场景"),
            ("生产核对基准", f"{PRODUCTION_SNAPSHOT_DATE} 已在线核对 {PRODUCTION_BASE_URL}、/healthz、知识搜索与分类接口"),
            ("推荐入口", f"生产 Web：{PRODUCTION_BASE_URL}；健康检查：{PRODUCTION_BASE_URL}healthz；小程序产物：{MINIAPP_BUILD_ENTRY}"),
            ("文档定位", "用于真实部署环境使用说明、互测交接和课程验收答辩；内容以线上事实和当前代码实现为准"),
        ],
    )

    add_spacer(doc)
    add_paragraph(doc, "文档变更历史记录", style="Heading 2")
    add_table(
        doc,
        ["序号", "日期", "编写人", "变更内容", "版本"],
        [["1", updated_at, "项目组 / Codex", "按 2026-06-07 线上部署事实与当前代码实现同步整理的正式用户使用说明书。", "V1.0"]],
        [0.55, 1.0, 1.0, 3.25, 0.7],
    )

    add_spacer(doc)
    add_paragraph(doc, "目录", style="Heading 1")
    add_bullets(
        doc,
        [
            "1 系统概述",
            "2 系统运行环境",
            "3 系统登录与退出",
            "4 公共功能说明",
            "5 管理员功能使用说明",
            "6 教师功能使用说明",
            "7 学生功能使用说明",
            "8 常见问题与解决方法",
            "9 联系方式与技术支持",
        ],
    )

    add_paragraph(doc, "1 系统概述", style="Heading 1")
    add_paragraph(doc, "1.1 编写目的", style="Heading 2")
    add_paragraph(
        doc,
        f"本文档面向系统实际使用者，说明 {PROJECT_NAME} 的访问方式、角色边界、核心功能和常见操作步骤。文档内容以 {PRODUCTION_SNAPSHOT_DATE} 对 {PRODUCTION_BASE_URL} 的在线核对结果和当前仓库代码实现为准，帮助管理员、教师和学生在统一口径下完成日常事务办理、通知触达、流程跟进和学业辅助工作。",
        style="Body Text",
    )

    add_paragraph(doc, "1.2 系统简介", style="Heading 2")
    add_paragraph(
        doc,
        "本系统由 Web 管理端与微信小程序学生端组成，围绕“知识库问答、党团事务流程、通知精准推送、事务申请与电子证明、学生画像与学业分析”五条业务主线设计，支持以学院侧自管方式沉淀政策、流程和数据，并通过金仓数据库与审计机制保证关键数据可追溯。当前生产环境已确认可直接访问 Web 管理端；学生端权威交付形态仍以微信小程序构建产物和部署维护方分发的测试包为准。",
        style="Body Text",
    )

    add_paragraph(doc, "1.3 系统功能概览", style="Heading 2")
    add_table(
        doc,
        ["模块", "主要能力", "典型使用者"],
        [
            ["知识库与模板", "政策知识检索、AI 辅助匹配、常用模板下载、知识条目与模板维护", "学生、教师、管理员"],
            ["党团流程", "流程模板维护、学生流程发起、节点推进、提醒记录、理论自测", "学生、辅导员、党团教师、班团骨干"],
            ["通知中心", "通知录入、标签化圈人、批次发送、站内消息查看、订阅消息配置", "管理员、教师、学生"],
            ["事务申请与证明", "请假/盖章/证明申请、附件上传、审批工作台、审批意见留痕、批准后 PDF 预览", "学生、教师"],
            ["画像与学业分析", "学生主档、成长事实、荣誉展示、成绩单核验、培养方案比对、课程建议", "学生、教师、管理员"],
        ],
        [1.2, 3.15, 2.15],
    )

    add_paragraph(doc, "1.4 用户角色说明", style="Heading 2")
    add_table(
        doc,
        ["角色", "说明", "主要权限"],
        [
            ["超级管理员", "学院系统全局维护角色", "后台账号、模板、知识库、通知、学生主档、审计与系统配置全局可管"],
            ["学院领导", "学院级业务查看与统筹角色", "查看全局统计、重点台账和部分审批/导出结果"],
            ["辅导员 / 班主任 / 党团教师", "一线业务处理角色", "在权限范围内审批申请、查看学生画像、推进流程、发送通知、核验成绩单"],
            ["班团骨干", "协同处理流程或提醒的学生干部角色", "在授权范围内查看班级/支部相关进度、接收并处理协同任务"],
            ["学生", "小程序主使用者", "查看通知、检索知识、下载模板、提交申请、查看画像与学业分析结果"],
        ],
        [1.2, 1.65, 3.65],
    )

    add_paragraph(doc, "2 系统运行环境", style="Heading 1")
    add_paragraph(doc, "2.1 客户端环境", style="Heading 2")
    add_table(
        doc,
        ["终端", "推荐环境", "说明"],
        [
            ["Web 管理端", "Windows 10/11 或 macOS，Chrome / Edge 新版本浏览器", "适合老师和管理员进行配置、审批、导入导出和统计分析"],
            ["微信小程序", "微信客户端测试包或微信开发者工具导入当前 mp-weixin 产物", f"适合学生查询通知、发起申请、查看画像和学业分析结果；当前权威产物目录为 {MINIAPP_BUILD_ENTRY}"],
            ["本地联调", "Docker Desktop + uv + pnpm + PowerShell", "适合课程答辩前联调、离线演示和问题复现"],
        ],
        [1.2, 2.4, 2.9],
    )

    add_paragraph(doc, "2.2 浏览器要求", style="Heading 2")
    add_bullets(
        doc,
        [
            "推荐使用最新版 Chrome 或 Microsoft Edge；不建议使用 IE 或兼容模式浏览器。",
            "首次访问远端环境时，请确认浏览器允许本地存储，以便保存登录态与权限信息。",
            "涉及文件导入、模板下载、PDF 预览时，建议使用桌面浏览器打开，以获得更稳定的上传与下载体验。",
            "微信小程序联调应使用部署维护方提供的测试包或当前仓库 mp-weixin 构建产物，不再假定旧版私有分享入口仍然有效。",
        ],
    )

    add_paragraph(doc, "2.3 系统访问方式", style="Heading 2")
    add_table(
        doc,
        ["场景", "访问方式", "备注"],
        [
            ["生产 Web", PRODUCTION_BASE_URL, "2026-06-07 在线核对可返回“信息学院管理后台”首页"],
            ["生产健康检查", f"{PRODUCTION_BASE_URL}healthz", "2026-06-07 在线核对返回 code=0、status=ok"],
            ["生产知识搜索", f"{PRODUCTION_BASE_URL}api/v1/knowledge/search?page=1&page_size=20", "当前线上可返回官方知识与模板下载条目"],
            ["生产知识分类", f"{PRODUCTION_BASE_URL}api/v1/knowledge/categories", "2026-06-07 在线核对返回 9 个分类"],
            ["本地开发 Web", "运行 scripts/dev/start-dev.ps1 后访问 http://localhost:4173", "适合离线演示与问题复现"],
            ["本地后端文档", "http://localhost:8080/docs", "适合接口联调和参数核对"],
            ["微信小程序", f"导入 {MINIAPP_BUILD_ENTRY} 或由部署维护方提供测试包", "适合学生视角验证；真实登录方式取决于当前部署配置"],
        ],
        [1.05, 2.75, 2.7],
    )

    add_paragraph(doc, "3 系统登录与退出", style="Heading 1")
    add_paragraph(doc, "3.1 用户登录", style="Heading 2")
    add_paragraph(doc, "操作步骤", style="Heading 3")
    add_steps_table(
        doc,
        [
            f"管理员或教师在浏览器中打开 Web 地址 {PRODUCTION_BASE_URL}，输入部署维护方分配的工号/账号与密码登录。",
            "若当前账号为系统生成的初始密码，登录后可能出现“请尽快修改当前初始密码”提醒，可根据是否属于共享验收账号决定立即修改或稍后处理。",
            "学生通过部署维护方提供的小程序测试包或当前 mp-weixin 产物进入“我的”页面完成登录/绑定；真实微信登录、测试态登录或绑定方式以当前部署配置为准。",
            "首次登录成功后，系统会根据角色自动加载可访问菜单；若角色具有范围限制，只会展示本角色允许访问的数据。",
        ],
    )
    add_paragraph(doc, "登录界面说明", style="Heading 3")
    add_bullets(
        doc,
        [
            "本轮 2026-06-07 在线核对中，历史互测文档曾使用的共享口令 admin / admin123 已不再可登录；正式使用必须以当前部署维护方发放的账号为准。",
            "登录后如果弹出“请尽快修改当前初始密码”，说明该账号仍处于系统生成密码阶段；正式使用建议尽快改密。",
            "小程序学生端的真实微信登录、绑定或测试态登录方式取决于当前部署配置；如无法进入，请先确认学生主档已导入且测试包来源正确。",
        ],
    )

    add_paragraph(doc, "3.2 密码修改", style="Heading 2")
    add_paragraph(doc, "操作步骤", style="Heading 3")
    add_steps_table(
        doc,
        [
            "登录后进入右上角个人菜单或个人中心页面。",
            "点击“修改密码”，输入旧密码、新密码和确认密码。",
            "提交成功后重新登录；如为正式管理账号，建议立即修改系统生成的初始密码。",
        ],
    )
    add_callout(doc, "说明：系统当前会区分“系统生成的初始密码”和“用户主动修改后的密码”，正式管理环境不建议长期保留默认口令。")

    add_paragraph(doc, "3.3 用户退出", style="Heading 2")
    add_paragraph(
        doc,
        "Web 端可通过右上角头像菜单点击“退出登录”；小程序端可在“我的”页面退出当前身份。退出后本地 token 会被清空，再次进入需要重新登录或重新绑定。",
        style="Body Text",
    )

    add_paragraph(doc, "4 公共功能说明", style="Heading 1")
    add_paragraph(doc, "4.1 首页功能", style="Heading 2")
    add_bullets(
        doc,
        [
            "Web 首页提供角色默认入口、统计卡片或快捷菜单，便于管理员快速进入通知、审批、学生管理和流程管理页面。",
            "小程序首页提供通知、知识库、事务申请、党团流程和学业分析等高频入口，并展示未读通知和待跟进事项。",
            "首页数据受当前角色和权限范围限制，不同用户看到的卡片与数量可能不同。",
        ],
    )

    add_paragraph(doc, "4.2 个人信息管理", style="Heading 2")
    add_bullets(
        doc,
        [
            "Web 端个人信息页用于查看当前账号、角色标签和登录态信息。",
            "学生端个人中心可查看个人基础信息、画像摘要、申请历史和敏感字段的“完整查看申请”状态。",
            "当个人信息存在错误时，学生可通过纠错申诉或成长补录入口提交补充信息，由教师侧审核处理。",
        ],
    )

    add_paragraph(doc, "4.3 修改密码", style="Heading 2")
    add_paragraph(
        doc,
        "除首次登录后的改密提示外，系统在个人中心中始终保留修改密码入口。管理员和教师建议定期更新密码；若账号由系统生成初始密码，应在首次登录后尽快完成改密。",
        style="Body Text",
    )

    add_paragraph(doc, "5 管理员功能使用说明", style="Heading 1")
    add_paragraph(doc, "5.1 学生信息与微信绑定管理", style="Heading 2")
    add_paragraph(doc, "功能说明", style="Heading 3")
    add_paragraph(
        doc,
        "管理员可维护学生主档、查看微信绑定状态、处理解绑请求，并通过导入导出能力批量维护学院学生基础数据。该模块是知识、通知、审批和画像等功能的基础数据入口。",
        style="Body Text",
    )
    add_paragraph(doc, "操作步骤", style="Heading 3")
    add_steps_table(
        doc,
        [
            "进入“用户管理 / 学生管理”页面，按学号、姓名、班级或年级检索学生。",
            "点击“新增学生”可录入单个学生主档；点击导入入口可按模板批量导入学生数据。",
            "打开学生详情后，可查看微信绑定状态；必要时执行解绑，允许学生重新绑定新微信。",
        ],
    )
    add_paragraph(doc, "学生信息字段", style="Heading 3")
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["学号", "学生唯一标识，也是学生端登录绑定的重要依据"],
            ["姓名", "用于审批、通知、画像和流程展示"],
            ["年级 / 专业 / 班级", "用于 scope 控制、通知圈人和学业分析"],
            ["学籍状态", "用于控制是否可编辑、是否可进入某些业务流程"],
            ["手机号 / 身份证号", "敏感字段，按权限控制并在存储和导出中执行脱敏或加密"],
            ["微信绑定状态", "用于判断当前学生是否已绑定微信和是否需要重新绑定"],
        ],
        [1.7, 4.8],
    )

    add_paragraph(doc, "5.2 知识库与常用模板管理", style="Heading 2")
    add_bullets(
        doc,
        [
            "进入“知识库管理”页面后，可在“知识条目”页签新增、编辑、发布或停用知识条目，维护标准答复、来源链接、标签和摘要。",
            "同一页面的“模板文件”页签支持上传、下载、停用党员证明、团员证明、教室借用等常用模板，并可把模板关联到知识条目。",
            "知识条目与模板文件都直接服务学生端检索和下载，因此更新政策或模板时应同步停用旧版本并核对关联关系。",
        ],
    )

    add_paragraph(doc, "5.3 党团流程模板与提醒管理", style="Heading 2")
    add_bullets(
        doc,
        [
            "管理员可维护党员发展、团员发展与团籍管理等流程模板，配置阶段、节点说明和推进顺序。",
            "在“学生流程”页可以为指定学生发起流程实例，并查看当前节点、已完成动作和提醒记录。",
            "提醒引擎会根据配置自动产生提醒；对特殊情况可人工标记延期或手动跟进。",
        ],
    )

    add_paragraph(doc, "5.4 通知中心与精准推送", style="Heading 2")
    add_bullets(
        doc,
        [
            "通知中心支持录入通知标题、正文、标签、目标规则和发送渠道，并可预览命中学生人数。",
            "发送后可查看批次记录、投递明细、阅读状态和失败原因，方便做结果回看和治理。",
            "通知推送默认支持站内消息；如启用微信订阅消息或邮件渠道，应提前完成对应配置与授权验证。",
        ],
    )

    add_paragraph(doc, "5.5 培养方案、成绩核验与学业分析", style="Heading 2")
    add_bullets(
        doc,
        [
            "管理员可维护培养方案、课程模块和当前学期开课信息，作为学业分析的对比基线。",
            "学生上传成绩单 PDF 后，系统会生成解析候选，教师或管理员核验通过后才进入正式成绩记录。",
            "管理端学业分析页展示培养方案缺口、课程建议与风险提示，但不会给出自动毕业结论。",
        ],
    )

    add_paragraph(doc, "5.6 用户权限、审计与系统配置", style="Heading 2")
    add_bullets(
        doc,
        [
            "后台账号支持按角色和范围配置权限，例如学院领导、辅导员、班主任、班团骨干等。",
            "审计日志会记录审批、导入、导出、配置修改和敏感访问等关键动作，便于追责与验收。",
            "系统配置项包括通知渠道、学业推荐学期、默认模板、归档任务等，修改前应确认影响范围。",
        ],
    )

    add_paragraph(doc, "6 教师功能使用说明", style="Heading 1")
    add_paragraph(doc, "6.1 事务审批工作台", style="Heading 2")
    add_paragraph(
        doc,
        "辅导员、班主任和业务教师可在审批工作台查看待处理申请，核对表单信息与附件后执行通过、驳回或转线下。驳回后学生可在原草稿基础上修改重提，状态与审批意见都会留痕。",
        style="Body Text",
    )

    add_paragraph(doc, "6.2 党团流程推进", style="Heading 2")
    add_paragraph(doc, "操作步骤", style="Heading 3")
    add_steps_table(
        doc,
        [
            "进入“党团流程管理 / 学生流程”，按学生姓名、学号或班级检索目标学生。",
            "打开流程详情后查看当前节点、历史节点和学生提交材料。",
            "确认节点已完成后点击推进；如节点需要学生补充材料，可先退回或标记待补充。",
        ],
    )

    add_paragraph(doc, "6.3 学生画像、荣誉与名单查看", style="Heading 2")
    add_bullets(
        doc,
        [
            "教师可在权限范围内查看学生画像，包括主档、成长事实、纠错申诉和完整查看申请状态。",
            "荣誉模块支持按类别、学年和历史状态查看公开荣誉，便于做展示或评优参考。",
            "敏感字段默认按最小必要原则展示，越权访问会被拒绝并记入审计日志。",
        ],
    )

    add_paragraph(doc, "6.4 学业分析与课程建议复核", style="Heading 2")
    add_bullets(
        doc,
        [
            "教师可查看学生成绩单解析候选，并将确认后的成绩写入正式记录。",
            "学业缺口结果用于辅助学生自查和老师答疑，课程建议分为“本学期开课”与“培养方案候选”两类来源。",
            "当开课数据不足时，页面会显示 warning，不应用该模块直接替代毕业资格审核。",
        ],
    )

    add_paragraph(doc, "6.5 通知查看与业务协同", style="Heading 2")
    add_bullets(
        doc,
        [
            "教师可查看自己范围内的通知批次、投递情况和阅读状态，便于督办关键通知。",
            "班团骨干可在授权范围内接收协同任务或提醒，但默认不拥有学院级全局数据权限。",
            "如需联合处理学生事务，建议通过通知或流程节点方式留痕，不建议脱离系统线下口头流转。",
        ],
    )

    add_paragraph(doc, "7 学生功能使用说明", style="Heading 1")
    add_paragraph(doc, "7.1 个人信息与画像查看", style="Heading 2")
    add_bullets(
        doc,
        [
            "学生可在小程序“我的”页面查看自己的基础信息、画像摘要和成长事实。",
            "如果发现个人信息不准确，可发起纠错申诉；如需补充竞赛、科研、实践等事实，可提交成长补录申请。",
            "对身份证号、手机号等敏感字段，学生端默认仅展示脱敏结果，如需完整查看需单独申请并等待审核。",
        ],
    )

    add_paragraph(doc, "7.2 知识检索与模板下载", style="Heading 2")
    add_steps_table(
        doc,
        [
            "进入“知识库”页面，在搜索框输入关键词，例如“奖学金”“请假”“入党流程”。",
            "查看系统返回的摘要、命中原因和来源文件，必要时点击来源链接或原文件继续核对。",
            "进入“常用模板”区域下载党员证明、团员证明、教室借用审批表等标准模板。",
        ],
    )

    add_paragraph(doc, "7.3 事务申请、审批状态与证明预览", style="Heading 2")
    add_bullets(
        doc,
        [
            "学生可发起请假、盖章、证明等事务申请，按要求填写用途说明并上传附件。",
            "申请详情页会显示状态流转、审批意见和历史动作；被驳回后可在原草稿基础上修改重提。",
            "证明 PDF 预览仅在申请审批通过后开放，草稿或已提交但未批准状态下不会生成 PDF。",
        ],
    )

    add_paragraph(doc, "7.4 党团流程与理论自测", style="Heading 2")
    add_bullets(
        doc,
        [
            "在“党团中心”可查看自己当前所处阶段、已完成节点和下一步待办。",
            "当节点要求提交材料时，页面会明确提示需要补充的内容和提交时机。",
            "如教师端已导入题库，学生还可在理论自测中完成党建或团务学习测验。",
        ],
    )

    add_paragraph(doc, "7.5 通知中心与学业分析", style="Heading 2")
    add_bullets(
        doc,
        [
            "通知中心会集中展示发送给当前学生的站内通知，支持查看已读状态和消息详情。",
            "学业分析页面可展示培养方案缺口、成绩单核验状态和课程建议，但结论仅供参考，不代表最终毕业审核结果。",
            "如果系统提示“暂无流程实例”“暂无通知”或“暂无学业建议”，应先确认对应数据是否已经由老师创建或导入。",
        ],
    )

    add_paragraph(doc, "8 常见问题与解决方法", style="Heading 1")
    add_paragraph(doc, "问题1：无法登录系统", style="Heading 2")
    add_bullets(
        doc,
        [
            "先确认访问地址是否正确：生产环境推荐使用 http://10.10.0.13/，本地环境使用 start-dev.ps1 输出的地址。",
            "若是 Web 管理端，请使用部署维护方当前发放的账号密码，不要继续参考旧版互测文档中的共享口令说明。",
            "若是小程序学生端，请确认测试包或开发者工具工程来自当前版本，且学生主档、微信绑定或测试登录配置已准备完成。",
        ],
    )

    add_paragraph(doc, "问题2：为什么申请详情里看不到证明 PDF", style="Heading 2")
    add_bullets(
        doc,
        [
            "当前系统只在申请状态为 APPROVED 时开放证明 PDF 预览。",
            "如果申请仍处于草稿、已提交或驳回状态，系统不会生成 PDF。",
            "若已批准仍无法打开，请联系教师确认该申请类型是否绑定了有效证明模板。",
        ],
    )

    add_paragraph(doc, "问题3：为什么通知、流程或学业结果为空", style="Heading 2")
    add_bullets(
        doc,
        [
            "通知为空通常表示当前学生尚未被纳入目标规则，或老师尚未发送对应通知。",
            "流程为空通常表示管理员还没有为该学生发起党团流程实例。",
            "学业建议为空通常表示成绩单尚未核验入库，或当前学期开课数据还未维护完成。",
        ],
    )

    add_paragraph(doc, "问题4：本地小程序调试连不上后端", style="Heading 2")
    add_bullets(
        doc,
        [
            "先确认 start-dev.ps1 已成功启动本地后端，并且 http://127.0.0.1:8080/healthz 返回 ok。",
            "重新编译小程序后，如仍命中旧环境，可清理本地存储并重新进入页面；当前代码已默认在开发态自动回本地接口。",
            "如微信开发者工具报网络或 WebSocket 相关错误，请先确认本机代理、Docker 和开发者工具端口状态。",
        ],
    )

    add_paragraph(doc, "9 联系方式与技术支持", style="Heading 1")
    add_table(
        doc,
        ["问题类型", "建议处理方式"],
        [
            ["访问与部署问题", "优先检查 /healthz；若仍异常，联系当前部署维护方确认服务、网络和数据库状态。"],
            ["账号与权限问题", "联系系统管理员核对角色、scope 和微信绑定状态。"],
            ["业务口径问题", "联系学院业务老师或按正式通知、流程说明、模板原文再次确认。"],
            ["功能缺陷反馈", "提交问题时附上账号、操作路径、输入内容、实际结果、期望结果和截图，便于复现。"],
        ],
        [1.7, 4.8],
    )
    return finalize(doc, output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build project user manual DOCX.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    print(build_doc(args.output))


if __name__ == "__main__":
    main()
